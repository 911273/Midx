# -*- coding: utf-8 -*-
"""
Module quản lý cơ sở dữ liệu SQLite cho phần mềm EPU Exam Manager.
- Lưu trữ ngân hàng câu hỏi (text + ảnh/blob + công thức)
- Lưu đợt trộn đề & ánh xạ câu hỏi
- Lưu kết quả chấm điểm & chi tiết từng sinh viên
- Phát hiện trùng lặp nội dung & gợi ý đồng bộ
"""

import os
import json
import base64
import hashlib
import sqlite3
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from contextlib import contextmanager
import utils

# ============= CẤU HÌNH =============
import sys
if getattr(sys, 'frozen', False):
    # Nếu chạy từ file EXE đóng gói
    BASE_DIR = os.path.dirname(sys.executable)
else:
    # Nếu chạy từ code nguồn
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

DB_FILE = os.path.join(BASE_DIR, "epu_exam.db")


# ============= KẾT NỐI =============
@contextmanager
def get_connection(db_path: str = None):
    """Context manager cho kết nối DB. Tự động commit/rollback."""
    conn = sqlite3.connect(db_path or DB_FILE)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


# ============= KHỞI TẠO SCHEMA =============
def init_db(db_path: str = None):
    """Tạo hoặc nâng cấp schema cơ sở dữ liệu."""
    with get_connection(db_path) as conn:
        conn.executescript("""
        -- Môn học
        CREATE TABLE IF NOT EXISTS subjects (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            code TEXT DEFAULT ''
        );

        -- Ngân hàng câu hỏi (file Word)
        CREATE TABLE IF NOT EXISTS question_banks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            subject_id INTEGER,
            file_name TEXT NOT NULL,
            file_path TEXT NOT NULL,
            file_hash TEXT DEFAULT '',
            file_modified_at TEXT DEFAULT '',
            total_questions INTEGER DEFAULT 0,
            imported_at TEXT DEFAULT (datetime('now', 'localtime')),
            FOREIGN KEY (subject_id) REFERENCES subjects(id) ON DELETE SET NULL
        );

        -- Câu hỏi
        CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            bank_id INTEGER NOT NULL,
            qid_in_file INTEGER DEFAULT 0,
            stem_text TEXT DEFAULT '',
            stem_spans_json TEXT DEFAULT '[]',
            stem_extra_spans_json TEXT DEFAULT '[]',
            diff_code TEXT DEFAULT '',
            chapter TEXT DEFAULT '',
            topic TEXT DEFAULT '',
            content_hash TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now', 'localtime')),
            FOREIGN KEY (bank_id) REFERENCES question_banks(id) ON DELETE CASCADE
        );


        CREATE TABLE IF NOT EXISTS question_options (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question_id INTEGER NOT NULL,
            label TEXT NOT NULL DEFAULT '',
            option_text TEXT DEFAULT '',
            option_spans_json TEXT DEFAULT '[]',
            is_correct INTEGER DEFAULT 0,
            FOREIGN KEY (question_id) REFERENCES questions(id) ON DELETE CASCADE
        );

        -- Index cho hiệu năng tìm kiếm
        CREATE INDEX IF NOT EXISTS idx_questions_bank ON questions(bank_id);
        CREATE INDEX IF NOT EXISTS idx_questions_diff ON questions(diff_code);
        CREATE INDEX IF NOT EXISTS idx_questions_ch ON questions(chapter);
        CREATE INDEX IF NOT EXISTS idx_q_options_qid ON question_options(question_id);

        CREATE TABLE IF NOT EXISTS exam_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            subject_id INTEGER,
            school TEXT DEFAULT '',
            faculty TEXT DEFAULT '',
            school_year TEXT DEFAULT '',
            semester TEXT DEFAULT '',
            class_name TEXT DEFAULT '',
            exam_title TEXT DEFAULT '',
            subject_name TEXT DEFAULT '',
            duration TEXT DEFAULT '',
            num_variants INTEGER DEFAULT 0,
            num_questions INTEGER DEFAULT 0,
            folder_path TEXT DEFAULT '',
            strategy TEXT DEFAULT '',
            shuffle_answers INTEGER DEFAULT 1,
            google_form_link TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now', 'localtime')),
            metadata_json TEXT DEFAULT '{}',
            FOREIGN KEY (subject_id) REFERENCES subjects(id) ON DELETE SET NULL
        );

        -- Mã đề (variant)
        CREATE TABLE IF NOT EXISTS exam_variants (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id INTEGER NOT NULL,
            exam_code INTEGER NOT NULL,
            FOREIGN KEY (session_id) REFERENCES exam_sessions(id) ON DELETE CASCADE
        );

        -- Câu hỏi trong mã đề
        CREATE TABLE IF NOT EXISTS variant_questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            variant_id INTEGER NOT NULL,
            question_id INTEGER,
            position INTEGER DEFAULT 0,
            correct_label TEXT DEFAULT '',
            FOREIGN KEY (variant_id) REFERENCES exam_variants(id) ON DELETE CASCADE,
            FOREIGN KEY (question_id) REFERENCES questions(id) ON DELETE SET NULL
        );

        -- Kết quả chấm điểm
        CREATE TABLE IF NOT EXISTS grading_results (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id INTEGER,
            student_name TEXT DEFAULT '',
            student_id TEXT DEFAULT '',
            exam_code INTEGER DEFAULT 0,
            score REAL DEFAULT 0.0,
            num_correct INTEGER DEFAULT 0,
            num_wrong INTEGER DEFAULT 0,
            total_questions INTEGER DEFAULT 0,
            school_year TEXT DEFAULT '',
            semester TEXT DEFAULT '',
            subject TEXT DEFAULT '',
            class_name TEXT DEFAULT '',
            graded_at TEXT DEFAULT (datetime('now', 'localtime')),
            result_file TEXT DEFAULT '',
            FOREIGN KEY (session_id) REFERENCES exam_sessions(id) ON DELETE SET NULL
        );

        -- Chi tiết câu trả lời
        CREATE TABLE IF NOT EXISTS student_answers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            result_id INTEGER NOT NULL,
            question_num INTEGER DEFAULT 0,
            student_answer TEXT DEFAULT '',
            correct_answer TEXT DEFAULT '',
            is_correct INTEGER DEFAULT 0,
            FOREIGN KEY (result_id) REFERENCES grading_results(id) ON DELETE CASCADE
        );

        -- Index tối ưu tìm kiếm
        CREATE INDEX IF NOT EXISTS idx_questions_bank ON questions(bank_id);
        CREATE INDEX IF NOT EXISTS idx_questions_hash ON questions(content_hash);
        CREATE INDEX IF NOT EXISTS idx_questions_diff ON questions(diff_code);
        CREATE INDEX IF NOT EXISTS idx_options_question ON question_options(question_id);
        CREATE INDEX IF NOT EXISTS idx_variants_session ON exam_variants(session_id);
        CREATE INDEX IF NOT EXISTS idx_vq_variant ON variant_questions(variant_id);
        CREATE INDEX IF NOT EXISTS idx_grading_session ON grading_results(session_id);
        CREATE INDEX IF NOT EXISTS idx_student_ans_result ON student_answers(result_id);
        CREATE INDEX IF NOT EXISTS idx_grading_year ON grading_results(school_year);
        CREATE INDEX IF NOT EXISTS idx_grading_subject ON grading_results(subject);
        CREATE INDEX IF NOT EXISTS idx_grading_class ON grading_results(class_name);
        CREATE INDEX IF NOT EXISTS idx_grading_class ON grading_results(class_name);


        """)

    # Migration: Thêm cột google_form_link vào bảng exam_sessions nếu chưa có
    try:
        with get_connection(db_path) as conn:
            conn.execute("ALTER TABLE exam_sessions ADD COLUMN google_form_link TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass # Cột đã tồn tại

    # Migration: Thêm cột chapter và topic vào bảng questions nếu chưa có
    for col in ["chapter", "topic"]:
        try:
            with get_connection(db_path) as conn:
                conn.execute(f"ALTER TABLE questions ADD COLUMN {col} TEXT DEFAULT ''")
        except sqlite3.OperationalError:
            pass

    # Sau khi đã có cột, tạo Index
    try:
        with get_connection(db_path) as conn:
            conn.execute("CREATE INDEX IF NOT EXISTS idx_questions_chapter ON questions(chapter)")
            conn.execute("CREATE INDEX IF NOT EXISTS idx_questions_topic ON questions(topic)")
    except sqlite3.OperationalError:
        pass




# ============= SERIALIZE / DESERIALIZE SPANS =============
def _serialize_spans(spans: List[Dict[str, Any]]) -> str:
    """Serialize danh sách spans (text/image/omml) thành JSON string.
    Blob ảnh → base64. OMML XML → string representation.
    """
    if not spans:
        return "[]"
    serialized = []
    for sp in spans:
        entry = {"type": sp.get("type", "text")}
        if sp["type"] == "text":
            entry["text"] = sp.get("text", "")
            entry["bold"] = sp.get("bold", False)
            entry["italic"] = sp.get("italic", False)
            entry["underline"] = sp.get("underline", False)
            entry["color_hex"] = sp.get("color_hex")
        elif sp["type"] == "image":
            blob = sp.get("blob", b"")
            entry["blob_b64"] = base64.b64encode(blob).decode("ascii") if blob else ""
        elif sp["type"] == "omml":
            # Serialize XML element thành string
            try:
                from lxml import etree
                entry["xml_str"] = etree.tostring(sp["xml"], encoding="unicode")
            except Exception:
                try:
                    import xml.etree.ElementTree as ET
                    entry["xml_str"] = ET.tostring(sp["xml"], encoding="unicode")
                except Exception:
                    entry["xml_str"] = ""
        serialized.append(entry)
    return json.dumps(serialized, ensure_ascii=False)


def _deserialize_spans(json_str: str) -> List[Dict[str, Any]]:
    """Deserialize JSON string thành danh sách spans."""
    if not json_str or json_str == "[]":
        return []
    try:
        data = json.loads(json_str)
    except (json.JSONDecodeError, TypeError):
        return []

    spans = []
    for entry in data:
        sp = {"type": entry.get("type", "text")}
        if sp["type"] == "text":
            sp["text"] = entry.get("text", "")
            sp["bold"] = entry.get("bold", False)
            sp["italic"] = entry.get("italic", False)
            sp["underline"] = entry.get("underline", False)
            sp["color_hex"] = entry.get("color_hex")
            sp["strip_black"] = False
        elif sp["type"] == "image":
            b64 = entry.get("blob_b64", "")
            sp["blob"] = base64.b64decode(b64) if b64 else b""
        elif sp["type"] == "omml":
            xml_str = entry.get("xml_str", "")
            if xml_str:
                try:
                    from lxml import etree
                    sp["xml"] = etree.fromstring(xml_str)
                except Exception:
                    try:
                        import xml.etree.ElementTree as ET
                        sp["xml"] = ET.fromstring(xml_str)
                    except Exception:
                        continue
            else:
                continue
        spans.append(sp)
    return spans


# ============= HASH NỘI DUNG =============
def _compute_content_hash(stem_text: str, options_texts: List[str]) -> str:
    """Tính hash SHA256 từ nội dung câu hỏi để phát hiện trùng lặp."""
    content = stem_text.strip().lower()
    for opt in sorted(options_texts):
        content += "|" + opt.strip().lower()
    return hashlib.sha256(content.encode("utf-8")).hexdigest()[:32]


def _compute_file_hash(file_path: str) -> str:
    """Tính hash SHA256 của file để phát hiện thay đổi."""
    if not os.path.exists(file_path):
        return ""
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        while True:
            chunk = f.read(8192)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()[:32]


# ============= QUẢN LÝ MÔN HỌC =============
def add_subject(name: str, code: str = "") -> int:
    """Thêm môn học. Nếu đã tồn tại thì trả về ID cũ."""
    with get_connection() as conn:
        row = conn.execute("SELECT id FROM subjects WHERE name = ?", (name,)).fetchone()
        if row:
            if code and not row.get("code"):
                conn.execute("UPDATE subjects SET code = ? WHERE id = ?", (code, row["id"]))
            return row["id"]
        cur = conn.execute("INSERT INTO subjects (name, code) VALUES (?, ?)", (name, code))
        return cur.lastrowid


def get_subjects() -> List[Dict]:
    """Lấy danh sách tất cả môn học."""
    with get_connection() as conn:
        rows = conn.execute("SELECT * FROM subjects ORDER BY name").fetchall()
        return [dict(r) for r in rows]


def delete_subject(subject_id: int):
    """Xóa môn học."""
    with get_connection() as conn:
        conn.execute("DELETE FROM subjects WHERE id = ?", (subject_id,))


def get_banks_by_subject(subject_id: int) -> List[Dict]:
    """Lấy danh sách ngân hàng thuộc một môn học."""
    with get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM question_banks WHERE subject_id = ? ORDER BY file_name",
            (subject_id,)
        ).fetchall()
        return [dict(r) for r in rows]


# ============= IMPORT NGÂN HÀNG CÂU HỎI =============
def import_bank_to_db(file_path: str, questions: List[Dict], subject_name: str = "") -> Dict[str, Any]:
    """
    Import câu hỏi từ tệp Word vào một ngân hàng mới.
    """
    file_name = os.path.basename(file_path)
    file_hash = _compute_file_hash(file_path)
    file_mtime = ""
    if os.path.exists(file_path):
        file_mtime = datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y-%m-%d %H:%M:%S")

    subject_id = None
    if subject_name:
        subject_id = add_subject(subject_name)

    with get_connection() as conn:
        # Kiểm tra xem file đã import chưa (theo file_hash)
        existing_bank = conn.execute(
            "SELECT id FROM question_banks WHERE file_hash = ? AND file_hash != ''",
            (file_hash,)
        ).fetchone()
        
        if existing_bank:
            return {
                "total_imported": 0,
                "duplicates_found": [],
                "bank_id": existing_bank["id"],
                "already_imported": True,
                "message": f"File '{file_name}' đã được import trước đó (nội dung chưa thay đổi)."
            }

        # Tạo bản ghi ngân hàng
        cur = conn.execute(
            """INSERT INTO question_banks 
               (subject_id, file_name, file_path, file_hash, file_modified_at, total_questions)
               VALUES (?, ?, ?, ?, ?, ?)""",
            (subject_id, file_name, file_path, file_hash, file_mtime, len(questions))
        )
        bank_id = cur.lastrowid
        
        imported_count, duplicates, duplicate_options_found = _add_questions_to_bank_internal(conn, bank_id, questions)

    return {
        "total_imported": imported_count,
        "duplicates_found": duplicates,
        "duplicate_options_found": duplicate_options_found,
        "bank_id": bank_id,
        "already_imported": False,
        "message": f"Đã import {imported_count} câu hỏi vào DB." +
                   (f" Phát hiện {len(duplicates)} câu trùng lặp." if duplicates else "")
    }


def append_questions_to_bank(bank_id: int, questions: List[Dict]) -> Dict[str, Any]:
    """
    Thêm câu hỏi vào một ngân hàng đã tồn tại.
    """
    with get_connection() as conn:
        imported_count, duplicates, duplicate_options_found = _add_questions_to_bank_internal(conn, bank_id, questions)
        
        # Cập nhật tổng số câu trong bảng ngân hàng
        conn.execute(
            "UPDATE question_banks SET total_questions = total_questions + ? WHERE id = ?",
            (imported_count, bank_id)
        )

    return {
        "total_imported": imported_count,
        "duplicates_found": duplicates,
        "duplicate_options_found": duplicate_options_found,
        "bank_id": bank_id,
        "message": f"Đã thêm {imported_count} câu hỏi vào ngân hàng ID {bank_id}." +
                   (f" Phát hiện {len(duplicates)} câu trùng lặp." if duplicates else "")
    }


def _add_questions_to_bank_internal(conn, bank_id: int, questions: List[Dict]) -> Tuple[int, List[Dict]]:
    """
    Logic dùng chung để lưu danh sách câu hỏi vào một bank_id.
    """
    imported_count = 0
    duplicates = []
    duplicate_options_found = [] # Câu hỏi có các đáp án trùng nhau
    
    # Lấy qid_in_file lớn nhất hiện có trong bank này để tiếp nối
    row = conn.execute("SELECT MAX(qid_in_file) as m FROM questions WHERE bank_id = ?", (bank_id,)).fetchone()
    current_max_qid = row["m"] or 0

    # Lấy tất cả content_hash hiện có trong DB để kiểm tra trùng lặp
    existing_hashes = set()
    for row in conn.execute("SELECT content_hash FROM questions WHERE content_hash != ''"):
        existing_hashes.add(row["content_hash"])

    letters = ["A", "B", "C", "D"]

    for q in questions:
        current_max_qid += 1

        # Sử dụng hàm tóm tắt từ utils để băm nội dung (bao gồm cả mô phỏng công thức)
        def _get_full_content(spans):
            return utils.get_spans_text_summary(spans)

        # Serialize stem spans
        stem_spans_json = _serialize_spans(q.get("stem_media_spans", []))
        stem_extra_json = _serialize_spans_list(q.get("stem_extra_media_spans", []))

        # Tính content hash
        stem_text = q.get("stem_text", "")
        # Dùng nội dung đầy đủ (text + math) để nhận diện trùng lặp chính xác hơn
        full_stem_for_hash = _get_full_content(q.get("stem_media_spans", [])) or stem_text
        
        opts_full_contents = []
        for opt in q.get("options", []):
            spans = opt.get("info", {}).get("spans", [])
            opts_full_contents.append(_get_full_content(spans))

        content_hash = _compute_content_hash(full_stem_for_hash, opts_full_contents)

        # Kiểm tra đáp án trùng lặp trong chính câu hỏi này
        opts_clean = [t.strip().lower() for t in opts_full_contents if t.strip()]
        if len(opts_clean) != len(set(opts_clean)):
            duplicate_options_found.append({
                "qid": q.get("qid"),
                "stem_text": stem_text[:80]
            })

        # Kiểm tra trùng lặp
        is_duplicate = content_hash in existing_hashes
        if is_duplicate:
            dup_row = conn.execute(
                """SELECT q.id, q.stem_text, b.file_name 
                   FROM questions q JOIN question_banks b ON q.bank_id = b.id
                   WHERE q.content_hash = ?""",
                (content_hash,)
            ).fetchone()
            duplicates.append({
                "qid": q.get("qid"),
                "stem_text": stem_text[:80],
                "existing_file": dup_row["file_name"] if dup_row else "?"
            })

        # Insert câu hỏi
        cur_q = conn.execute(
            """INSERT INTO questions 
               (bank_id, qid_in_file, stem_text, stem_spans_json, stem_extra_spans_json, 
                diff_code, chapter, topic, content_hash)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (bank_id, current_max_qid, stem_text, stem_spans_json, stem_extra_json,
             q.get("diff_code", ""), q.get("chapter", ""), q.get("topic", ""), content_hash)
        )

        q_id = cur_q.lastrowid
        existing_hashes.add(content_hash)

        # Lưu phương án
        correct_idx = q.get("correct_index", -1)
        for i, opt in enumerate(q.get("options", [])[:4]):
            info = opt.get("info", {})
            opt_spans_json = _serialize_spans(info.get("spans", []))
            opt_text = "".join(sp.get("text", "") for sp in info.get("spans", []) if sp.get("type") == "text")
            label = letters[i] if i < len(letters) else str(i)

            conn.execute(
                """INSERT INTO question_options 
                   (question_id, label, option_text, option_spans_json, is_correct)
                   VALUES (?, ?, ?, ?, ?)""",
                (q_id, label, opt_text.strip(), opt_spans_json, 1 if i == correct_idx else 0)
            )
        imported_count += 1
        
    return imported_count, duplicates, duplicate_options_found


def _serialize_spans_list(spans_list: List[List[Dict]]) -> str:
    """Serialize danh sách các danh sách spans (stem_extra_media_spans)."""
    if not spans_list:
        return "[]"
    result = []
    for spans in spans_list:
        result.append(json.loads(_serialize_spans(spans)))
    return json.dumps(result, ensure_ascii=False)


def _deserialize_spans_list(json_str: str) -> List[List[Dict]]:
    """Deserialize danh sách các danh sách spans."""
    if not json_str or json_str == "[]":
        return []
    try:
        data = json.loads(json_str)
        return [_deserialize_spans(json.dumps(item)) for item in data]
    except Exception:
        return []


# ============= KIỂM TRA ĐỒNG BỘ =============
def check_sync_status(file_path: str) -> Dict[str, Any]:
    """
    Kiểm tra trạng thái đồng bộ giữa file Word và DB.
    
    Returns:
        Dict: status ('not_imported', 'up_to_date', 'modified'), bank_id, details
    """
    file_name = os.path.basename(file_path)
    current_hash = _compute_file_hash(file_path)

    with get_connection() as conn:
        # Tìm theo đường dẫn hoặc tên file
        row = conn.execute(
            "SELECT * FROM question_banks WHERE file_name = ? ORDER BY imported_at DESC LIMIT 1",
            (file_name,)
        ).fetchone()

        if not row:
            return {"status": "not_imported", "bank_id": None, "message": "Chưa import vào DB"}

        bank = dict(row)
        if bank["file_hash"] == current_hash:
            return {
                "status": "up_to_date",
                "bank_id": bank["id"],
                "message": f"Đã đồng bộ ({bank['total_questions']} câu, import lúc {bank['imported_at']})"
            }
        else:
            return {
                "status": "modified",
                "bank_id": bank["id"],
                "message": f"File đã thay đổi từ lần import cuối ({bank['imported_at']}). Nên import lại!"
            }


# ============= TRA CỨU CÂU HỎI =============
def search_questions(keyword: str = "", diff_code: str = "", subject_name: str = "",
                     bank_id: int = None, chapter: str = "", topic: str = "", limit: int = 100) -> List[Dict]:
    """Tìm kiếm câu hỏi trong DB."""
    conditions = []
    params = []

    if keyword:
        conditions.append("q.stem_text LIKE ?")
        params.append(f"%{keyword}%")
    if diff_code:
        conditions.append("q.diff_code = ?")
        params.append(diff_code)
    if chapter:
        conditions.append("q.chapter = ?")
        params.append(chapter)
    if topic:
        conditions.append("q.topic LIKE ?")
        params.append(f"%{topic}%")
    if subject_name:
        conditions.append("s.name LIKE ?")
        params.append(f"%{subject_name}%")
    if bank_id:
        conditions.append("q.bank_id = ?")
        params.append(bank_id)

    where_clause = " AND ".join(conditions) if conditions else "1=1"

    with get_connection() as conn:
        sql = f"""
        SELECT q.*, b.file_name, COALESCE(s.name, '') as subject_name
        FROM questions q
        JOIN question_banks b ON q.bank_id = b.id
        LEFT JOIN subjects s ON b.subject_id = s.id
        WHERE {where_clause}
        ORDER BY q.id DESC
        LIMIT ?
        """

        params.append(limit)
        rows = conn.execute(sql, params).fetchall()

        results = []
        if not rows:
            return results

        # Tối ưu hóa: Lấy toàn bộ options của danh sách câu hỏi trong 1 lần query (Tránh N+1)
        q_ids = [r["id"] for r in rows]
        placeholders = ",".join(["?"] * len(q_ids))
        all_opts = conn.execute(
            f"SELECT * FROM question_options WHERE question_id IN ({placeholders}) ORDER BY question_id, label",
            q_ids
        ).fetchall()

        # Group options theo question_id
        opts_map = {}
        for opt in all_opts:
            qid = opt["question_id"]
            if qid not in opts_map:
                opts_map[qid] = []
            opts_map[qid].append(dict(opt))

        for r in rows:
            q_dict = dict(r)
            q_dict["options"] = opts_map.get(q_dict["id"], [])
            results.append(q_dict)

        return results


def get_bank_chapters_and_topics(bank_id: int = None) -> Dict[str, List[str]]:
    """Lấy danh sách các Chương và Chủ đề hiện có trong ngân hàng (hoặc toàn DB)."""
    where_clause = "WHERE bank_id = ?" if bank_id else ""
    params = (bank_id,) if bank_id else ()
    
    with get_connection() as conn:
        # Lấy Chapters
        ch_rows = conn.execute(
            f"SELECT DISTINCT chapter FROM questions {where_clause} ORDER BY chapter",
            params
        ).fetchall()
        chapters = [r["chapter"] for r in ch_rows if r["chapter"]]
        
        # Lấy Topics
        tp_rows = conn.execute(
            f"SELECT DISTINCT topic FROM questions {where_clause} ORDER BY topic",
            params
        ).fetchall()
        topics = [r["topic"] for r in tp_rows if r["topic"]]
        
        return {"chapters": chapters, "topics": topics}


def get_question_with_spans(question_id: int) -> Optional[Dict]:
    """Lấy câu hỏi đầy đủ với spans đã deserialize (để render/trộn đề)."""
    with get_connection() as conn:
        row = conn.execute("SELECT * FROM questions WHERE id = ?", (question_id,)).fetchone()
        if not row:
            return None

        q = dict(row)
        q["stem_media_spans"] = _deserialize_spans(q.get("stem_spans_json", "[]"))
        q["stem_extra_media_spans"] = _deserialize_spans_list(q.get("stem_extra_spans_json", "[]"))

        opts = conn.execute(
            "SELECT * FROM question_options WHERE question_id = ? ORDER BY label",
            (question_id,)
        ).fetchall()

        q["options"] = []
        q["correct_index"] = 0
        for i, opt in enumerate(opts):
            opt_dict = dict(opt)
            opt_dict["info"] = {"spans": _deserialize_spans(opt_dict.get("option_spans_json", "[]"))}
            if opt_dict.get("is_correct"):
                q["correct_index"] = i
            q["options"].append(opt_dict)


def batch_update_questions(question_ids, chapter=None, diff_code=None, topic=None):
    """Cập nhật hàng loạt các trường thông tin cho danh sách câu hỏi."""
    if not question_ids:
        return
    
    updates = []
    params = []
    if chapter is not None:
        updates.append("chapter = ?")
        params.append(chapter)
    if diff_code is not None:
        updates.append("diff_code = ?")
        params.append(diff_code)
    if topic is not None:
        updates.append("topic = ?")
        params.append(topic)
    
    if not updates:
        return

    # Chuẩn bị SQL động
    sql = f"UPDATE questions SET {', '.join(updates)} WHERE id IN ({','.join(['?']*len(question_ids))})"
    params.extend(question_ids)
    
    with get_connection() as conn:
        conn.execute(sql, params)


def get_bank_stats(bank_id: int = None) -> Dict[str, Any]:
    """Thống kê ngân hàng câu hỏi."""
    with get_connection() as conn:
        if bank_id:
            total = conn.execute("SELECT COUNT(*) as c FROM questions WHERE bank_id = ?", (bank_id,)).fetchone()["c"]
            diff_stats = conn.execute(
                "SELECT diff_code, COUNT(*) as c FROM questions WHERE bank_id = ? GROUP BY diff_code",
                (bank_id,)
            ).fetchall()
        else:
            total = conn.execute("SELECT COUNT(*) as c FROM questions").fetchone()["c"]
            diff_stats = conn.execute(
                "SELECT diff_code, COUNT(*) as c FROM questions GROUP BY diff_code"
            ).fetchall()

        banks_count = conn.execute("SELECT COUNT(*) as c FROM question_banks").fetchone()["c"]
        subjects_count = conn.execute("SELECT COUNT(*) as c FROM subjects").fetchone()["c"]

        return {
            "total_questions": total,
            "total_banks": banks_count,
            "total_subjects": subjects_count,
            "by_difficulty": {r["diff_code"] or "N/A": r["c"] for r in diff_stats}
        }


def get_all_banks() -> List[Dict]:
    """Lấy danh sách tất cả ngân hàng đã import."""
    with get_connection() as conn:
        rows = conn.execute("""
            SELECT b.*, COALESCE(s.name, '') as subject_name 
            FROM question_banks b 
            LEFT JOIN subjects s ON b.subject_id = s.id
            ORDER BY b.imported_at DESC
        """).fetchall()
        return [dict(r) for r in rows]


# ============= QUẢN LÝ ĐỢT TRỘN ĐỀ =============
def save_exam_session(school: str, faculty: str, school_year: str, semester: str,
                      class_name: str, exam_title: str, subject_name: str, duration: str,
                      num_variants: int, num_questions: int, folder_path: str,
                      strategy: str, shuffle_answers: bool,
                      variants_data: List[Dict] = None,
                      metadata_json: str = "{}") -> int:
    """
    Lưu đợt trộn đề vào DB.
    
    Args:
        variants_data: List of {exam_code, questions: [{question_db_id, position, correct_label}]}
    
    Returns:
        session_id
    """
    subject_id = add_subject(subject_name) if subject_name else None

    with get_connection() as conn:
        cur = conn.execute(
            """INSERT INTO exam_sessions 
               (subject_id, school, faculty, school_year, semester, class_name, exam_title,
                subject_name, duration, num_variants, num_questions, folder_path, strategy,
                shuffle_answers, metadata_json)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (subject_id, school, faculty, school_year, semester, class_name, exam_title,
             subject_name, duration, num_variants, num_questions, folder_path, strategy,
             1 if shuffle_answers else 0, metadata_json)
        )
        session_id = cur.lastrowid

        # Lưu từng mã đề và câu hỏi
        if variants_data:
            for v in variants_data:
                cur_v = conn.execute(
                    "INSERT INTO exam_variants (session_id, exam_code) VALUES (?, ?)",
                    (session_id, v["exam_code"])
                )
                variant_id = cur_v.lastrowid

                for vq in v.get("questions", []):
                    conn.execute(
                        """INSERT INTO variant_questions 
                           (variant_id, question_id, position, correct_label)
                           VALUES (?, ?, ?, ?)""",
                        (variant_id, vq.get("question_db_id"), vq.get("position", 0),
                         vq.get("correct_label", ""))
                    )

    return session_id


def get_exam_sessions(school_year: str = "", semester: str = "",
                      subject: str = "", class_name: str = "") -> List[Dict]:
    """Lấy danh sách đợt trộn đề, có thể lọc theo năm/kỳ/môn/lớp."""
    conditions = []
    params = []
    if school_year:
        conditions.append("school_year LIKE ?")
        params.append(f"%{school_year}%")
    if semester:
        conditions.append("semester = ?")
        params.append(semester)
    if subject:
        conditions.append("subject_name LIKE ?")
        params.append(f"%{subject}%")
    if class_name:
        conditions.append("class_name LIKE ?")
        params.append(f"%{class_name}%")

    where = " AND ".join(conditions) if conditions else "1=1"
    with get_connection() as conn:
        rows = conn.execute(
            f"SELECT * FROM exam_sessions WHERE {where} ORDER BY created_at DESC", params
        ).fetchall()
        return [dict(r) for r in rows]


def delete_exam_session(session_id: int):
    """Xóa đợt trộn đề và dữ liệu liên quan."""
    with get_connection() as conn:
        conn.execute("DELETE FROM exam_sessions WHERE id = ?", (session_id,))


def update_exam_session_link(session_id: int, link: str):
    """Cập nhật link Google Form cho một đợt trộn đề."""
    with get_connection() as conn:
        conn.execute(
            "UPDATE exam_sessions SET google_form_link = ? WHERE id = ?",
            (link, session_id)
        )


# ============= CHẤM ĐIỂM =============
def save_grading_results(session_id: int, school_year: str, semester: str,
                         subject: str, class_name: str,
                         results: List[Dict], result_file: str = "") -> int:
    """
    Lưu kết quả chấm điểm cho nhiều sinh viên.
    
    Args:
        results: List of {
            student_name, student_id, exam_code, score, num_correct, num_wrong,
            total_questions, answers: [{question_num, student_answer, correct_answer, is_correct}]
        }
    
    Returns:
        Số bản ghi đã lưu
    """
    saved = 0
    with get_connection() as conn:
        for r in results:
            cur = conn.execute(
                """INSERT INTO grading_results 
                   (session_id, student_name, student_id, exam_code, score,
                    num_correct, num_wrong, total_questions,
                    school_year, semester, subject, class_name, result_file)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (session_id, r.get("student_name", ""), r.get("student_id", ""),
                 r.get("exam_code", 0), r.get("score", 0.0),
                 r.get("num_correct", 0), r.get("num_wrong", 0),
                 r.get("total_questions", 0),
                 school_year, semester, subject, class_name, result_file)
            )
            result_id = cur.lastrowid

            for ans in r.get("answers", []):
                conn.execute(
                    """INSERT INTO student_answers 
                       (result_id, question_num, student_answer, correct_answer, is_correct)
                       VALUES (?, ?, ?, ?, ?)""",
                    (result_id, ans.get("question_num", 0),
                     ans.get("student_answer", ""), ans.get("correct_answer", ""),
                     1 if ans.get("is_correct") else 0)
                )
            saved += 1
    return saved


def get_grading_history(school_year: str = "", semester: str = "",
                        subject: str = "", class_name: str = "",
                        limit: int = 200) -> List[Dict]:
    """Lấy lịch sử chấm điểm, có thể lọc theo năm/kỳ/môn/lớp."""
    conditions = []
    params = []
    if school_year:
        conditions.append("school_year LIKE ?")
        params.append(f"%{school_year}%")
    if semester:
        conditions.append("semester = ?")
        params.append(semester)
    if subject:
        conditions.append("subject LIKE ?")
        params.append(f"%{subject}%")
    if class_name:
        conditions.append("class_name LIKE ?")
        params.append(f"%{class_name}%")

    where = " AND ".join(conditions) if conditions else "1=1"
    with get_connection() as conn:
        rows = conn.execute(
            f"SELECT * FROM grading_results WHERE {where} ORDER BY graded_at DESC LIMIT ?",
            params + [limit]
        ).fetchall()
        return [dict(r) for r in rows]


def get_grading_summary(school_year: str = "", semester: str = "",
                        subject: str = "", class_name: str = "") -> Dict[str, Any]:
    """Thống kê tổng quan kết quả chấm điểm."""
    conditions = []
    params = []
    if school_year:
        conditions.append("school_year LIKE ?")
        params.append(f"%{school_year}%")
    if semester:
        conditions.append("semester = ?")
        params.append(semester)
    if subject:
        conditions.append("subject LIKE ?")
        params.append(f"%{subject}%")
    if class_name:
        conditions.append("class_name LIKE ?")
        params.append(f"%{class_name}%")

    where = " AND ".join(conditions) if conditions else "1=1"
    with get_connection() as conn:
        row = conn.execute(
            f"""SELECT COUNT(*) as total_students,
                       AVG(score) as avg_score,
                       MAX(score) as max_score,
                       MIN(score) as min_score,
                       COUNT(DISTINCT class_name) as total_classes,
                       COUNT(DISTINCT subject) as total_subjects
                FROM grading_results WHERE {where}""",
            params
        ).fetchone()
        return dict(row) if row else {}


def get_unique_grading_values() -> Dict[str, List[str]]:
    """Lấy danh sách các giá trị duy nhất để làm bộ lọc."""
    res = {
        "years": [],
        "semesters": [],
        "subjects": [],
        "classes": []
    }
    with get_connection() as conn:
        # Năm học
        rows = conn.execute("SELECT DISTINCT school_year FROM grading_results WHERE school_year != '' ORDER BY school_year DESC").fetchall()
        res["years"] = [r[0] for r in rows]
        
        # Học kỳ
        rows = conn.execute("SELECT DISTINCT semester FROM grading_results WHERE semester != '' ORDER BY semester").fetchall()
        res["semesters"] = [str(r[0]) for r in rows]
        
        # Môn học
        rows = conn.execute("SELECT DISTINCT subject FROM grading_results WHERE subject != '' ORDER BY subject").fetchall()
        res["subjects"] = [r[0] for r in rows]
        
        # Lớp
        rows = conn.execute("SELECT DISTINCT class_name FROM grading_results WHERE class_name != '' ORDER BY class_name").fetchall()
        res["classes"] = [r[0] for r in rows]
        
    return res



# ============= QUẢN LÝ NGÂN HÀNG THEO FILE =============
def get_bank_by_filename(file_name: str) -> Optional[Dict]:
    """Tìm ngân hàng theo tên file."""
    with get_connection() as conn:
        row = conn.execute(
            "SELECT * FROM question_banks WHERE file_name = ? ORDER BY imported_at DESC LIMIT 1",
            (file_name,)
        ).fetchone()
        return dict(row) if row else None


def delete_bank(bank_id: int):
    """Xóa ngân hàng và tất cả câu hỏi liên quan."""
    with get_connection() as conn:
        conn.execute("DELETE FROM question_banks WHERE id = ?", (bank_id,))


def reimport_bank(file_path: str, questions: List[Dict], subject_name: str = "") -> Dict[str, Any]:
    """Xóa bản ghi cũ và import lại file Word."""
    file_name = os.path.basename(file_path)
    with get_connection() as conn:
        # Xóa ngân hàng cũ theo tên file
        conn.execute("DELETE FROM question_banks WHERE file_name = ?", (file_name,))
    # Import mới
    return import_bank_to_db(file_path, questions, subject_name)


# ============= TÌM CÂU HỎI TRÙNG LẶP =============
def find_duplicates(bank_id: int = None) -> List[Dict]:
    """Tìm các câu hỏi trùng lặp nội dung trong DB."""
    with get_connection() as conn:
        if bank_id:
            sql = """
            SELECT q1.id as id1, q1.stem_text as text1, b1.file_name as file1,
                   q2.id as id2, q2.stem_text as text2, b2.file_name as file2
            FROM questions q1
            JOIN questions q2 ON q1.content_hash = q2.content_hash AND q1.id < q2.id
            JOIN question_banks b1 ON q1.bank_id = b1.id
            JOIN question_banks b2 ON q2.bank_id = b2.id
            WHERE q1.bank_id = ? OR q2.bank_id = ?
            ORDER BY q1.content_hash
            """
            rows = conn.execute(sql, (bank_id, bank_id)).fetchall()
        else:
            sql = """
            SELECT q1.id as id1, q1.stem_text as text1, b1.file_name as file1,
                   q2.id as id2, q2.stem_text as text2, b2.file_name as file2
            FROM questions q1
            JOIN questions q2 ON q1.content_hash = q2.content_hash AND q1.id < q2.id
            JOIN question_banks b1 ON q1.bank_id = b1.id
            JOIN question_banks b2 ON q2.bank_id = b2.id
            ORDER BY q1.content_hash
            """
            rows = conn.execute(sql).fetchall()
        return [dict(r) for r in rows]


# ============= SESSION ↔ FOLDER MAPPING =============
def find_session_by_folder(folder_path: str) -> Optional[Dict]:
    """Tìm exam session theo đường dẫn thư mục."""
    with get_connection() as conn:
        row = conn.execute(
            "SELECT * FROM exam_sessions WHERE folder_path = ?", (folder_path,)
        ).fetchone()
        return dict(row) if row else None


def delete_session_by_folder(folder_path: str):
    """Xóa exam session theo thư mục."""
    with get_connection() as conn:
        conn.execute("DELETE FROM exam_sessions WHERE folder_path = ?", (folder_path,))


def delete_question(question_id: int):
    """Xóa một câu hỏi và các phương án liên quan."""
    with get_connection() as conn:
        conn.execute("DELETE FROM questions WHERE id = ?", (question_id,))


# ============= SOẠN THẢO CÂU HỎI (MANUAL) =============
def save_question_manually(bank_id: int, stem_text: str, stem_media_spans: List[Dict],
                            stem_extra_media_spans: List[List[Dict]], diff_code: str,
                            options_data: List[Dict], question_id: Optional[int] = None) -> int:
    """
    Lưu hoặc cập nhật một câu hỏi soạn thảo tay.
    
    Args:
        options_data: List of {label, text, spans, is_correct}
        question_id: Nếu cung cấp, thực hiện UPDATE thay vì INSERT.
    """
    stem_spans_json = _serialize_spans(stem_media_spans)
    stem_extra_json = _serialize_spans_list(stem_extra_media_spans)
    
    # Tính content hash
    opts_texts = [opt.get("text", "") for opt in options_data[:4]]
    content_hash = _compute_content_hash(stem_text, opts_texts)

    with get_connection() as conn:
        if question_id:
            conn.execute(
                """UPDATE questions 
                   SET bank_id = ?, stem_text = ?, stem_spans_json = ?, 
                       stem_extra_spans_json = ?, diff_code = ?, content_hash = ?
                   WHERE id = ?""",
                (bank_id, stem_text, stem_spans_json, stem_extra_json, diff_code, content_hash, question_id)
            )
            # Xóa options cũ
            conn.execute("DELETE FROM question_options WHERE question_id = ?", (question_id,))
            q_id = question_id
        else:
            # Lấy qid_in_file lớn nhất hiện tại trong bank này
            row = conn.execute("SELECT MAX(qid_in_file) as m FROM questions WHERE bank_id = ?", (bank_id,)).fetchone()
            next_qid = (row["m"] or 0) + 1
            
            cur = conn.execute(
                """INSERT INTO questions 
                   (bank_id, qid_in_file, stem_text, stem_spans_json, stem_extra_spans_json, diff_code, content_hash)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (bank_id, next_qid, stem_text, stem_spans_json, stem_extra_json, diff_code, content_hash)
            )
            q_id = cur.lastrowid

        # Lưu options mới
        for opt in options_data[:4]:
            opt_spans_json = _serialize_spans(opt.get("spans", []))
            conn.execute(
                """INSERT INTO question_options 
                   (question_id, label, option_text, option_spans_json, is_correct)
                   VALUES (?, ?, ?, ?, ?)""",
                (q_id, opt.get("label", ""), opt.get("text", "").strip(), 
                 opt_spans_json, 1 if opt.get("is_correct") else 0)
            )
            
    return q_id


def create_virtual_bank(subject_id: int, bank_name: str = "NGÂN HÀNG TỰ SOẠN") -> int:
    """Tạo một ngân hàng ảo để lưu câu hỏi soạn thảo tay."""
    with get_connection() as conn:
        # Kiểm tra xem đã có ngân hàng cùng tên chưa
        row = conn.execute(
            "SELECT id FROM question_banks WHERE subject_id = ? AND file_name = ?",
            (subject_id, bank_name)
        ).fetchone()
        if row:
            return row["id"]
            
        cur = conn.execute(
            """INSERT INTO question_banks 
               (subject_id, file_name, file_path, file_hash, total_questions)
               VALUES (?, ?, ?, ?, ?)""",
            (subject_id, bank_name, "INTERNAL", "VIRTUAL", 0)
        )
        return cur.lastrowid


def rename_subject(subject_id: int, new_name: str):
    """Đổi tên môn học."""
    with get_connection() as conn:
        conn.execute("UPDATE subjects SET name = ? WHERE id = ?", (new_name, subject_id))


def move_questions(question_ids: List[int], target_bank_id: int):
    """Di chuyển danh sách câu hỏi sang ngân hàng khác."""
    with get_connection() as conn:
        # Lấy qid_in_file lớn nhất của bank đích
        row = conn.execute("SELECT MAX(qid_in_file) as m FROM questions WHERE bank_id = ?", (target_bank_id,)).fetchone()
        current_max = row["m"] or 0
        
        for qid in question_ids:
            current_max += 1
            conn.execute(
                "UPDATE questions SET bank_id = ?, qid_in_file = ? WHERE id = ?",
                (target_bank_id, current_max, qid)
            )
        
        _refresh_bank_counts(conn)


def copy_questions(question_ids: List[int], target_bank_id: int):
    """Sao chép danh sách câu hỏi sang ngân hàng khác."""
    with get_connection() as conn:
        row = conn.execute("SELECT MAX(qid_in_file) as m FROM questions WHERE bank_id = ?", (target_bank_id,)).fetchone()
        current_max = row["m"] or 0
        
        for qid in question_ids:
            current_max += 1
            # Lấy thông tin câu hỏi cũ
            q_row = conn.execute("SELECT * FROM questions WHERE id = ?", (qid,)).fetchone()
            if not q_row: continue
            q = dict(q_row)
            
            # Chèn câu hỏi mới
            cur = conn.execute(
                """INSERT INTO questions 
                   (bank_id, qid_in_file, stem_text, stem_spans_json, stem_extra_spans_json, diff_code, content_hash)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (target_bank_id, current_max, q["stem_text"], q["stem_spans_json"], 
                 q["stem_extra_spans_json"], q["diff_code"], q["content_hash"])
            )
            new_q_id = cur.lastrowid
            
            # Sao chép options
            opts_rows = conn.execute("SELECT * FROM question_options WHERE question_id = ?", (qid,)).fetchall()
            for opt in opts_rows:
                conn.execute(
                    """INSERT INTO question_options 
                       (question_id, label, option_text, option_spans_json, is_correct)
                       VALUES (?, ?, ?, ?, ?)""",
                    (new_q_id, opt["label"], opt["option_text"], opt["option_spans_json"], opt["is_correct"])
                )
        
        _refresh_bank_counts(conn)


def merge_banks(source_bank_ids: List[int], target_bank_id: int):
    """Gộp nhiều ngân hàng vào một ngân hàng đích."""
    with get_connection() as conn:
        # Lấy qid_in_file lớn nhất của bank đích
        row = conn.execute("SELECT MAX(qid_in_file) as m FROM questions WHERE bank_id = ?", (target_bank_id,)).fetchone()
        current_max = row["m"] or 0

        for src_id in source_bank_ids:
            if src_id == target_bank_id: continue
            
            # Lấy danh sách câu hỏi từ bank nguồn
            q_rows = conn.execute("SELECT id FROM questions WHERE bank_id = ? ORDER BY qid_in_file", (src_id,)).fetchall()
            for r in q_rows:
                current_max += 1
                conn.execute(
                    "UPDATE questions SET bank_id = ?, qid_in_file = ? WHERE id = ?",
                    (target_bank_id, current_max, r["id"])
                )
            
            # Xóa bank nguồn sau khi đã hết câu hỏi
            conn.execute("DELETE FROM question_banks WHERE id = ?", (src_id,))
        
        _refresh_bank_counts(conn)


def get_duplicate_conflicts() -> List[Dict]:
    """
    Tìm các nhóm câu hỏi trùng nội dung nhưng có mức độ khó khác nhau.
    Trả về: List of {content_hash, stem_preview, diff_codes: [], count}
    """
    with get_connection() as conn:
        sql = """
            SELECT content_hash, MIN(stem_text) as stem_preview, 
                   GROUP_CONCAT(DISTINCT COALESCE(diff_code, 'TB')) as diff_codes_str,
                   COUNT(*) as count
            FROM questions 
            WHERE content_hash != ''
            GROUP BY content_hash
            HAVING COUNT(DISTINCT diff_code) > 1
        """
        rows = conn.execute(sql).fetchall()
        results = []
        for r in rows:
            results.append({
                "content_hash": r["content_hash"],
                "stem_preview": r["stem_preview"][:100],
                "diff_codes": r["diff_codes_str"].split(","),
                "count": r["count"]
            })
        return results


def merge_duplicate_questions(conflict_resolutions: Dict[str, str] = None) -> int:
    """
    Tự động tìm và gộp các câu hỏi trùng content_hash.
    
    Args:
        conflict_resolutions: Dict mapping content_hash -> target_diff_code.
                             Nếu không cung cấp, sẽ giữ nguyên mức độ khó của câu được chọn làm mốc.
    """
    conflict_resolutions = conflict_resolutions or {}
    
    with get_connection() as conn:
        # Tìm các content_hash bị trùng
        dups = conn.execute("""
            SELECT content_hash, MIN(id) as keep_id, GROUP_CONCAT(id) as all_ids
            FROM questions 
            WHERE content_hash != ''
            GROUP BY content_hash
            HAVING COUNT(*) > 1
        """).fetchall()
        
        merged_count = 0
        for row in dups:
            chash = row["content_hash"]
            keep_id = row["keep_id"]
            all_ids = [int(x) for x in row["all_ids"].split(",") if int(x) != keep_id]
            
            # Cập nhật mức độ khó nếu có yêu cầu giải quyết xung đột
            if chash in conflict_resolutions:
                conn.execute(
                    "UPDATE questions SET diff_code = ? WHERE id = ?",
                    (conflict_resolutions[chash], keep_id)
                )
            
            for delete_id in all_ids:
                # Cập nhật variant_questions để trỏ về keep_id
                conn.execute(
                    "UPDATE variant_questions SET question_id = ? WHERE question_id = ?",
                    (keep_id, delete_id)
                )
                # Xóa câu hỏi trùng
                conn.execute("DELETE FROM questions WHERE id = ?", (delete_id,))
                merged_count += 1
        
        _refresh_bank_counts(conn)
        return merged_count


def _refresh_bank_counts(conn):
    """Cập nhật lại trường total_questions trong bảng question_banks."""
    conn.execute("""
        UPDATE question_banks 
        SET total_questions = (
            SELECT COUNT(*) FROM questions WHERE questions.bank_id = question_banks.id
        )
    """)


def recompute_all_hashes() -> int:
    """Cập nhật lại content_hash cho toàn bộ câu hỏi trong DB dựa trên logic mới."""
    count = 0
    with get_connection() as conn:
        rows = conn.execute("SELECT id, stem_spans_json, stem_text FROM questions").fetchall()
        for r in rows:
            q_id = r["id"]
            spans = _deserialize_spans(r["stem_spans_json"] or "[]")
            
            # Lấy options
            opt_rows = conn.execute("SELECT option_spans_json, option_text FROM question_options WHERE question_id = ?", (q_id,)).fetchall()
            
            # Tính lại full content for hashing
            full_stem = utils.get_spans_text_summary(spans) or (r["stem_text"] or "")
            opt_contents = []
            for o in opt_rows:
                o_spans = _deserialize_spans(o["option_spans_json"] or "[]")
                o_summary = utils.get_spans_text_summary(o_spans) or (o["option_text"] or "")
                opt_contents.append(o_summary)
                
                # Cập nhật tóm tắt phương án để SQL Audit hoạt động đúng
                conn.execute("UPDATE question_options SET option_text = ? WHERE question_id = ? AND option_spans_json = ?", 
                             (o_summary, q_id, o["option_spans_json"]))
                
            new_hash = _compute_content_hash(full_stem, opt_contents)
            
            # Cập nhật DB
            conn.execute("UPDATE questions SET content_hash = ?, stem_text = ? WHERE id = ?", (new_hash, full_stem, q_id))
            count += 1
    return count

# ============= TIỆN ÍCH HỆ THỐNG (BACKUP & EXPORT) =============
def backup_database(dest_path: str) -> bool:
    """Sao lưu file CSDL đến vị trí mới."""
    import shutil
    try:
        shutil.copy2(DB_FILE, dest_path)
        return True
    except Exception:
        return False


def restore_database(src_path: str) -> bool:
    """Khôi phục CSDL từ file backup."""
    import shutil
    try:
        # Cần đảm bảo không có kết nối nào đang mở nếu muốn ghi đè an toàn nhất
        # Nhưng thông qua context manager get_connection(), ta thường đóng ngay sau khi dùng.
        shutil.copy2(src_path, DB_FILE)
        return True
    except Exception:
        return False


def export_bank_to_docx(bank_id: int, output_path: str) -> bool:
    """Xuất nội dung ngân hàng câu hỏi từ DB ra file Word theo mẫu EPU."""
    from docx import Document
    from docx.shared import RGBColor
    import utils # Giả định utils.py nằm cùng thư mục
    
    try:
        with get_connection() as conn:
            bank_row = conn.execute("SELECT file_name FROM question_banks WHERE id = ?", (bank_id,)).fetchone()
            bank_name = bank_row["file_name"] if bank_row else "NGAN_HANG"
            
            # Lấy danh sách câu hỏi
            questions = search_questions(bank_id=bank_id, limit=5000)
            if not questions:
                return False
                
            doc = Document()
            utils.set_page_layout(doc)
            utils.set_single_line_spacing(doc)
            
            for i, q in enumerate(questions, 1):
                # Reload đầy đủ spans
                q_full = get_question_with_spans(q["id"])
                diff = q_full.get('diff_code') or "TB"
                
                p_stem = doc.add_paragraph()
                r_label = p_stem.add_run(f"Câu {i} [<{diff}>]: ")
                r_label.bold = True
                
                if q_full.get("stem_media_spans"):
                    utils.render_spans_into_paragraph(p_stem, q_full["stem_media_spans"], default_black=False)
                else:
                    p_stem.add_run(q_full.get("stem_text", ""))
                
                for extra in q_full.get("stem_extra_media_spans", []):
                    p_ex = doc.add_paragraph()
                    utils.render_spans_into_paragraph(p_ex, extra, default_black=False)
                
                # Options
                letters = ["A", "B", "C", "D"]
                for j, opt in enumerate(q_full.get("options", [])):
                    p_opt = doc.add_paragraph()
                    r_opt_lbl = p_opt.add_run(f"{letters[j]}. ")
                    if opt.get("is_correct"):
                        r_opt_lbl.font.color.rgb = RGBColor(0xFF, 0, 0)
                        r_opt_lbl.bold = True
                    
                    spans = opt.get("info", {}).get("spans", [])
                    if spans:
                        utils.render_spans_into_paragraph(p_opt, spans, default_black=False)
                    else:
                        p_opt.add_run(opt.get("option_text", ""))
                        
            doc.save(output_path)
            return True
    except Exception as e:
        print(f"Lỗi xuất Word: {e}")
        return False
def audit_all_questions() -> List[Dict]:
    """
    Thực hiện kiểm tra toàn diện CSDL và trả về danh sách các sự cố.
    """
    issues = []
    with get_connection() as conn:
        # 1. Chưa có đáp án đúng
        sql_no_correct = """
            SELECT q.id, q.stem_text, q.qid_in_file, b.file_name, s.name as subject_name
            FROM questions q
            JOIN question_banks b ON q.bank_id = b.id
            LEFT JOIN subjects s ON b.subject_id = s.id
            WHERE q.id NOT IN (SELECT question_id FROM question_options WHERE is_correct = 1)
        """
        for r in conn.execute(sql_no_correct).fetchall():
            issues.append({**dict(r), "type": "Thiếu Đ/A đúng", "detail": "Chưa chọn đáp án đúng."})

        # 2. Thiếu đáp án/Đáp án rỗng
        sql_incomplete = """
            SELECT q.id, q.stem_text, q.qid_in_file, b.file_name, s.name as subject_name
            FROM questions q
            JOIN question_banks b ON q.bank_id = b.id
            LEFT JOIN subjects s ON b.subject_id = s.id
            WHERE q.id IN (
                SELECT question_id FROM question_options GROUP BY question_id HAVING COUNT(*) < 4
                UNION
                SELECT question_id FROM question_options 
                WHERE (option_text IS NULL OR TRIM(option_text) = '') 
                AND (option_spans_json IS NULL OR option_spans_json = '[]' OR option_spans_json = '')
            )
        """
        for r in conn.execute(sql_incomplete).fetchall():
            # Tránh lặp nếu đã dính lỗi thiếu đáp án đúng bên trên
            if not any(iss["id"] == r["id"] and iss["type"] == "Thiếu đáp án" for iss in issues):
                issues.append({**dict(r), "type": "Lỗi SL đáp án", "detail": "Ít hơn 4 đáp án hoặc có đáp án rỗng."})

        # 3. Trùng lặp đáp án trong cùng 1 câu
        sql_dup_opts = """
            SELECT q.id, q.stem_text, q.qid_in_file, b.file_name, s.name as subject_name
            FROM questions q
            JOIN question_banks b ON q.bank_id = b.id
            LEFT JOIN subjects s ON b.subject_id = s.id
            WHERE q.id IN (
                SELECT question_id FROM question_options 
                GROUP BY question_id, LOWER(TRIM(option_text)), option_spans_json 
                HAVING COUNT(*) > 1
            )
        """
        for r in conn.execute(sql_dup_opts).fetchall():
            issues.append({**dict(r), "type": "Trùng đáp án", "detail": "Có các phương án trùng nội dung."})

        # 4. Trùng lặp nội dung câu hỏi (Hash trùng)
        sql_dup_hash = """
            SELECT q.id, q.stem_text, q.qid_in_file, b.file_name, s.name as subject_name
            FROM questions q
            JOIN question_banks b ON q.bank_id = b.id
            LEFT JOIN subjects s ON b.subject_id = s.id
            WHERE q.content_hash IN (
                SELECT content_hash FROM questions WHERE content_hash != '' GROUP BY content_hash HAVING COUNT(*) > 1
            )
        """
        for r in conn.execute(sql_dup_hash).fetchall():
            issues.append({**dict(r), "type": "Trùng câu hỏi", "detail": "Nội dung câu hỏi bị trùng với câu khác."})

    return issues


def get_grading_history_by_session(session_id: int) -> List[Dict]:
    """Lấy lịch sử chấm điểm của một session cụ thể."""
    with get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM grading_results WHERE session_id = ? ORDER BY student_id",
            (session_id,)
        ).fetchall()
        return [dict(r) for r in rows]
