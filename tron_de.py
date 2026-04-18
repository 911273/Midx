# -*- coding: utf-8 -*-
"""
Tab Trộn đề – bản tối ưu hiệu năng
- Precompile nội dung câu/đáp án: text/img/OMML thành spans để render nhanh
- Không còn duyệt XPath/runs lặp lại ở giai đoạn sinh đề
- Không tô font từng run (dựa vào style 'Normal' ⇒ Times New Roman), tăng tốc mạnh
- Trình bày: A4, lề 0.75/0.75/0.5/0.5 inch, TNR 11pt, line spacing = 1
- Tiêu ngữ 2 cột; dòng Họ tên/MSSV/Mã đề; đường kẻ ngang
- “Text cuối đề” (người dùng nhập) đặt NGAY DƯỚI đường kẻ, căn giữa & nghiêng (KHÔNG in “Hết”)
- Đáp án có thể dàn 2 cột; loại nhãn A./B./C./D. gốc; in nhãn mới đậm
- Tăng độ ngẫu nhiên: chiến lược chọn câu (ngẫu nhiên / xoay vòng / phân tầng [<XX>]) + đảo đáp án
- Xuất dap_an.xlsx: cột A=Mã đề, sau đó Câu 1..N
"""

import os, re, random, io, copy, json
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, Text, END

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageTk

# Import tiện ích chung
import utils
import db

# ================= CẤU HÌNH =================
DEFAULT_NUM_OPTIONS    = 4
OUTPUT_FOLDER          = "DE_TRON"
COMPACT_TEXT_THRESHOLD = 80   # ngưỡng coi đáp án gọn để dàn 2 cột
SUPER_COMPACT_THRESHOLD = 25  # ngưỡng coi đáp án cực gọn để dàn 4 cột
QUESTION_FONT_PT       = 11
BOLD_RATIO_THRESHOLD   = 0.5  # ngưỡng tỷ lệ bold sau nhãn khi nhận diện đáp án đúng

# ============= PHÂN TÍCH RUN & MEDIA (1 LẦN) =============
def _run_non_black(run) -> bool:
    return utils.get_run_color(run) != "NONE"

def _run_bold(run) -> bool: return bool(getattr(run.font, "bold", False))
def _run_ul(run) -> bool:   return bool(getattr(run.font, "underline", False))

def _concat_runs(paragraph):
    spans, full, pos = [], [], 0
    for run in paragraph.runs:
        t = run.text or ""
        if t:
            spans.append((pos, pos+len(t), run, t))
            full.append(t); pos += len(t)
    return spans, "".join(full)

# _LABEL_REGEX đã được chuyển lên phần cấu hình REGEX

def _label_span_end(paragraph) -> int:
    return utils.label_span_end(paragraph)

def looks_like_option(p) -> bool:
    txt = (p.text or "")
    if txt.strip().startswith("[<$>]"):
        return True
    return utils._LABEL_REGEX.match(txt or "") is not None

def _extract_media_from_element(elem, part, seen_rids: set, seen_hash: set) -> List[Dict[str, Any]]:
    """
    Trích xuất hình ảnh từ một element (run hoặc paragraph level).
    Hỗ trợ drawing (hiện đại), imagedata/pict (cổ điển, MathType).
    """
    spans = []

    def _push_blob(blob, rid_hint=None):
        if rid_hint and rid_hint in seen_rids:
            return
        h = None
        try:
            import hashlib
            h = hashlib.sha1(blob).hexdigest()
        except Exception: pass
        if h and h in seen_hash:
            return
        spans.append({"type": "image", "blob": blob})
        if rid_hint: seen_rids.add(rid_hint)
        if h: seen_hash.add(h)

    # 1) Hiện đại: drawing/blip
    # Sử dụng local-name để bỏ qua vấn đề namespace prefix
    for blip in elem.xpath('.//*[local-name()="blip"]'):
        rid = blip.get(qn('r:embed')) or blip.get('r:embed')
        if rid and part and rid in part.related_parts:
            try: _push_blob(part.related_parts[rid].blob, rid_hint=rid)
            except: pass

    # 2) Cổ điển hoặc MathType: imagedata/pict/object/shape
    for imd in elem.xpath('.//*[local-name()="imagedata"]'):
        rid = imd.get(qn('r:id')) or imd.get('r:id') or imd.get(qn('r:embed'))
        if rid and part and rid in part.related_parts:
            try: _push_blob(part.related_parts[rid].blob, rid_hint=rid)
            except: pass
            
    for ole in elem.xpath('.//*[local-name()="OLEObject"]'):
        rid = ole.get(qn('r:id')) or ole.get('r:id')
        if rid and part and rid in part.related_parts:
            # OLEObject thường đi kèm với một hình ảnh preview (imagedata) ở trên
            # nhưng đôi khi ta cần kiểm tra thêm link
            pass

    return spans

def collect_para_spans(p, strip_to_black=True, skip_len=0) -> List[Dict[str, Any]]:
    """
    Duyệt tuần tự XML để lấy Spans (Text, Image, OMML) theo đúng thứ tự.
    skip_len: số ký tự văn bản ở đầu paragraph cần bỏ qua (dùng cho nhãn A./B./C.)
    """
    spans_out = []
    part = getattr(p, 'part', None)
    seen_rids = set()
    seen_hash = set()
    
    consumed_txt = 0
    
    for child in p._element.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        
        # 1. OMML (Công thức toán học)
        if tag == "oMath":
            spans_out.append({"type": "omml", "xml": copy.deepcopy(child)})
        elif tag == "oMathPara":
            # oMathPara thường chứa oMath bên trong, ta lấy các oMath để đồng nhất dữ liệu
            for m in child.xpath('.//*[local-name()="oMath"]'):
                spans_out.append({"type": "omml", "xml": copy.deepcopy(m)})
            
        # 2. Run (Chứa văn bản hoặc hình ảnh)
        elif tag == "r":
            # Kiểm tra ảnh trong run trước
            r_media = _extract_media_from_element(child, part, seen_rids, seen_hash)
            spans_out.extend(r_media)
            
            # Kiểm tra văn bản trong run
            # Một run có thể có nhiều <w:t>
            from docx.text.run import Run
            run_obj = Run(child, p)
            t = run_obj.text or ""
            if t:
                # Logic cắt nhãn
                cut = min(len(t), max(0, skip_len - consumed_txt))
                t_show = t[cut:]
                consumed_txt += len(t)
                
                if t_show:
                    # Lấy màu hex
                    color_hex = None
                    try:
                        if run_obj.font and run_obj.font.color and run_obj.font.color.rgb:
                            color_hex = str(run_obj.font.color.rgb).replace("#", "").upper()
                    except: pass
                    
                    spans_out.append({
                        "type": "text",
                        "text": utils.remove_marker_text(t_show),
                        "bold": bool(run_obj.bold),
                        "italic": bool(run_obj.italic),
                        "underline": bool(run_obj.underline),
                        "subscript": bool(run_obj.font.subscript),
                        "superscript": bool(run_obj.font.superscript),
                        "font_name": run_obj.font.name, # Lưu font gốc
                        "strip_black": bool(strip_to_black),
                        "color_hex": color_hex
                    })
                    
            # Kiểm tra công thức inline bên trong run (nếu có)
            for m in child.xpath('.//*[local-name()="oMath"]'):
                spans_out.append({"type": "omml", "xml": copy.deepcopy(m)})
        
        # 3. Drawing/Pict/Object ở paragraph-level (MathType hoặc ảnh lồng)
        elif tag in ("drawing", "pict", "object"):
             spans_out.extend(_extract_media_from_element(child, part, seen_rids, seen_hash))

    return spans_out

def para_collect_media_spans(p) -> List[Dict[str, Any]]:
    return collect_para_spans(p, strip_to_black=False)

def _paragraph_has_media(spans: List[Dict[str, Any]]) -> bool:
    return any(sp["type"] in ("image", "omml") for sp in spans)

def precompile_option_paragraph(p, strip_to_black=True) -> Dict[str, Any]:
    """
    Dùng collect_para_spans để lấy dữ liệu đáp án đồng nhất và đúng thứ tự.
    """
    # Tính nhãn để bỏ (skip_len) - Đồng bộ với bộ nhận diện chung
    full_text = p.text or ""
    skip = 0
    end_lbl = _label_span_end(p)
    if end_lbl > 0:
        skip = end_lbl

    # Thu thập spans
    spans_out = collect_para_spans(p, strip_to_black=strip_to_black, skip_len=skip)
    
    # Metrics nhận diện đáp án đúng (vẫn cần tính)
    # Lưu ý: Chúng ta duyệt lại runs để tính metrics cho chính xác theo logic cũ (bold ratio, color...)
    any_color = False
    any_ul = False
    content_hi = False
    label_colored = False
    label_color = "NONE"
    content_color = "NONE"
    label_color_chars = 0 # Số ký tự có màu trong nhãn
    label_total_chars = 0 # Tổng số ký tự trong nhãn (loại trừ space)
    bold_chars = 0
    total_chars = 0
    consumed = 0
    for run in p.runs:
        t = run.text or ""
        run_start = consumed
        consumed += len(t)
        
        r_col = utils.get_run_color(run)
        if r_col != "NONE":
            any_color = True
            # Ưu tiên màu nhãn
            if end_lbl >= 0 and run_start < end_lbl:
                lbl_end_in_run = min(len(t), end_lbl - run_start)
                lbl_chars = sum(1 for ch in t[:lbl_end_in_run] if not ch.isspace())
                label_color_chars += lbl_chars
                label_colored = True
                label_color = r_col
            else:
                if content_color == "NONE":
                    content_color = r_col
        
        # Tính tổng ký tự nhãn (để tính ratio)
        if end_lbl >= 0 and run_start < end_lbl:
            lbl_end_in_run = min(len(t), end_lbl - run_start)
            label_total_chars += sum(1 for ch in t[:lbl_end_in_run] if not ch.isspace())

        if _run_ul(run): any_ul = True
        
        # sau nhãn
        if end_lbl >= 0 and consumed > end_lbl:
            # văn bản thực sự sau nhãn
            offset = max(0, end_lbl - run_start)
            after_txt = t[offset:]
            if any(ch.isalnum() for ch in after_txt):
                if _run_non_black(run) or _run_ul(run) or _run_bold(run):
                    content_hi = True
                total_chars += sum(1 for ch in after_txt if ch.isalnum())
                if _run_bold(run):
                    bold_chars += sum(1 for ch in after_txt if ch.isalnum())

    ratio = (bold_chars / total_chars) if total_chars else 0.0
    l_ratio = (label_color_chars / label_total_chars) if label_total_chars else 0.0
    metrics = {
        "any_color": any_color, 
        "any_ul": any_ul, 
        "content_hi": content_hi, 
        "content_bold_ratio": ratio,
        "label_colored": label_colored,
        "label_color": label_color,
        "label_color_ratio": l_ratio,
        "content_color": content_color
    }

    # Độ gọn
    plain_text = utils.remove_marker_text(p.text or "", trim=True)
    is_compact = (len(plain_text) <= COMPACT_TEXT_THRESHOLD) and (not _paragraph_has_media(spans_out))

    return {"spans": spans_out, "is_compact": is_compact, "metrics": metrics}

def _paragraph_has_media(spans: List[Dict[str, Any]]) -> bool:
    return any(sp["type"] in ("image", "omml") for sp in spans)

def precompile_media_paragraph(p) -> List[Dict[str, Any]]:
    """Media-only spans cho thân câu/đoạn phụ."""
    # Bây giờ collect_para_spans lấy cả text, nên hạm này sẽ trả về spans hỗn hợp.
    # Trong render_question, stem_text được in trước, sau đó mới đến media spans.
    # ĐỂ TRÁNH LẶP TEXT: chúng ta chỉ lấy những spans KHÔNG phải text? 
    # KHÔNG, tốt nhất là gom stem_text và media thành 1 luồng duy nhất.
    return collect_para_spans(p, strip_to_black=False)

# ============= NHẬN DIỆN ĐÁP ÁN ĐÚNG (1 LẦN) =============
def pick_correct_option_index(metrics_list: List[Dict[str, Any]]) -> Optional[int]:
    """
    Xác định đáp án đúng dựa trên sự khác biệt duy nhất (Unique One Out).
    """
    if not metrics_list: return None
    n = len(metrics_list)

    # 1) LOGIC UNIQUE COLOR & COVERAGE (Ưu tiên Cao nhất)
    # Kiểm tra tỷ lệ bao phủ màu của Nhãn (label_color_ratio)
    l_ratios = [m.get("label_color_ratio", 0) for m in metrics_list]
    if any(r > 0.4 for r in l_ratios): # Có ít nhất 1 nhãn bôi màu trên 40% diện tích
        max_r = max(l_ratios)
        # Nếu chỉ có 1 thằng đạt Max và Max đủ lớn (>50%)
        if max_r > 0.5 and l_ratios.count(max_r) == 1:
            return l_ratios.index(max_r)
        
        # Nếu có sự khác biệt rõ rệt (ví dụ 100% vs 20%)
        # Ta lấy những thằng có ratio > 0.5
        idx_strong = [i for i, r in enumerate(l_ratios) if r > 0.5]
        if len(idx_strong) == 1:
            return idx_strong[0]
            
    # Kiểm tra màu Nhãn (label_color) - Theo Unique Color cũ nhưng lọc nhiễu
    # (Chỉ xét các nhãn có ratio đáng kể > 20%)
    l_colors = [m.get("label_color", "NONE") if m.get("label_color_ratio", 0) > 0.2 else "NONE" 
                for m in metrics_list]
    from collections import Counter
    c_counts = Counter(l_colors)
    if len(c_counts) > 1:
        # Tìm màu chỉ xuất hiện đúng 1 lần
        uniques = [col for col, count in c_counts.items() if count == 1 and col != "NONE"]
        if len(uniques) == 1:
            return l_colors.index(uniques[0])
        # Nếu có 3 thằng Đỏ, 1 thằng Đen -> thằng Đen là duy nhất
        uniques_all = [col for col, count in c_counts.items() if count == 1]
        if len(uniques_all) == 1:
            return l_colors.index(uniques_all[0])

    # Kiểm tra màu Nội dung (content_color)
    c_colors = [m.get("content_color", "NONE") for m in metrics_list]
    cc_counts = Counter(c_colors)
    if len(cc_counts) > 1:
        uniques = [col for col, count in cc_counts.items() if count == 1 and col != "NONE"]
        if len(uniques) == 1:
            return c_colors.index(uniques[0])
        uniques_all = [col for col, count in cc_counts.items() if count == 1]
        if len(uniques_all) == 1:
            return c_colors.index(uniques_all[0])

    # 2) DỰ PHÒNG: Logic cũ (nếu không có sự khác biệt màu rõ ràng)
    # Duy nhất label_colored (Màu bất kỳ ở nhãn, nếu các thằng khác NONE)
    idx_lbl_col = [i for i, m in enumerate(metrics_list) if m.get("label_colored")]
    if len(idx_lbl_col) == 1:
        return idx_lbl_col[0]

    # content_hi
    idx_hi = [i for i,m in enumerate(metrics_list) if m.get("content_hi")]
    if len(idx_hi) == 1:
        return idx_hi[0]
    
    # duy nhất any_color
    idx_col = [i for i,m in enumerate(metrics_list) if m.get("any_color")]
    if len(idx_col) == 1:
        return idx_col[0]
    
    # duy nhất any_ul
    idx_ul = [i for i,m in enumerate(metrics_list) if m.get("any_ul")]
    if len(idx_ul) == 1:
        return idx_ul[0]
    
    # bold_ratio vượt ngưỡng
    ratios = [m.get("content_bold_ratio", 0) for m in metrics_list]
    if ratios:
        max_r = max(ratios)
        if max_r >= BOLD_RATIO_THRESHOLD and ratios.count(max_r) == 1:
            return ratios.index(max_r)
            
    return None

# ============= TÁCH & PRECOMPILE NGÂN HÀNG =============
def split_questions_from_docx(docx_path: str, num_options: int = DEFAULT_NUM_OPTIONS):
    doc = Document(docx_path)
    paras = doc.paragraphs
    n = len(paras)
    questions = []
    warnings = []
    i = 0
    while i < n:
        qid, diff_code, stem_text = utils.parse_question_header(paras[i].text)
        if qid is None:
            i += 1; continue

        stem_para_idx = i
        stem_extra_idxs: List[int] = []
        opt_idxs: List[int] = []
        i += 1; options_started = False

        while i < n:
            p2 = paras[i]; t2 = (p2.text or "").strip()
            if utils.QUESTION_HEADER_PAT.match(t2): break
            if not options_started:
                if looks_like_option(p2):
                    options_started = True
                    opt_idxs.append(i)
                else:
                    stem_extra_idxs.append(i)
                i += 1
            else:
                # gom tất cả đoạn thuộc block phương án (mỗi phương án 1 paragraph; nếu kéo dài vẫn gộp)
                if looks_like_option(p2):
                    opt_idxs.append(i)
                else:
                    opt_idxs.append(i)
                i += 1

        if not opt_idxs:
            continue

        # PRECOMPILE phương án
        opt_infos = [precompile_option_paragraph(paras[idx], strip_to_black=True) for idx in opt_idxs]
        # Chọn đáp án đúng
        correct_pick = pick_correct_option_index([info["metrics"] for info in opt_infos])
        if correct_pick is None:
            warnings.append(f"Câu hỏi tại đoạn {stem_para_idx + 1} (qid: {qid}): Không tìm thấy đáp án đúng.")
            continue

        # Cắt còn 4 phương án đầu; đảm bảo có đáp án đúng trong 4 cái
        k = min(num_options, len(opt_idxs))
        slice_idxs = opt_idxs[:k]
        slice_infos = opt_infos[:k]
        if correct_pick >= k:
            slice_idxs[-1]  = opt_idxs[correct_pick]
            slice_infos[-1] = opt_infos[correct_pick]
            correct_pick = k - 1

        full_stem_para_txt = paras[stem_para_idx].text or ""
        raw_match = utils.QUESTION_HEADER_PAT.search(full_stem_para_txt)
        skip_len_stem = raw_match.start(3) if raw_match else 0
        
        stem_media_spans = collect_para_spans(paras[stem_para_idx], strip_to_black=False, skip_len=skip_len_stem)

        extra_media_spans: List[List[Dict[str, Any]]] = []
        for ex in stem_extra_idxs:
            extra_media_spans.append(collect_para_spans(paras[ex], strip_to_black=False))

        # Nếu phần chữ trống (do chỉ có công thức), hãy lấy tóm tắt từ spans
        final_stem_text = utils.remove_marker_text(stem_text or "", trim=True)
        if not final_stem_text and stem_media_spans:
            final_stem_text = utils.get_spans_text_summary(stem_media_spans)

        questions.append({
            "qid": qid,
            "diff_code": diff_code,
            "stem_para_idx": stem_para_idx,
            "stem_text": final_stem_text,
            "stem_media_spans": stem_media_spans,
            "stem_extra_media_spans": extra_media_spans,
            "options": [
                {"para_idx": pi, "info": inf} for pi, inf in zip(slice_idxs, slice_infos)
            ],
            "correct_index": int(correct_pick)
        })

    return questions, doc, warnings

# ============= CÁC HÀM TRÌNH BÀY (WORD) =============
def _add_tab_char(p):
    try: p.add_run().add_tab()
    except: p.add_run('\t')

def add_info_line_with_leaders(doc, exam_code: int, tab1=4200, tab2=8200, tab3=9000):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._element.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is not None: pPr.remove(tabs)
    tabs = OxmlElement('w:tabs')
    def _add_tab(pos, val='left', leader='dot'):
        t = OxmlElement('w:tab'); t.set(qn('w:val'), val); t.set(qn('w:pos'), str(pos))
        if leader: t.set(qn('w:leader'), leader); tabs.append(t)
    _add_tab(tab1, 'left', 'dot'); _add_tab(tab2, 'left', 'dot'); _add_tab(tab3, 'right', None)
    pPr.append(tabs)

    r = p.add_run("Họ và tên sinh viên: "); utils.set_run_font(r, 'Times New Roman'); r.font.size = Pt(11); _add_tab_char(p)
    r = p.add_run("  Mã số sinh viên: ");   utils.set_run_font(r, 'Times New Roman'); r.font.size = Pt(11); _add_tab_char(p)
    _add_tab_char(p)
    r = p.add_run(f"Mã đề: {exam_code}"); utils.set_run_font(r, 'Times New Roman'); r.bold = True; r.font.size = Pt(11)

def add_header_2cols(doc: Document, school: str, faculty: str, exam_title: str, school_year: str,
                     subject: str, duration_text: str, exam_code: int):
    tbl = doc.add_table(rows=1, cols=2); utils.remove_table_borders(tbl); tbl.autofit = True
    left = tbl.cell(0, 0); lp = left.paragraphs[0]; lp.text = ""
    utils.add_styled_text(lp,  school,  size=13, bold=False, uppercase=True,  align=WD_ALIGN_PARAGRAPH.CENTER)
    lp2 = left.add_paragraph(); utils.add_styled_text(lp2, faculty, size=12, bold=True, uppercase=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    right = tbl.cell(0, 1); rp = right.paragraphs[0]; rp.text = ""
    utils.add_styled_text(rp,  exam_title,                size=13, bold=True, uppercase=True,  align=WD_ALIGN_PARAGRAPH.CENTER)
    rp2 = right.add_paragraph(); utils.add_styled_text(rp2, f"NĂM HỌC {school_year}", size=12, bold=True, uppercase=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    rp3 = right.add_paragraph(); utils.add_styled_text(rp3, f"Môn: {subject}",        size=12, bold=True, uppercase=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    rp4 = right.add_paragraph(); utils.add_styled_text(rp4, duration_text,            size=11, bold=False, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph("")
    add_info_line_with_leaders(doc, exam_code)
    utils.add_horizontal_rule(doc, thickness_eighths=8)
    doc.add_paragraph("")

def add_exam_end(doc: Document, tail_text: Optional[str] = None, use_section_pages: bool = False):
    utils.add_horizontal_rule(doc, thickness_eighths=8)
    
    if tail_text and tail_text.strip():
        p_tail = doc.add_paragraph()
        p_tail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_tail.add_run(tail_text.strip())
        utils.set_run_font(r, 'Times New Roman')
        r.font.size = Pt(QUESTION_FONT_PT)
        r.italic = True

def option_is_compact_precompiled(opt_info: Dict[str, Any]) -> bool:
    if _paragraph_has_media(opt_info["spans"]):
        return False
    return bool(opt_info.get("is_compact", False))

# ============= VẼ CÂU HỎI =============
def render_question(dst_doc: Document, qnum: int, q: Dict, ans_layout: str = "Tự động"):
    p_title = dst_doc.add_paragraph(); p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_label = p_title.add_run(f"Câu {qnum}. "); r_label.bold = True; utils.set_run_font(r_label, 'Times New Roman'); r_label.font.size = Pt(QUESTION_FONT_PT)
    
    if q.get("stem_media_spans"):
        utils.render_spans_into_paragraph(p_title, q["stem_media_spans"], font_size_pt=QUESTION_FONT_PT)
    else:
        r_stem = p_title.add_run((q['stem_text'] or '').strip()); utils.set_run_font(r_stem, 'Times New Roman'); r_stem.font.size = Pt(QUESTION_FONT_PT)

    for extra_spans in q.get("stem_extra_media_spans", []):
        if extra_spans:
            p_media2 = dst_doc.add_paragraph()
            utils.render_spans_into_paragraph(p_media2, extra_spans, font_size_pt=QUESTION_FONT_PT)

    opts = q["options"][:4]
    letters = ["A", "B", "C", "D"]
    
    # Xác định bố cục thực tế
    final_layout = ans_layout
    if final_layout == "Tự động":
        # Heuristic tự động
        has_media = any(_paragraph_has_media(o["info"]["spans"]) for o in opts)
        if has_media or len(opts) != 4:
            final_layout = "Dưới nhau (1x4)"
        else:
            all_txt_lens = [len(utils.get_clean_text("".join(sp.get("text", "") for sp in o["info"]["spans"]))) for o in opts]
            max_len = max(all_txt_lens) if all_txt_lens else 0
            if max_len <= SUPER_COMPACT_THRESHOLD:
                final_layout = "Dàn 4 cột (4x1)"
            elif max_len <= COMPACT_TEXT_THRESHOLD:
                final_layout = "Dàn 2 cột (2x2)"
            else:
                final_layout = "Dưới nhau (1x4)"

    if final_layout == "Dàn 4 cột (4x1)" and len(opts) == 4:
        tbl = dst_doc.add_table(rows=1, cols=4); tbl.autofit = True; utils.remove_table_borders(tbl)
        # Thiết lập độ rộng cột tương đối đều nhau
        for i, opt in enumerate(opts):
            cell = tbl.cell(0, i); p_cell = cell.paragraphs[0]
            lab = p_cell.add_run(f" {letters[i]}. "); lab.bold=True; utils.set_run_font(lab, 'Times New Roman'); lab.font.size=Pt(QUESTION_FONT_PT)
            utils.render_spans_into_paragraph(p_cell, opt["info"]["spans"], font_size_pt=QUESTION_FONT_PT)
    elif final_layout == "Dàn 2 cột (2x2)" and len(opts) == 4:
        tbl = dst_doc.add_table(rows=2, cols=2); tbl.autofit = True; utils.remove_table_borders(tbl)
        positions = [(0,0), (0,1), (1,0), (1,1)]
        for i, opt in enumerate(opts):
            r0,c0 = positions[i]; cell = tbl.cell(r0,c0); p_cell = cell.paragraphs[0]
            lab = p_cell.add_run(f"   {letters[i]}. "); lab.bold=True; utils.set_run_font(lab, 'Times New Roman'); lab.font.size=Pt(QUESTION_FONT_PT)
            utils.render_spans_into_paragraph(p_cell, opt["info"]["spans"], font_size_pt=QUESTION_FONT_PT)
    else:
        # Mặc định: Dưới nhau (1x4)
        for i, opt in enumerate(opts):
            p_opt = dst_doc.add_paragraph()
            lab = p_opt.add_run(f"   {letters[i]}. "); lab.bold=True; utils.set_run_font(lab, 'Times New Roman'); lab.font.size=Pt(QUESTION_FONT_PT)
            utils.render_spans_into_paragraph(p_opt, opt["info"]["spans"], font_size_pt=QUESTION_FONT_PT)

# ============= ĐẢO THỨ TỰ ĐÁP ÁN =============
def shuffle_options_in_place(q: dict, rng: random.Random, enable: bool = True):
    if not enable:
        return q
    opts = q["options"]
    old_correct = q["correct_index"]
    order = list(range(len(opts)))
    rng.shuffle(order)
    new_opts = [opts[i] for i in order]
    new_correct = order.index(old_correct)
    q["options"] = new_opts
    q["correct_index"] = new_correct
    return q

# ============= CHIẾN LƯỢC CHỌN CÂU =============
def choose_questions(all_questions: List[dict], n_questions: int, strategy: str,
                     exam_idx: int, total_exams: int, base_perm: List[int],
                     rng: random.Random) -> List[dict]:
    N = len(all_questions)
    if n_questions > N:
        n_questions = N

    if strategy == "Ngẫu nhiên thuần":
        return rng.sample(all_questions, n_questions)

    if strategy == "Xoay vòng theo mã đề":
        start = ((exam_idx - 1) * n_questions) % N
        take = []
        i = 0
        while len(take) < n_questions and i < N:
            take.append(all_questions[ base_perm[(start + i) % N] ])
            i += 1
        return take

    if strategy == "Phân tầng theo độ khó (nếu có)":
        buckets: Dict[str, List[dict]] = {}
        for q in all_questions:
            key = (q.get("diff_code") or "UNK")
            buckets.setdefault(key, []).append(q)
        if len(buckets) <= 1:
            return rng.sample(all_questions, n_questions)

        keys = sorted(buckets.keys())
        k = len(keys)
        base = n_questions // k
        rem = n_questions % k

        chosen = []
        for idx, key in enumerate(keys):
            need = base + (1 if idx < rem else 0)
            pool = buckets[key][:]
            rng.shuffle(pool)
            chosen.extend(pool[:need])

        if len(chosen) < n_questions:
            remaining = [q for q in all_questions if q not in chosen]
            if remaining:
                extra = rng.sample(remaining, min(len(remaining), n_questions - len(chosen)))
                chosen.extend(extra)

        rng.shuffle(chosen)
        return chosen

    return rng.sample(all_questions, n_questions)

# ============= XÂY 1 ĐỀ TỪ DANH SÁCH CÂU + (TUỲ CHỌN) ĐẢO =============
def build_exam_from_selected(selected_questions: List[dict],
                             school, faculty, subject, duration_text,
                             exam_title, school_year, exam_code,
                             tail_text, shuffle_answers: bool, seed: int,
                             ans_layout: str = "Tự động"):
    rng = random.Random(seed)

    # TỐI ƯU HÓA: Không dùng deepcopy toàn bộ (nặng). 
    # Chỉ copy cấu trúc và trộn thứ tự options.
    selected = []
    for q_orig in selected_questions:
        q_copy = q_orig.copy() # Shallow copy cho các key cấp 1
        q_copy["options"] = list(q_orig["options"]) # Copy list options để shuffle
        if shuffle_answers:
            shuffle_options_in_place(q_copy, rng, enable=True)
        selected.append(q_copy)

    dst_doc = Document()
    utils.set_page_layout(dst_doc)
    utils.set_single_line_spacing(dst_doc)
    utils.add_page_number_to_footer(dst_doc, use_section_pages=False)
    add_header_2cols(dst_doc, school, faculty, exam_title, school_year, subject, duration_text, exam_code)

    for i, q in enumerate(selected, 1):
        render_question(dst_doc, i, q, ans_layout=ans_layout)

    add_exam_end(dst_doc, tail_text=tail_text, use_section_pages=False)
    # Khôi phục đánh số trang footer
    utils.add_page_number_to_footer(dst_doc, use_section_pages=False)
    return dst_doc, selected

# ================= GUI – TAB TRỘN ĐỀ =================
class MixTab:
    def __init__(self, parent):
        self.parent = parent
        self.questions: List[Dict] = []
        self.all_questions: List[Dict] = []
        self.selected_indices: List[int] = []
        self.pinned_indices: List[int] = [] # Câu hỏi cố định
        self.matrix_data: List[Dict] = [] # Danh sách các tệp nguồn và số lượng câu
        self.src_doc: Optional[Document] = None

        try:
            from tkinterdnd2 import DND_FILES
            HAS_DND = True
        except ImportError:
            HAS_DND = False

        self.var_code_start = tk.StringVar(value="100")
        self.var_num_tests = tk.StringVar(value="5")
        self.var_num_questions = tk.StringVar(value="20")
        self.var_export_mode = tk.StringVar(value="Chỉ tổng hợp")
        self.var_strategy = tk.StringVar(value="Ngẫu nhiên thuần")
        self.var_ans_layout = tk.StringVar(value="Tự động")

        # ====== Layout chính: 2 cột đều nhau ======
        parent.columnconfigure(0, weight=1, uniform="col")
        parent.columnconfigure(1, weight=1, uniform="col")

        # --- Khung trái: Thông tin đề ---
        frm_left = ttk.LabelFrame(parent, text="Thông tin đề", padding=utils.PAD_M)
        frm_left.grid(row=0, column=0, sticky="nsew", padx=utils.PAD_S, pady=utils.PAD_S)
        frm_left.columnconfigure(1, weight=1)

        r = 0
        def _add_entry(label, var, row):
            ttk.Label(frm_left, text=label).grid(row=row, column=0, sticky="e", padx=utils.PAD_S, pady=utils.PAD_XS)
            ttk.Entry(frm_left, textvariable=var).grid(row=row, column=1, sticky="ew", padx=utils.PAD_S, pady=utils.PAD_XS)

        self.var_school = tk.StringVar(value="ĐẠI HỌC ĐIỆN LỰC")
        _add_entry("Trường:", self.var_school, r); r+=1

        self.var_faculty = tk.StringVar(value="KHOA NĂNG LƯỢNG MỚI")
        _add_entry("Khoa:", self.var_faculty, r); r+=1

        self.var_examtitle = tk.StringVar(value="KIỂM TRA SỐ 1")
        _add_entry("Tên bài:", self.var_examtitle, r); r+=1

        self.var_year = tk.StringVar(value="2025 - 2026")
        _add_entry("Năm học:", self.var_year, r); r+=1

        self.var_subject = tk.StringVar(value="Tiết kiệm năng lượng")
        _add_entry("Môn:", self.var_subject, r); r+=1

        self.var_semester = tk.StringVar(value="1")
        _add_entry("Học kỳ:", self.var_semester, r); r+=1

        self.var_class = tk.StringVar(value="D17DIENLANH1")
        _add_entry("Lớp:", self.var_class, r); r+=1

        self.var_duration = tk.StringVar(value="60 phút")
        _add_entry("Thời gian:", self.var_duration, r); r+=1

        ttk.Label(frm_left, text="Ghi chú:").grid(row=r, column=0, sticky="ne", padx=utils.PAD_S, pady=utils.PAD_XS)
        self.txt_tail = Text(frm_left, height=3, wrap="word", font=utils.FONT_MAIN)
        self.txt_tail.grid(row=r, column=1, sticky="ew", padx=utils.PAD_S, pady=utils.PAD_XS); r+=1

        # --- Khung phải: Cấu hình sinh đề ---
        frm_right = ttk.LabelFrame(parent, text="Cấu hình sinh đề", padding=utils.PAD_M)
        frm_right.grid(row=0, column=1, sticky="nsew", padx=utils.PAD_S, pady=utils.PAD_S)
        frm_right.columnconfigure(1, weight=1)

        r2 = 0
        def _add_config_row(label, var, row):
            ttk.Label(frm_right, text=label).grid(row=row, column=0, sticky="e", padx=utils.PAD_S, pady=utils.PAD_XS)
            ttk.Entry(frm_right, textvariable=var).grid(row=row, column=1, sticky="ew", padx=utils.PAD_S, pady=utils.PAD_XS)

        ttk.Label(frm_right, text="Nguồn câu hỏi:").grid(row=r2, column=0, sticky="e", padx=utils.PAD_S, pady=utils.PAD_XS)
        frm_btns = ttk.Frame(frm_right)
        frm_btns.grid(row=r2, column=1, sticky="w", padx=utils.PAD_S, pady=utils.PAD_XS)
        ttk.Button(frm_btns, text="File", command=self.load_file, width=8).pack(side=tk.LEFT)
        ttk.Button(frm_btns, text="Kho", command=self.load_from_bank, width=8).pack(side=tk.LEFT, padx=2)
        ttk.Button(frm_btns, text="CSDL", command=self.load_from_db, width=8, bootstyle="info").pack(side=tk.LEFT)
        r2 += 1
        
        self.lbl_bank = ttk.Label(frm_right, text="Chưa chọn file", foreground="orange", font=utils.FONT_MAIN)
        self.lbl_bank.grid(row=r2, column=1, sticky="w", padx=utils.PAD_S); r2+=1

        self.btn_matrix = ttk.Button(frm_right, text="📋 Thiết lập Ma trận đề (Nhiều tệp)", command=self.manage_matrix)
        self.btn_matrix.grid(row=r2, column=0, columnspan=2, sticky="ew", padx=utils.PAD_S, pady=utils.PAD_S); r2+=1

        self.btn_select_qs = ttk.Button(frm_right, text="🎯 Chọn câu hỏi cụ thể", command=self.select_questions, state="disabled")
        self.btn_select_qs.grid(row=r2, column=0, columnspan=2, sticky="ew", padx=utils.PAD_S, pady=utils.PAD_XS); r2+=1

        _add_config_row("Mã đề bắt đầu:", self.var_code_start, r2); r2+=1
        _add_config_row("Số lượng đề:", self.var_num_tests, r2); r2+=1
        _add_config_row("Số câu/đề:", self.var_num_questions, r2); r2+=1

        ttk.Label(frm_right, text="Loại file xuất:").grid(row=r2, column=0, sticky="e", padx=utils.PAD_S, pady=utils.PAD_XS)
        ttk.OptionMenu(frm_right, self.var_export_mode, "Chỉ tổng hợp", 
                       "Chỉ tổng hợp", "Cả hai", "Chỉ từng đề").grid(row=r2, column=1, sticky="ew", padx=utils.PAD_S); r2+=1
        
        ttk.OptionMenu(frm_right, self.var_strategy, "Ngẫu nhiên thuần", 
                       "Ngẫu nhiên thuần", "Xoay vòng theo mã đề", "Phân tầng độ khó").grid(row=r2, column=1, sticky="ew", padx=utils.PAD_S); r2+=1

        ttk.Label(frm_right, text="Bố cục đáp án:").grid(row=r2, column=0, sticky="e", padx=utils.PAD_S, pady=utils.PAD_XS)
        ttk.OptionMenu(frm_right, self.var_ans_layout, "Tự động", 
                       "Tự động", "Dàn 4 cột (4x1)", "Dàn 2 cột (2x2)", "Dưới nhau (1x4)").grid(row=r2, column=1, sticky="ew", padx=utils.PAD_S); r2+=1

        self.var_shuffle_ans = tk.BooleanVar(value=True)
        ttk.Checkbutton(frm_right, text="Đảo thứ tự đáp án (A–D)", variable=self.var_shuffle_ans)\
            .grid(row=r2, column=1, sticky="w", padx=utils.PAD_S); r2 += 1

        self.btn_run = ttk.Button(frm_right, text="🚀 Bắt đầu Trộn đề", command=self.run, state="disabled", bootstyle="success")
        self.btn_run.grid(row=r2, column=0, columnspan=2, sticky="ew", padx=utils.PAD_S, pady=utils.PAD_S); r2+=1


        # --- Khung tiến độ ---
        frm_progress = ttk.Frame(parent, padding=utils.PAD_S)
        frm_progress.grid(row=1, column=0, columnspan=2, sticky="ew")
        
        self.lbl_status = ttk.Label(frm_progress, text="Sẵn sàng", font=("Arial", 10, "italic"))
        self.lbl_status.pack(side=tk.TOP, anchor="w", padx=utils.PAD_S)
        
        self.progress = ttk.Progressbar(frm_progress, orient=tk.HORIZONTAL, mode='determinate', bootstyle="success")
        self.progress.pack(side=tk.TOP, fill=tk.X, padx=utils.PAD_S, pady=utils.PAD_XS)

        # Nhật ký gọn gàng
        frm_log = ttk.LabelFrame(parent, text="Nhật ký")
        frm_log.grid(row=2, column=0, columnspan=2, sticky="nsew", padx=utils.PAD_S, pady=utils.PAD_S)
        parent.rowconfigure(2, weight=1)
        frm_log.rowconfigure(0, weight=1)
        frm_log.columnconfigure(0, weight=1)

        self.log = Text(frm_log, height=5, wrap="word", font=("Consolas", 11))
        self.log.grid(row=0, column=0, sticky="nsew", padx=utils.PAD_XS, pady=utils.PAD_XS)
        vsb = ttk.Scrollbar(frm_log, orient="vertical", command=self.log.yview)
        vsb.grid(row=0, column=1, sticky="ns")
        self.log.configure(yscrollcommand=vsb.set)

    def logmsg(self, s): 
        self.log.insert(END, s + "\n"); self.log.see(END)

    def manage_matrix(self):
        """Mở hộp thoại thiết lập ma trận đề từ nhiều tệp nguồn."""
        dlg = tk.Toplevel(self.parent)
        utils.setup_dialog(dlg, width_pct=0.75, height_pct=0.7, title="Thiết lập Ma trận đề", parent=self.parent)

        frm_top = ttk.Frame(dlg, padding=10)
        frm_top.pack(fill=tk.X)
        ttk.Label(frm_top, text="Quản lý các nguồn câu hỏi và số lượng cần lấy:", font=("Arial", 11, "bold")).pack(side=tk.LEFT)
        
        frm_body = ttk.Frame(dlg, padding=10)
        frm_body.pack(fill=tk.BOTH, expand=True)

        cols = ("filename", "total", "pick")
        tree = ttk.Treeview(frm_body, columns=cols, show="headings", height=10)
        tree.heading("filename", text="Tên file ngân hàng")
        tree.heading("total", text="Tổng số câu")
        tree.heading("pick", text="Số câu cần lấy")
        tree.column("filename", width=400)
        tree.column("total", width=100, anchor="center")
        tree.column("pick", width=120, anchor="center")
        
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(frm_body, orient="vertical", command=tree.yview)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        tree.config(yscrollcommand=vsb.set)

        # Trạng thái hiện tại của ma trận
        # self.matrix_data = [{ "path": ..., "questions": [...], "pick_count": 10 }, ...]
        
        def refresh_tree():
            tree.delete(*tree.get_children())
            for item in self.matrix_data:
                tree.insert("", "end", values=(os.path.basename(item["path"]), len(item["questions"]), item["pick_count"]))

        refresh_tree()

        def add_file(from_bank=True):
            if from_bank:
                bank_dir = "NGAN_HANG_CAU_HOI"
                if not os.path.exists(bank_dir): return
                f = filedialog.askopenfilename(initialdir=bank_dir, filetypes=[("Word files", "*.docx")])
            else:
                f = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
            
            if not f: return
            try:
                qs, _, warns = split_questions_from_docx(f)
                if warns:
                    self.logmsg(f"Cảnh báo khi nạp {os.path.basename(f)}: " + "; ".join(warns))
                self.matrix_data.append({"path": f, "questions": qs, "pick_count": len(qs)})
                refresh_tree()
            except Exception as e:
                messagebox.showerror("Lỗi", str(e))

        def remove_file():
            sel = tree.selection()
            if not sel: return
            idx = tree.index(sel[0])
            self.matrix_data.pop(idx)
            refresh_tree()

        def edit_count():
            sel = tree.selection()
            if not sel: return
            idx = tree.index(sel[0])
            item = self.matrix_data[idx]
            from tkinter import simpledialog
            c = simpledialog.askinteger("Số lượng", f"Nhập số lượng câu muốn lấy từ {os.path.basename(item['path'])}:", 
                                         initialvalue=item["pick_count"], minvalue=0, maxvalue=len(item["questions"]))
            if c is not None:
                item["pick_count"] = c
                refresh_tree()

        def confirm():
            # Gom toàn bộ câu hỏi từ ma trận
            new_all_qs = []
            for item in self.matrix_data:
                # Đánh dấu nguồn để dễ truy vết nếu cần
                for q in item["questions"]:
                    q["source_file"] = os.path.basename(item["path"])
                new_all_qs.extend(item["questions"])
            
            self.all_questions = new_all_qs
            self.questions = new_all_qs.copy()
            self.selected_indices = list(range(len(new_all_qs)))
            self.pinned_indices = []
            
            self.lbl_bank.config(text=f"Đã nạp Ma trận ({len(self.matrix_data)} file)")
            self.btn_select_qs["state"] = "normal"
            self.btn_run["state"] = "normal"
            self.logmsg(f"Ma trận đề: Nạp {len(new_all_qs)} câu hỏi từ {len(self.matrix_data)} file.")
            dlg.destroy()

        frm_btns = ttk.Frame(dlg, padding=10)
        frm_btns.pack(fill=tk.X)
        ttk.Button(frm_btns, text="➕ Thêm file", command=add_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(frm_btns, text="💾 Thêm từ CSDL", command=lambda: self._add_db_bank_to_matrix(self.matrix_data, refresh_tree)).pack(side=tk.LEFT, padx=5)
        ttk.Button(frm_btns, text="✏️ Đổi số lượng", command=edit_count).pack(side=tk.LEFT, padx=5)
        ttk.Button(frm_btns, text="❌ Xóa", command=remove_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(frm_btns, text="✅ Xác nhận ma trận", command=confirm, bootstyle="success").pack(side=tk.RIGHT, padx=5)

    def _do_load_file(self, fname):
        try:
            qs, self.src_doc, warns = split_questions_from_docx(fname)
            if warns:
                self.logmsg("Cảnh báo: " + "; ".join(warns))
            if not qs:
                messagebox.showerror("Lỗi", "Không tìm thấy câu hỏi hợp lệ trong file này.")
                return
            self.all_questions = qs
            self.questions = qs.copy()
            self.selected_indices = list(range(len(qs)))
            self.pinned_indices = []
            self.matrix_data = [{"path": fname, "questions": qs, "pick_count": len(qs)}]
            self.lbl_bank.config(text=os.path.basename(fname))
            self.logmsg(f"Đã nạp {len(self.questions)} câu hỏi hợp lệ.")
            self.btn_run["state"] = "normal"
            self.btn_select_qs["state"] = "normal"
            self.btn_select_qs.config(text=f"Chọn câu hỏi & Cài đặt cố định")
        except Exception as e:
            messagebox.showerror("Lỗi", str(e)); self.logmsg(f"Lỗi: {e}")


    def load_from_bank(self):
        bank_dir = "NGAN_HANG_CAU_HOI"
        if not os.path.exists(bank_dir):
            messagebox.showinfo("Thông báo", "Thư mục Ngân hàng câu hỏi chưa được tạo hoặc trống.")
            return
        files = [f for f in os.listdir(bank_dir) if f.lower().endswith('.docx') and not f.startswith('~')]
        if not files:
            messagebox.showinfo("Thông báo", "Thư mục Ngân hàng câu hỏi hiện không có file nào.")
            return

        def on_select():
            sel = listbox.curselection()
            if not sel: return
            fname = os.path.join(bank_dir, files[sel[0]])
            dlg.destroy()
            self._do_load_file(fname)

        dlg = tk.Toplevel(self.parent)
        dlg.title("Chọn file từ Ngân hàng")
        utils.set_window_icon(dlg)
        dlg.geometry("600x450")
        dlg.transient(self.parent.winfo_toplevel())
        dlg.grab_set()

        ttk.Label(dlg, text="Chọn 1 file trong ngân hàng:").pack(padx=10, pady=(10, 5), anchor="w")
        listbox = tk.Listbox(dlg)
        listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        for f in files:
            listbox.insert(tk.END, f)
            
        ttk.Button(dlg, text="Chọn", command=on_select).pack(pady=10)

    def load_file(self):
        fname = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if not fname: return
        self._do_load_file(fname)

    # ============= TRỘN ĐỀ TỪ CSDL =============
    def _convert_db_questions(self, db_questions: List[Dict]) -> List[Dict]:
        """
        Chuyển đổi câu hỏi từ format DB sang format tương thích với split_questions_from_docx().
        Mỗi câu hỏi cần có: qid, stem_text, stem_media_spans, options, correct_index, diff_code...
        """
        converted = []
        for q in db_questions:
            # Deserialize spans
            stem_spans = db._deserialize_spans(q.get("stem_spans_json", "[]"))
            stem_extra = db._deserialize_spans_list(q.get("stem_extra_spans_json", "[]"))
            
            # Chuyển đổi options
            options = []
            correct_index = 0
            for i, opt in enumerate(q.get("options", [])):
                opt_spans = db._deserialize_spans(opt.get("option_spans_json", "[]"))
                
                # Tạo text từ spans nếu option_text trống
                opt_text = opt.get("option_text", "")
                if not opt_text and opt_spans:
                    opt_text = "".join(sp.get("text", "") for sp in opt_spans if sp.get("type") == "text")
                
                options.append({
                    "info": {
                        "spans": opt_spans,
                        "text": opt_text
                    }
                })
                if opt.get("is_correct"):
                    correct_index = i
            
            question = {
                "qid": q.get("qid_in_file", 0),
                "db_id": q.get("id"),
                "stem_text": q.get("stem_text", ""),
                "stem_media_spans": stem_spans,
                "stem_extra_media_spans": stem_extra,
                "options": options,
                "correct_index": correct_index,
                "diff_code": q.get("diff_code", ""),
                "source_file": f"[DB] {q.get('file_name', '')}",
                "from_db": True
            }
            converted.append(question)
        return converted

    def load_from_db(self):
        """Mở dialog chọn câu hỏi từ cơ sở dữ liệu để trộn đề."""
        banks = db.get_all_banks()
        if not banks:
            messagebox.showinfo("Thông báo", 
                "Chưa có ngân hàng nào trong CSDL.\n"
                "Hãy vào tab 'Ngân hàng' để import file Word vào DB trước.")
            return

        dlg = tk.Toplevel(self.parent)
        dlg.title("Chọn câu hỏi từ Cơ sở dữ liệu")
        utils.set_window_icon(dlg)
        dlg.geometry("950x600")
        dlg.transient(self.parent.winfo_toplevel())
        dlg.grab_set()

        # --- Khung trên: Bộ lọc ---
        frm_filter = ttk.LabelFrame(dlg, text="Bộ lọc")
        frm_filter.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(frm_filter, text="Môn học:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        var_subject = tk.StringVar()
        ttk.Entry(frm_filter, textvariable=var_subject, width=20).grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(frm_filter, text="Độ khó:").grid(row=0, column=2, padx=5, pady=5, sticky="e")
        var_diff = tk.StringVar()
        ttk.Entry(frm_filter, textvariable=var_diff, width=8).grid(row=0, column=3, padx=5, pady=5)

        ttk.Label(frm_filter, text="Từ khóa:").grid(row=0, column=4, padx=5, pady=5, sticky="e")
        var_kw = tk.StringVar()
        ttk.Entry(frm_filter, textvariable=var_kw, width=20).grid(row=0, column=5, padx=5, pady=5)

        # --- Danh sách ngân hàng (bên trái) ---
        frm_body = ttk.Frame(dlg)
        frm_body.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        frm_banks = ttk.LabelFrame(frm_body, text="Ngân hàng đã import")
        frm_banks.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        bank_cols = ("check", "name", "subject", "total", "pick")
        bank_tree = ttk.Treeview(frm_banks, columns=bank_cols, show="headings", height=15)
        bank_tree.heading("check", text="☑")
        bank_tree.heading("name", text="Tên ngân hàng")
        bank_tree.heading("subject", text="Môn học")
        bank_tree.heading("total", text="Tổng câu")
        bank_tree.heading("pick", text="Số lấy")
        bank_tree.column("check", width=35, anchor=tk.CENTER)
        bank_tree.column("name", width=250, anchor=tk.W)
        bank_tree.column("subject", width=120, anchor=tk.W)
        bank_tree.column("total", width=70, anchor=tk.CENTER)
        bank_tree.column("pick", width=70, anchor=tk.CENTER)

        vsb = ttk.Scrollbar(frm_banks, orient="vertical", command=bank_tree.yview)
        bank_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        bank_tree.config(yscrollcommand=vsb.set)

        # Trạng thái chọn và số câu
        bank_selection = {}  # bank_id -> {"selected": bool, "pick_count": int, "total": int}
        for b in banks:
            bank_selection[b["id"]] = {
                "selected": True,
                "pick_count": b["total_questions"],
                "total": b["total_questions"],
                "name": b["file_name"],
                "subject": b.get("subject_name", "")
            }

        def refresh_bank_tree():
            bank_tree.delete(*bank_tree.get_children())
            kw_filter = var_kw.get().strip().lower()
            subj_filter = var_subject.get().strip().lower()
            diff_filter = var_diff.get().strip()
            
            for b in banks:
                info = bank_selection[b["id"]]
                # Lọc theo môn học
                if subj_filter and subj_filter not in info["subject"].lower():
                    continue
                
                check_mark = "☑" if info["selected"] else "☐"
                bank_tree.insert("", "end", iid=str(b["id"]), values=(
                    check_mark, info["name"], info["subject"],
                    info["total"], info["pick_count"]
                ))

        refresh_bank_tree()

        def toggle_bank(event):
            sel = bank_tree.selection()
            if not sel: return
            bid = int(sel[0])
            col = bank_tree.identify_column(event.x)
            if col == "#1":  # Cột check
                bank_selection[bid]["selected"] = not bank_selection[bid]["selected"]
                refresh_bank_tree()
            elif col == "#5":  # Cột pick
                from tkinter import simpledialog
                c = simpledialog.askinteger(
                    "Số lượng", f"Số câu lấy từ '{bank_selection[bid]['name']}':",
                    initialvalue=bank_selection[bid]["pick_count"],
                    minvalue=0, maxvalue=bank_selection[bid]["total"]
                )
                if c is not None:
                    bank_selection[bid]["pick_count"] = c
                    refresh_bank_tree()

        bank_tree.bind("<Button-1>", toggle_bank)

        # Thông tin tổng hợp
        lbl_total = ttk.Label(dlg, text="", font=("Arial", 11, "bold"))
        lbl_total.pack(padx=10, pady=(0, 5), anchor="w")

        def update_total():
            total = sum(info["pick_count"] for info in bank_selection.values() if info["selected"])
            n_banks = sum(1 for info in bank_selection.values() if info["selected"])
            lbl_total.config(text=f"Tổng cộng: {total} câu từ {n_banks} ngân hàng")
        update_total()

        def confirm_load():
            selected_banks = [
                (bid, info) for bid, info in bank_selection.items() 
                if info["selected"] and info["pick_count"] > 0
            ]
            if not selected_banks:
                messagebox.showwarning("Cảnh báo", "Chưa chọn ngân hàng nào!", parent=dlg)
                return

            self.lbl_status.config(text="Đang nạp câu hỏi từ CSDL...")
            self.parent.update_idletasks()
            
            all_qs = []
            self.matrix_data = []
            
            for bid, info in selected_banks:
                # Lấy câu hỏi từ DB
                diff_filter = var_diff.get().strip()
                db_qs = db.search_questions(
                    bank_id=bid,
                    diff_code=diff_filter,
                    keyword=var_kw.get().strip(),
                    limit=9999
                )
                converted = self._convert_db_questions(db_qs)
                
                self.matrix_data.append({
                    "path": f"[DB] {info['name']}",
                    "questions": converted,
                    "pick_count": min(info["pick_count"], len(converted))
                })
                all_qs.extend(converted)
            
            if not all_qs:
                messagebox.showwarning("Cảnh báo", "Không tìm thấy câu hỏi nào với bộ lọc hiện tại!", parent=dlg)
                return

            self.all_questions = all_qs
            self.questions = all_qs.copy()
            self.selected_indices = list(range(len(all_qs)))
            self.pinned_indices = []
            self.src_doc = None  # Không có source document khi từ DB

            self.lbl_bank.config(text=f"💾 CSDL: {len(all_qs)} câu từ {len(selected_banks)} ngân hàng")
            self.btn_select_qs["state"] = "normal"
            self.btn_run["state"] = "normal"
            self.logmsg(f"💾 Đã nạp {len(all_qs)} câu hỏi từ CSDL ({len(selected_banks)} ngân hàng).")
            self.lbl_status.config(text="Sẵn sàng")
            dlg.destroy()

        frm_actions = ttk.Frame(dlg)
        frm_actions.pack(fill=tk.X, padx=10, pady=5)
        ttk.Button(frm_actions, text="🔍 Lọc", command=lambda: [refresh_bank_tree(), update_total()]).pack(side=tk.LEFT, padx=5)
        ttk.Button(frm_actions, text="☑ Chọn tất cả", 
                   command=lambda: [info.update({"selected": True}) for info in bank_selection.values()] or [refresh_bank_tree(), update_total()]).pack(side=tk.LEFT, padx=5)
        ttk.Button(frm_actions, text="☐ Bỏ chọn tất cả",
                   command=lambda: [info.update({"selected": False}) for info in bank_selection.values()] or [refresh_bank_tree(), update_total()]).pack(side=tk.LEFT, padx=5)
        ttk.Button(frm_actions, text="✅ Nạp câu hỏi đã chọn", command=confirm_load, bootstyle="success").pack(side=tk.RIGHT, padx=5)

    def _add_db_bank_to_matrix(self, matrix_data, refresh_callback):
        """Thêm ngân hàng từ CSDL vào ma trận đề."""
        banks = db.get_all_banks()
        if not banks:
            messagebox.showinfo("Thông báo", "Chưa có ngân hàng nào trong CSDL.")
            return

        dlg = tk.Toplevel(self.parent)
        dlg.title("Chọn ngân hàng từ CSDL")
        utils.set_window_icon(dlg)
        dlg.geometry("600x400")
        dlg.transient(self.parent.winfo_toplevel())
        dlg.grab_set()

        ttk.Label(dlg, text="Chọn ngân hàng để thêm vào ma trận:").pack(padx=10, pady=(10, 5), anchor="w")

        lb = tk.Listbox(dlg, selectmode=tk.MULTIPLE)
        lb.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        for b in banks:
            lb.insert(tk.END, f"{b['file_name']} ({b['total_questions']} câu, Môn: {b.get('subject_name', 'N/A')})")

        def on_add():
            sels = lb.curselection()
            if not sels: return
            for idx in sels:
                b = banks[idx]
                db_qs = db.search_questions(bank_id=b["id"], limit=9999)
                converted = self._convert_db_questions(db_qs)
                matrix_data.append({
                    "path": f"[DB] {b['file_name']}",
                    "questions": converted,
                    "pick_count": len(converted)
                })
            refresh_callback()
            dlg.destroy()

        ttk.Button(dlg, text="✅ Thêm vào ma trận", command=on_add).pack(pady=10)


    def select_questions(self):
        if not hasattr(self, 'all_questions') or not self.all_questions:
            return
            
        dlg = tk.Toplevel(self.parent)
        dlg.title("Chọn tập câu hỏi cho trộn đề")
        utils.set_window_icon(dlg)
        dlg.geometry("1100x700")
        dlg.transient(self.parent.winfo_toplevel())
        dlg.grab_set()

        frm_top = ttk.Frame(dlg)
        frm_top.pack(fill=tk.X, padx=10, pady=10)
        ttk.Label(frm_top, text="Click vào [ ] để chọn/bỏ chọn.").pack(side=tk.LEFT)
        
        lbl_count = ttk.Label(frm_top, text="", font=("Arial", 11, "bold"), foreground="blue")
        lbl_count.pack(side=tk.RIGHT)
        
        frm_body = ttk.Frame(dlg)
        frm_body.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        frm_left = ttk.Frame(frm_body)
        frm_left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        
        cols = ("check", "pinned", "question")
        tree = ttk.Treeview(frm_left, columns=cols, show="headings", selectmode="browse")
        tree.heading("check", text="Chọn")
        tree.heading("pinned", text="Ghim")
        tree.heading("question", text="Nội dung câu hỏi")
        tree.column("check", width=50, anchor=tk.CENTER, stretch=False)
        tree.column("pinned", width=50, anchor=tk.CENTER, stretch=False)
        tree.column("question", width=350, anchor=tk.W)
        
        vsb = ttk.Scrollbar(frm_left, orient="vertical", command=tree.yview)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        tree.config(yscrollcommand=vsb.set)
        
        frm_right = ttk.LabelFrame(frm_body, text="Chi tiết câu hỏi")
        frm_right.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))
        
        txt_detail = Text(frm_right, wrap="word", state="disabled", font=("Arial", 11))
        vsb_rt = ttk.Scrollbar(frm_right, orient="vertical", command=txt_detail.yview)
        txt_detail.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        vsb_rt.pack(side=tk.RIGHT, fill=tk.Y)
        txt_detail.config(yscrollcommand=vsb_rt.set)
        
        selected_set = set(getattr(self, 'selected_indices', list(range(len(self.all_questions)))))
        pinned_set = set(getattr(self, 'pinned_indices', []))
        
        def update_count_label():
            pinned_count = len(pinned_set)
            lbl_count.config(text=f"Đã chọn: {len(selected_set)} / {len(self.all_questions)} câu (Cố định: {pinned_count})")
            
        for i, q in enumerate(self.all_questions):
            text_preview = q['stem_text'][:80] + "..." if len(q['stem_text']) > 80 else q['stem_text']
            text_preview = text_preview.replace('\n', ' | ')
            if q.get("source_file"):
                text_preview = f"[{q['source_file']}] " + text_preview
            
            chk_sym = "✅" if i in selected_set else "⬜"
            pin_sym = "📌" if i in pinned_set else "➖"
            tree.insert("", tk.END, values=(chk_sym, pin_sym, f"Q{q['qid']}: {text_preview}"), tags=(str(i),))
            
        update_count_label()
            
        def on_tree_click(event):
            region = tree.identify("region", event.x, event.y)
            if region == "cell":
                col = tree.identify_column(event.x)
                if col == "#1":  # Clicked the checkbox column
                    row = tree.identify_row(event.y)
                    if not row: return
                    vals = list(tree.item(row, "values"))
                    idx = int(tree.item(row, "tags")[0])
                    
                    is_checked = (vals[0] == "[x]")
                    if is_checked:
                        vals[0] = "[ ]"
                        selected_set.discard(idx)
                    else:
                        vals[0] = "[x]"
                        selected_set.add(idx)
                    tree.item(row, values=vals)
                    update_count_label()

        def on_tree_select(event):
            sel = tree.selection()
            if not sel: return
            row = sel[0]
            idx = int(tree.item(row, "tags")[0])
            q = self.all_questions[idx]
            
            txt_detail.config(state="normal")
            txt_detail.delete("1.0", tk.END)
            
            if not hasattr(self, "preview_images"):
                self.preview_images = []
            self.preview_images.clear()

            def render_preview_spans(spans_list, txt_widget):
                for sp in spans_list:
                    if sp["type"] == "text":
                        # Hiển thị văn bản với các định dạng cơ bản
                        tag_name = ""
                        if sp.get("bold") and sp.get("italic"): tag_name = "bold_italic"
                        elif sp.get("bold"): tag_name = "bold"
                        elif sp.get("italic"): tag_name = "italic"
                        
                        txt_widget.insert(tk.END, sp["text"], tag_name)
                    elif sp["type"] == "image":
                        try:
                            pil_img = Image.open(io.BytesIO(sp["blob"]))
                            w, h = pil_img.size
                            max_w = 500  # tăng kích thước xem trước
                            if w > max_w:
                                pil_img = pil_img.resize((max_w, int(h * max_w / w)), Image.LANCZOS)
                            photo = ImageTk.PhotoImage(pil_img)
                            self.preview_images.append(photo)
                            txt_widget.insert(tk.END, "\n")
                            txt_widget.image_create(tk.END, image=photo)
                            txt_widget.insert(tk.END, "\n")
                        except Exception:
                            txt_widget.insert(tk.END, " [Lỗi hiển thị hình] ")
                    elif sp["type"] == "omml":
                        txt_widget.insert(tk.END, " [Công thức] ")
            
            txt_detail.insert(tk.END, f"Câu {q['qid']} [Độ khó: {q.get('diff_code') or 'N/A'}]\n\n", "bold_header")
            
            # Dòng 1 của thân câu đã có trong stem_media_spans (bao gồm text sau nhãn)
            if q.get("stem_media_spans"):
                render_preview_spans(q["stem_media_spans"], txt_detail)
            else:
                txt_detail.insert(tk.END, f"{q['stem_text']}\n")
                
            for extra_spans in q.get("stem_extra_media_spans", []):
                if extra_spans:
                    txt_detail.insert(tk.END, "\n")
                    render_preview_spans(extra_spans, txt_detail)
            txt_detail.insert(tk.END, "\n\n")
            
            letters = ["A", "B", "C", "D"]
            correct = q.get("correct_index", 0)
            
            for i, opt in enumerate(q.get("options", [])):
                info = opt.get("info", {})
                spans = info.get("spans", [])
                
                is_correct = (i == correct)
                prefix = f"{letters[i]}. "
                
                start_idx = txt_detail.index("end-1c")
                txt_detail.insert(tk.END, prefix)
                
                render_preview_spans(spans, txt_detail)
                
                if is_correct:
                    txt_detail.insert(tk.END, "  <-- ĐÚNG")
                    
                end_idx = txt_detail.index("end-1c")
                txt_detail.insert(tk.END, "\n")
                
                if is_correct:
                    txt_detail.tag_add("correct", start_idx, end_idx)
                    
            txt_detail.tag_config("bold_header", font=("Arial", 12, "bold"), foreground="#2C3E50")
            txt_detail.tag_config("bold", font=("Arial", 11, "bold"))
            txt_detail.tag_config("italic", font=("Arial", 11, "italic"))
            txt_detail.tag_config("bold_italic", font=("Arial", 11, "bold", "italic"))
            txt_detail.tag_config("correct", foreground="red", font=("Arial", 11, "bold"))
            txt_detail.config(state="disabled")

        tree.bind("<ButtonRelease-1>", on_tree_click)
        tree.bind("<<TreeviewSelect>>", on_tree_select)
        
        def on_ok():
            if not selected_set:
                messagebox.showwarning("Cảnh báo", "Bạn chưa chọn câu hỏi nào!")
                return
            self.selected_indices = sorted(list(selected_set))
            self.pinned_indices = sorted(list(pinned_set))
            self.questions = [self.all_questions[i] for i in self.selected_indices]
            self.btn_select_qs.config(text=f"Quỹ câu: {len(self.questions)} (Ghim {len(self.pinned_indices)})")
            self.logmsg(f"Đã cập nhật quỹ: {len(self.questions)} câu. Cố định: {len(self.pinned_indices)} câu.")
            dlg.destroy()
            
        def select_all():
            for row in tree.get_children():
                vals = list(tree.item(row, "values"))
                idx = int(tree.item(row, "tags")[0])
                vals[0] = "[x]"
                tree.item(row, values=vals)
                selected_set.add(idx)
            update_count_label()
            
        def deselect_all():
            for row in tree.get_children():
                vals = list(tree.item(row, "values"))
                idx = int(tree.item(row, "tags")[0])
                vals[0] = "[ ]"
                tree.item(row, values=vals)
                selected_set.discard(idx)
            update_count_label()

        frm_bot = ttk.Frame(dlg)
        frm_bot.pack(fill=tk.X, padx=10, pady=10)
        ttk.Button(frm_bot, text="Chọn tất cả", command=select_all).pack(side=tk.LEFT, padx=5)
        ttk.Button(frm_bot, text="Bỏ chọn tất cả", command=deselect_all).pack(side=tk.LEFT, padx=5)
        ttk.Button(frm_bot, text="Xác nhận", command=on_ok).pack(side=tk.RIGHT, padx=5)

    def run(self):
        if not self.questions:
            return
        try:
            n_tests  = int(self.var_num_tests.get().strip())
            n_qs     = int(self.var_num_questions.get().strip())
            code_sta = int(self.var_code_start.get().strip())
        except ValueError:
            messagebox.showerror("Lỗi", "Số đề, số câu, mã đề bắt đầu phải là số nguyên.")
            return
        if n_tests <= 0 or n_qs <= 0:
            messagebox.showerror("Lỗi", "Số đề/Số câu phải > 0.")
            return
        if n_qs > len(self.questions):
            messagebox.showerror("Lỗi", f"Số câu ({n_qs}) vượt quá số câu hợp lệ ({len(self.questions)}).")
            return

        school   = self.var_school.get().strip()
        faculty  = self.var_faculty.get().strip()
        ex_title = self.var_examtitle.get().strip()
        year_txt = self.var_year.get().strip()
        semester = self.var_semester.get().strip()
        class_name = self.var_class.get().strip()
        subject  = self.var_subject.get().strip()
        dur_txt  = self.var_duration.get().strip()
        tail_text= self.txt_tail.get("1.0", "end").strip()

        mode = self.var_export_mode.get()
        export_single   = (mode in ("Cả hai", "Chỉ từng đề"))
        export_aggregate= (mode in ("Cả hai", "Chỉ tổng hợp"))

        strategy = self.var_strategy.get()
        shuffle_answers = self.var_shuffle_ans.get()
        ans_layout = self.var_ans_layout.get()
        

        import re
        safe_subject = re.sub(r'[\\/*?:"<>|]', "", subject).strip()
        if not safe_subject:
            safe_subject = "Run"
        folder_name = f"{safe_subject} - {datetime.now().strftime('%d-%m-%Y_%Hh%Mm%Ss')}"
        
        out_dir = os.path.join(os.getcwd(), OUTPUT_FOLDER, folder_name)
        os.makedirs(out_dir, exist_ok=True)

        self.logmsg(f"Thư mục xuất: {out_dir}")
        self.logmsg(f"{school} | {faculty}")
        self.logmsg(f"{ex_title} – Năm học {year_txt} - Học kỳ {semester}")
        self.logmsg(f"Lớp: {class_name} - Môn: {subject} | {dur_txt}")
        self.logmsg(f"Tạo {n_tests} đề, mỗi đề {n_qs} câu. Mã đề bắt đầu: {code_sta}")
        self.logmsg(f"Chế độ xuất: {mode}")
        self.logmsg(f"Chiến lược chọn câu: {strategy} | Đảo đáp án: {'Có' if shuffle_answers else 'Không'}")

        # Lưu metadata.json
        try:
            import json
            meta = {
                "school": school,
                "faculty": faculty,
                "exam_title": ex_title,
                "school_year": year_txt,
                "semester": semester,
                "class_name": class_name,
                "subject": subject,
                "duration": dur_txt,
                "num_tests": n_tests,
                "num_questions": n_qs,
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            with open(os.path.join(out_dir, "metadata.json"), "w", encoding="utf-8") as f:
                json.dump(meta, f, ensure_ascii=False, indent=4)
        except Exception as e:
            self.logmsg(f"Cảnh báo: Không thể lưu metadata.json: {e}")

        answers_by_code: Dict[int, List[str]] = {}
        variants_db_data: List[Dict] = []  # Dữ liệu cho DB

        # Hoán vị nền (cho "Xoay vòng")
        rng_global = random.Random(20250923)
        base_perm = list(range(len(self.questions)))
        rng_global.shuffle(base_perm)

        # Khởi tạo thanh tiến độ
        self.progress["maximum"] = n_tests
        self.progress["value"] = 0
        self.lbl_status.config(text=f"Đang chuẩn bị {n_tests} đề...")
        self.parent.update_idletasks()

        all_doc = Document() if export_aggregate else None
        if export_aggregate:
            utils.set_page_layout(all_doc)
            utils.set_single_line_spacing(all_doc)
            # Bật phân biệt trang chẵn lẻ để hỗ trợ việc ngắt trang vật lý chính xác
            all_doc.settings.odd_and_even_pages_header_footer = True
            # Ép phần đầu tiên bắt đầu từ trang lẻ (mặc định là vậy nhưng đặt rõ cho chắc chắn)
            if all_doc.sections:
                all_doc.sections[0].start_type = WD_SECTION_START.ODD_PAGE

        for t in range(1, n_tests + 1):
            exam_code = code_sta + (t - 1)
            rng_exam = random.Random(10_000 + exam_code)

            # 1) Xử lý câu hỏi cố định (Pinned)
            pinned_qs = [self.all_questions[idx] for idx in self.pinned_indices]
            
            # Số câu cần lấy thêm từ quỹ đã chọn (loại bỏ những câu đã được ghim)
            needed = max(0, n_qs - len(pinned_qs))
            pool = [q for i, q in enumerate(self.questions) if self.selected_indices[i] not in self.pinned_indices]
            
            # 2) Chọn câu thêm theo chiến lược
            extra_qs = choose_questions(
                pool, needed, strategy=strategy,
                exam_idx=t, total_exams=n_tests, 
                base_perm=list(range(len(pool))), rng=rng_exam
            )
            
            # Trộn chung pinned và extra
            combined_qs = pinned_qs + extra_qs
            # Đảo thứ tự câu hỏi trong đề (đề phòng pinned luôn ở đầu)
            rng_exam.shuffle(combined_qs)

            # 3) Xây đề & (tuỳ chọn) đảo đáp án
            doc, qs_for_answers = build_exam_from_selected(
                combined_qs,
                school=school, faculty=faculty, subject=subject, duration_text=dur_txt,
                exam_title=ex_title, school_year=year_txt, exam_code=exam_code,
                tail_text=tail_text, shuffle_answers=shuffle_answers, seed=exam_code,
                ans_layout=ans_layout
            )

            # 4) Lưu
            if export_single:
                fn = os.path.join(out_dir, f"De_{exam_code}.docx")
                doc.save(fn)
                self.logmsg(f"Đã ghi {fn}")

            # 4) Gom đáp án để xuất Excel
            letters = ["A", "B", "C", "D"]
            answers = [letters[q["correct_index"]] for q in qs_for_answers]
            answers_by_code[exam_code] = answers

            # Thu thập dữ liệu cho DB
            variant_qs_for_db = []
            for pos, q in enumerate(qs_for_answers, 1):
                variant_qs_for_db.append({
                    "question_db_id": q.get("db_id"),  # Có thể None nếu chưa import DB
                    "position": pos,
                    "correct_label": letters[q["correct_index"]] if q["correct_index"] < len(letters) else ""
                })
            variants_db_data.append({
                "exam_code": exam_code,
                "questions": variant_qs_for_db
            })

            # 5) Tổng hợp (nếu bật)
            if export_aggregate:
                add_header_2cols(all_doc, school, faculty, ex_title, year_txt, subject, dur_txt, exam_code)
                for i, q in enumerate(qs_for_answers, 1):
                    render_question(all_doc, i, q, ans_layout=ans_layout)
                add_exam_end(all_doc, tail_text=tail_text, use_section_pages=True)

                # Chỉ ngắt section nếu chưa phải đề cuối cùng
                if t < n_tests:
                    # Ngắt section (bắt đầu ở trang Lẻ) để đề tiếp theo luôn bắt đầu ở mặt trước tờ giấy mới
                    new_sec = all_doc.add_section(WD_SECTION_START.ODD_PAGE)
                    
                    # Khởi tạo lại số trang bắt đầu từ 1 cho mỗi đề mới
                    sectPr = new_sec._sectPr
                    pgNumType = sectPr.find(qn('w:pgNumType'))
                    if pgNumType is None:
                        pgNumType = OxmlElement('w:pgNumType')
                        pgSz = sectPr.find(qn('w:pgSz'))
                        if pgSz is not None: pgSz.addprevious(pgNumType)
                        else: sectPr.append(pgNumType)
                    pgNumType.set(qn('w:start'), '1')

            # Cập nhật tiến độ
            self.progress["value"] = t
            self.lbl_status.config(text=f"Đang trộn: {t}/{n_tests} đề...")
            self.parent.update_idletasks()

        if export_aggregate:
            # Khôi phục đánh số trang footer cho toàn bộ đề trong file tổng hợp
            utils.add_page_number_to_footer(all_doc, use_section_pages=True)
            # Áp dụng lại layout (Lề đối xứng) cho toàn bộ các đề
            utils.set_page_layout(all_doc)
            all_fn = os.path.join(out_dir, "Tong_hop_de.docx")
            all_doc.save(all_fn)
            self.logmsg(f"Đã ghi {all_fn}")

        # Xuất dap_an.xlsx
        cols = ["Mã đề"] + [f"Câu {i}" for i in range(1, n_qs + 1)]
        rows = []
        for code in sorted(answers_by_code.keys()):
            row = [code] + answers_by_code[code]
            if len(row) < len(cols):
                row += [""] * (len(cols) - len(row))
            rows.append(row)
        df = pd.DataFrame(rows, columns=cols)
        ans_path = os.path.join(out_dir, "dap_an.xlsx")
        df.to_excel(ans_path, index=False)
        self.logmsg(f"Đã ghi {ans_path}")
        
        # Lưu vào cơ sở dữ liệu
        try:
            db.save_exam_session(
                school=school, faculty=faculty, school_year=year_txt,
                semester=semester, class_name=class_name, exam_title=ex_title,
                subject_name=subject, duration=dur_txt,
                num_variants=n_tests, num_questions=n_qs,
                folder_path=out_dir, strategy=strategy,
                shuffle_answers=shuffle_answers,
                variants_data=variants_db_data,
                metadata_json=json.dumps(meta, ensure_ascii=False) if 'meta' in locals() else '{}'
            )
            self.logmsg("✅ Đã lưu thông tin đợt trộn vào cơ sở dữ liệu.")
        except Exception as e:
            self.logmsg(f"⚠️ Cảnh báo: Không thể lưu vào DB: {e}")
        
        self.lbl_status.config(text="Hoàn tất.")
        self.parent.update_idletasks()
        self.logmsg("Hoàn tất.")
        
        # Tự động mở thư mục kết quả
        try:
            os.startfile(out_dir)
        except Exception:
            pass

