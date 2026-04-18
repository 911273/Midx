import os
import ctypes
import io
import copy
import json
import re
import tkinter as tk
from tkinter import ttk
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.run import Run
from typing import List, Dict, Any, Optional, Tuple
import base64

# Đảm bảo namespace Toán học (OMML) được đăng ký
from docx.oxml.ns import nsmap
nsmap['m'] = "http://schemas.openxmlformats.org/officeDocument/2006/math"
# ============= CẤU HÌNH & REGEX DÙNG CHUNG =============
QUESTION_HEADER_PAT = re.compile(r"^\s*Câu\s*(\d+)[\s:.-]*\s*(?:\[\s*<([A-Z0-9]{1,3})>\s*\])?[\s:.-]*\s*(.*)$", re.IGNORECASE)
# Regex nhận diện phần nhãn/tiền tố của phương án (hỗ trợ [<$>], A., - A., v.v.)
_LABEL_REGEX = re.compile(
    r'^\s*(?:\[\s*\<\s*\$\s*\>\s*\]\s*(?:[-•●▪▫]?\s*)?(?:(?:\([A-Fa-f]\)|[A-Fa-f]\.)\s*)?|' # TH 1: Có [<$>]
    r'(?:[-•●▪▫]?\s*)?(?:\([A-Fa-f]\)|[A-Fa-f]\.)\s*)' # TH 2: Chỉ có nhãn A./B./...
)
RAW_OPTION_MARKER = re.compile(r"\[\s*\<\s*\$\s*\>\s*\]\s*")

# ============= TIỆN ÍCH PHÂN TÍCH WORD =============
def remove_marker_text(s: str, trim: bool=False) -> str:
    t = RAW_OPTION_MARKER.sub("", s or "").replace("[<$>]", "")
    return t.strip() if trim else t

def parse_question_header(text: str):
    m = QUESTION_HEADER_PAT.match((text or "").strip())
    if not m: return None, None, None
    qid, diff_code, stem = m.groups()
    return int(qid), diff_code, (stem or "").strip()

def get_spans_text_summary(spans: List[Dict[str, Any]]) -> str:
    """Trả về văn bản tổng hợp từ danh sách spans, hỗ trợ mô phỏng công thức OMML."""
    if not spans: return ""
    res = []
    
    def _xml_to_pseudo_latex(node):
        if node is None: return ""
        out = ""
        tag = node.tag.rsplit('}', 1)[-1]
        if tag == 't': return node.text or ""
        elif tag == 'f':
            num = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="num"]'))
            den = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="den"]'))
            return f"({num}/{den})"
        elif tag == 'rad':
            e = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="e"]'))
            deg = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="deg"]'))
            return f"root({deg},{e})" if deg and deg.strip() else f"sqrt({e})"
        elif tag == 'sup':
            e = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="e"]'))
            sup = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="sup"]'))
            return f"{e}^{{{sup}}}"
        
        # Nếu không phải tag đặc biệt, hãy lấy tất cả text thô bên trong nó
        # (Đảm bảo không bỏ sót văn bản trong các tag m:r, m:t...)
        raw_t = "".join(t.text for t in node.xpath('.//*[local-name()="t"]') if t.text)
        if raw_t: return raw_t

        for child in node.iterchildren():
            out += _xml_to_pseudo_latex(child)
        return out

    for sp in spans:
        stype = sp.get("type", "text")
        if stype == "text":
            res.append(sp.get("text", ""))
        elif stype == "omml":
            xml_node = sp.get("xml")
            raw_txt = ""
            if xml_node is not None:
                try: raw_txt = _xml_to_pseudo_latex(xml_node)
                except: pass
            if raw_txt: res.append(f" [{raw_txt}] ")
        elif stype == "image":
            res.append(" [Ảnh] ")
    
    return "".join(res).strip()

def get_run_color(run) -> str:
    """Trả về định danh màu sắc của run (Hex, Highlighter hoặc Theme)."""
    if getattr(run.font, "highlight_color", None) is not None:
        return f"HL_{run.font.highlight_color}"
    
    col = getattr(run.font, "color", None)
    if not col: return "NONE"
    
    rgb = getattr(col, "rgb", None)
    if rgb is not None:
        s = str(rgb).replace("#","").upper()
        if s == "000000": return "NONE" # Coi đen là không màu
        return s
        
    theme = getattr(col, "theme_color", None)
    if theme is not None:
        # 1 là Black, 2 là White trong mặc định Office
        if theme in (1, 2): return "NONE"
        return f"T_{theme}"
    
    return "NONE"

def label_span_end(paragraph) -> int:
    txt = paragraph.text or ""
    m = _LABEL_REGEX.match(txt)
    return m.end() if m else -1

# ============= ĐỊNH DẠNG TRANG & STYLE =============
def set_page_layout(doc):
    for sec in doc.sections:
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)
        sec.gutter = Cm(0.5)
        
        # Bật lề đối xứng (Mirror Margins) qua OXML
        sectPr = sec._sectPr
        mirrorMargins = sectPr.find(qn('w:mirrorMargins'))
        if mirrorMargins is None:
            mirrorMargins = OxmlElement('w:mirrorMargins')
            sectPr.append(mirrorMargins)
    
    # Bật chế độ phân biệt Header/Footer trang chẵn lẻ cho tài liệu
    doc.settings.odd_and_even_pages_header_footer = True

def set_single_line_spacing(doc, font_name='Times New Roman'):
    """
    Thiết lập giãn dòng đơn và font chữ chuẩn cho toàn bộ tài liệu.
    Loại bỏ hoàn toàn Space After.
    """
    try:
        # 1. Cấu hình style Normal
        style = doc.styles['Normal']
        pf = style.paragraph_format
        pf.line_spacing = 1.0
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        style.font.name = font_name
        rPr = style._element.get_or_add_rPr()
        
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)

        for k in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rFonts.set(qn(k), font_name)
            
        # 2. Thiết lập Document Defaults (quan trọng cho Header/Footer và Table)
        try:
            # Truy cập trực tiếp vào OXML của styles
            styles_element = doc.styles.element
            doc_defaults = styles_element.find(qn('w:docDefaults'))
            if doc_defaults is not None:
                rPr_def = doc_defaults.find(qn('w:rPrDefault'))
                if rPr_def is not None:
                    rPr_inner = rPr_def.find(qn('w:rPr'))
                    if rPr_inner is not None:
                        # Kiểm tra và lấy w:rFonts (get_or_add_rFonts không tồn tại trên OxmlElement thường)
                        fonts = rPr_inner.find(qn('w:rFonts'))
                        if fonts is None:
                            fonts = OxmlElement('w:rFonts')
                            rPr_inner.append(fonts)
                        for k in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                            fonts.set(qn(k), font_name)
        except Exception:
            pass
            
    except Exception:
        pass

def add_horizontal_rule(doc, thickness_eighths=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness_eighths))
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV'):
        edge = OxmlElement(side)
        edge.set(qn('w:val'), 'nil')
        tblBorders.append(edge)
    tblPr.append(tblBorders)
def set_run_font(run, font_name='Times New Roman'):
    """Thiết lập font chữ cho một run, bao gồm cả các thuộc tính OXML để hỗ trợ đa ngôn ngữ."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for k in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(k), font_name)

def add_styled_text(p, text, size=12, bold=False, italic=False, align=None, uppercase=False, color_rgb=None, font_name='Times New Roman'):
    if align is not None:
        p.alignment = align
    r = p.add_run(text.upper() if uppercase else text)
    set_run_font(r, font_name)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color_rgb:
        r.font.color.rgb = color_rgb
    return r

# ============= XỬ LÝ MEDIA (ẢNH & CÔNG THỨC) =============
def append_omml(dst_para, xml_elem):
    """
    Chèn công thức toán học vào paragraph.
    Xử lý đặc biệt cho oMathPara để tránh việc lồng paragraph không hợp lệ.
    """
    try:
        tag = xml_elem.tag.rsplit('}', 1)[-1]
        
        # Hàm hỗ trợ chèn một node OMML và dọn dẹp XML
        def _append_node(parent, node):
            # Tạo bản sao sâu để tránh lỗi liên kết document
            new_node = copy.deepcopy(node)
            parent.append(new_node)

        if tag == 'oMathPara':
            # oMathPara có thể chứa nhiều oMath (vd: hệ phương trình)
            maths = xml_elem.xpath('.//*[local-name()="oMath"]')
            if maths:
                for m in maths:
                    _append_node(dst_para._element, m)
            else:
                # Fallback nếu không dùng xpath được
                for child in xml_elem.iterchildren():
                    _append_node(dst_para._element, child)
        else:
            _append_node(dst_para._element, xml_elem)
            
    except Exception as e:
        # Nếu lỗi XML, thử ghi nhận vào log hoặc chèn thông báo ẩn
        pass

def append_image(dst_para, blob, width_cm=None):
    """
    Chèn hình ảnh vào paragraph với cơ chế fallback sang PNG nếu lỗi.
    Nếu width_cm là None, sử dụng kích thước gốc của ảnh.
    """
    try:
        width = Cm(width_cm) if width_cm else None
        dst_para.add_run().add_picture(io.BytesIO(blob), width=width)
    except Exception:
        try:
            from PIL import Image
            pil_img = Image.open(io.BytesIO(blob))
            if pil_img.mode in ("RGBA", "P"):
                pil_img = pil_img.convert("RGB")
            tmp_img = io.BytesIO()
            pil_img.save(tmp_img, format="PNG")
            tmp_img.seek(0)
            dst_para.add_run().add_picture(tmp_img, width=Cm(width_cm))
        except Exception as e:
            dst_para.add_run(f" [Lỗi ảnh: {str(e)[:20]}...] ")

def render_spans_into_paragraph(dst_para, spans: List[Dict[str, Any]], font_size_pt=11, default_black=True):
    """
    Render danh sách spans (text/image/omml) vào paragraph.
    Gom nhóm text trùng định dạng để tối ưu file output.
    """
    first_text_written = False
    for sp in spans:
        stype = sp.get("type")
        if stype == "omml":
            append_omml(dst_para, sp["xml"])
        elif stype == "image":
            append_image(dst_para, sp["blob"])
        elif stype == "text":
            txt = sp.get("text", "")
            if txt:
                if not first_text_written:
                    txt = txt.lstrip()
                    first_text_written = True
                r = dst_para.add_run(txt)
                r.bold = bool(sp.get("bold"))
                r.italic = bool(sp.get("italic"))
                r.underline = bool(sp.get("underline"))
                r.font.subscript = bool(sp.get("subscript"))
                r.font.superscript = bool(sp.get("superscript"))
                
                # Chỉ ép font nếu span có thông tin font cụ thể, tránh ghi đè font mặc định của Style
                f_name = sp.get("font_name")
                if f_name:
                    set_run_font(r, f_name)
                
                r.font.size = Pt(font_size_pt)

                if default_black and sp.get("strip_black", True):
                    r.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    color_hex = sp.get("color_hex")
                    if color_hex:
                        try:
                            r.font.color.rgb = RGBColor.from_string(color_hex)
                        except Exception:
                            pass


# ============= HỖ TRỢ SOẠN THẢO (EDITOR) =============
def render_latex_to_png(latex_str: str, font_size: int = 12, dpi: int = 150) -> Optional[bytes]:
    """Render LaTeX string thành bytes hình ảnh PNG bằng matplotlib."""
    if not latex_str:
        return None
    
    # Đảm bảo mã LaTeX nằm trong dấu $$
    trimmed = latex_str.strip()
    if not trimmed.startswith("$"):
        trimmed = f"${trimmed}$"

    try:
        import matplotlib
        matplotlib.use('Agg') # Tránh hiện cửa sổ GUI
        import matplotlib.pyplot as plt
        from io import BytesIO

        # Thiết lập font và style
        fig = plt.figure(figsize=(0.1, 0.1)) # Kích thước tạm
        plt.axis('off')
        
        # Render
        t = plt.text(0, 0, trimmed, fontsize=font_size)
        
        # Tính toán kích thước tự động
        fig.canvas.draw()
        bbox = t.get_window_extent()
        width, height = bbox.width / dpi, bbox.height / dpi
        
        # Cập nhật kích thước thực tế
        fig.set_size_inches(width, height)
        
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=dpi, transparent=True, bbox_inches='tight', pad_inches=0.05)
        plt.close(fig)
        return buf.getvalue()
    except Exception as e:
        print(f"Lỗi render LaTeX: {e}")
        return None


def extract_spans_from_text_widget(text_widget, media_map: Dict[str, Dict]) -> List[Dict]:
    """
    Trích xuất danh sách spans (text, image) từ widget tk.Text.
    media_map: Dict ánh xạ tên ảnh trong widget (vd: 'pyimage1') sang dữ liệu {type: 'image'/'latex', blob: bytes, latex_code: str}
    """
    spans = []
    
    # tk.Text.dump trả về danh sách (key, value, index)
    # Ví dụ: ('text', 'Nội dung ', '1.0'), ('image', 'pyimage1', '1.8')
    content = text_widget.dump("1.0", "end-1c", text=True, image=True)
    
    for key, value, index in content:
        if key == "text":
            if value:
                spans.append({
                    "type": "text",
                    "text": value,
                    "bold": False,
                    "italic": False,
                    "underline": False,
                    "color_hex": None
                })
        elif key == "image":
            image_name = str(value)
            if image_name in media_map:
                m = media_map[image_name]
                if m["type"] == "image":
                    spans.append({
                        "type": "image",
                        "blob": m["blob"]
                    })
                elif m["type"] == "latex":
                    # Lưu công thức dưới dạng OMML nếu có thể, hoặc đơn giản là lưu ảnh render
                    # Hiện tại logic trộn đề hỗ trợ OMML (xml) và image (blob).
                    # Ta lưu ảnh render để đảm bảo tương thích mọi nơi.
                    spans.append({
                        "type": "image",
                        "blob": m["blob"],
                        "latex_code": m.get("latex_code") # Lưu kèm để sau này sửa đổi
                    })
    
    # Hậu xử lý: Gộp các span văn bản liên tiếp
    merged = []
    for s in spans:
        if merged and merged[-1]["type"] == "text" and s["type"] == "text":
            merged[-1]["text"] += s["text"]
        else:
            merged.append(s)
            
    return merged

# ============= TIỆN ÍCH KHÁC =============
def get_clean_text(text: str) -> str:
    """Loại bỏ các ký tự điều khiển hoặc marker đặc biệt."""
    if not text: return ""
    return text.replace("[<$>]", "").strip()

# ============= ĐÁNH SỐ TRANG & FIELD CODES =============
def _insert_field(run, field_code):
    """
    Chèn một Word Field Code (ví dụ: PAGE, NUMPAGES) vào một run.
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f" {field_code} " # Thêm khoảng trắng để Word nhận diện chuẩn hơn

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    # Chèn một placeholder tạm thời (khoảng trắng) để mã trường không bị lộ
    t = OxmlElement('w:t')
    t.text = " "

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)

def _add_num_to_footer_obj(footer, use_section_pages):
    # Xóa các paragraph cũ nếu có
    for p in footer.paragraphs:
        p.text = ""
    
    # Nếu chưa có paragraph nào thì thêm mới
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thêm text "Trang "
    r = p.add_run("Trang ")
    set_run_font(r, 'Times New Roman')
    r.font.size = Pt(10)
    
    # Chèn số trang hiện tại { PAGE }
    r_page = p.add_run()
    set_run_font(r_page, 'Times New Roman')
    r_page.font.size = Pt(10)
    _insert_field(r_page, 'PAGE')
    
    # Thêm text " / "
    r = p.add_run(" / ")
    set_run_font(r, 'Times New Roman')
    r.font.size = Pt(10)
    
    # Chèn tổng số trang { NUMPAGES } hoặc { SECTIONPAGES }
    r_total = p.add_run()
    set_run_font(r_total, 'Times New Roman')
    r_total.font.size = Pt(10)
    total_field = 'SECTIONPAGES' if use_section_pages else 'NUMPAGES'
    _insert_field(r_total, total_field)

def add_page_number_to_footer(doc, use_section_pages=False):
    """
    Thêm đánh số trang "Trang X / Y" vào footer của tất cả các section.
    Áp dụng cho cả trang chẵn và lẻ.
    """
    for section in doc.sections:
        # Áp dụng cho footer mặc định (trang lẻ)
        _add_num_to_footer_obj(section.footer, use_section_pages)
        # Áp dụng cho footer trang chẵn
        _add_num_to_footer_obj(section.even_page_footer, use_section_pages)

def get_assets_path():
    """Lấy đường dẫn thư mục assets, hỗ trợ PyInstaller."""
    try:
        import sys
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, "assets")

def set_window_icon(window):
    """Thiết lập logo cho cửa sổ tkinter và thanh taskbar."""
    try:
        # Đăng ký AppUserModelID để Windows nhận diện icon ứng dụng trên taskbar
        # Thay vì hiển thị icon mặc định của Python
        myappid = 'vupq.epu.mide.v9' 
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
    except Exception:
        pass

    assets_dir = get_assets_path()
    ico_path = os.path.join(assets_dir, "logo.ico")
    jpeg_path = os.path.join(assets_dir, "logo.jpeg")

    # 1. Thử thiết lập iconbitmap bằng file .ico (chuẩn cho Windows title bar)
    if os.path.exists(ico_path):
        try:
            window.iconbitmap(ico_path)
        except Exception:
            pass

    # 2. Sử dụng ImageTk làm fallback hoặc để đồng bộ cho iconphoto
    if os.path.exists(jpeg_path):
        try:
            from PIL import Image, ImageTk
            icon_image = Image.open(jpeg_path)
            # Resize để đảm bảo tương thích tốt
            icon_image = icon_image.resize((64, 64))
            icon_photo = ImageTk.PhotoImage(icon_image)
            
            # iconphoto(True, ...) áp dụng cho cả các cửa sổ con và thanh taskbar
            window.iconphoto(True, icon_photo)
            
            # Giữ tham chiếu để tránh bị Garbage Collection
            window._icon_photo = icon_photo 
        except Exception as e:
            print(f"Lỗi khi tải logo (ImageTk): {e}")

# ============= QUẢN LÝ CẤU HÌNH & TIỆN ÍCH MỚI =============
CONFIG_FILE = "config.json"

# Quy chuẩn giao diện (UI Standards)
PAD_XS = 2
PAD_S = 5
PAD_M = 10
PAD_L = 15

FONT_MAIN = ("Arial", 10)
FONT_BOLD = ("Arial", 10, "bold")
FONT_HEADER = ("Arial", 14, "bold")

def load_config() -> Dict[str, Any]:
    """Tải cấu hình người dùng từ file config.json."""
    default_config = {
        "theme": "litera",
        "last_out_dir": os.path.join(os.getcwd(), "DE_TRON"),
        "school": "ĐẠI HỌC ĐIỆN LỰC",
        "faculty": "KHOA NĂNG LƯỢNG MỚI"
    }
    if not os.path.exists(CONFIG_FILE):
        return default_config
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            user_config = json.load(f)
            # Merge với mặc định để tránh thiếu key
            for k, v in default_config.items():
                if k not in user_config:
                    user_config[k] = v
            return user_config
    except Exception:
        return default_config

def save_config(config: Dict[str, Any]):
    """Lưu cấu hình người dùng xuống file config.json."""
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=4, ensure_ascii=False)
    except Exception:
        pass

def setup_dialog(win, width_pct=0.7, height_pct=0.7, title=None, parent=None):
    """
    Khởi tạo cửa sổ con (dialog) với các thiết lập chuẩn:
    - Title, Icon
    - Kích thước theo % màn hình
    - Căn giữa màn hình
    - Modal (nếu có parent)
    """
    if title: win.title(title)
    set_window_icon(win)
    
    win.update_idletasks()
    sw = win.winfo_screenwidth()
    sh = win.winfo_screenheight()
    
    w = int(sw * width_pct)
    h = int(sh * height_pct)
    
    # Giới hạn kích thước hợp lý cho desktop
    w = min(max(w, 800), 1600)
    h = min(max(h, 600), 900)
    
    x = (sw - w) // 2
    y = (sh - h) // 3
    
    win.geometry(f"{w}x{h}+{x}+{y}")
    
    if parent:
        win.transient(parent.winfo_toplevel())
        win.grab_set()

def center_window(win, width=1280, height=750):
    """Căn giữa cửa sổ trên màn hình, có tính đến scaling."""
    win.update_idletasks()
    sw = win.winfo_screenwidth()
    sh = win.winfo_screenheight()
    
    # Nếu cửa sổ đã được scale qua 'tk scaling', winfo_screenwidth có thể trả về giá trị đã scale
    # Ta cố gắng căn giữa tương đối chính xác nhất
    x = (sw - width) // 2
    y = max((sh - height) // 3, 40)
    win.geometry(f"{width}x{height}+{x}+{y}")


# ============= GIAO DIỆN XEM TRƯỚC CÂU HỎI (DÙNG CHUNG) =============
class QuestionPreviewDialog:
    def __init__(self, parent, questions: List[Dict], title="Xem trước nội dung"):
        self.parent = parent
        self.questions = questions
        self.preview_images = []
        
        self.win = tk.Toplevel(parent)
        setup_dialog(self.win, width_pct=0.75, height_pct=0.85, title=title, parent=parent)
        
        self.setup_ui()
        if self.questions:
            self.lb.selection_set(0)
            self.show_q(None)

    def setup_ui(self):
        frm_body = ttk.Frame(self.win)
        frm_body.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Sidebar: Danh sách câu
        frm_list = ttk.Frame(frm_body)
        frm_list.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 5))
        
        ttk.Label(frm_list, text="Danh sách câu:", font=FONT_BOLD).pack(anchor="w", pady=(0, 5))
        
        self.lb = tk.Listbox(frm_list, width=20, font=("Arial", 10))
        self.lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(frm_list, orient="vertical", command=self.lb.yview)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.lb.config(yscrollcommand=vsb.set)
        
        for i, q in enumerate(self.questions, 1):
            # Hiển thị số câu và mã ID nếu có
            label = f"Câu {i}"
            if q.get('qid'): label += f" (ID: {q['qid']})"
            self.lb.insert(tk.END, label)
        
        # Content: Chi tiết câu hỏi
        frm_content = ttk.Frame(frm_body)
        frm_content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.txt_detail = tk.Text(frm_content, wrap="word", font=("Arial", 11), state="disabled", padx=10, pady=10)
        self.txt_detail.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb_rt = ttk.Scrollbar(frm_content, orient="vertical", command=self.txt_detail.yview)
        vsb_rt.pack(side=tk.RIGHT, fill=tk.Y)
        self.txt_detail.config(yscrollcommand=vsb_rt.set)
        
        # Tags định dạng
        self.txt_detail.tag_config("bold_header", font=("Arial", 12, "bold"), foreground="#2C3E50")
        self.txt_detail.tag_config("bold", font=("Arial", 11, "bold"))
        self.txt_detail.tag_config("italic", font=("Arial", 11, "italic"))
        self.txt_detail.tag_config("bold_italic", font=("Arial", 11, "bold", "italic"))
        self.txt_detail.tag_config("correct", foreground="red", font=("Arial", 11, "bold"))
        self.txt_detail.tag_config("sub", font=("Arial", 8), offset=-4)
        self.txt_detail.tag_config("sup", font=("Arial", 8), offset=4)
        
        self.lb.bind("<<ListboxSelect>>", self.show_q)

    def render_spans(self, spans_list):
        """Render danh sách spans vào widget Text."""
        for sp in spans_list:
            stype = sp.get("type")
            if stype == "text":
                tags = []
                if sp.get("bold") and sp.get("italic"): tags.append("bold_italic")
                elif sp.get("bold"): tags.append("bold")
                elif sp.get("italic"): tags.append("italic")
                
                if sp.get("subscript"): tags.append("sub")
                elif sp.get("superscript"): tags.append("sup")
                
                self.txt_detail.insert(tk.END, sp.get("text", ""), tags)
            elif stype == "image":
                try:
                    from PIL import Image, ImageTk
                    pil_img = Image.open(io.BytesIO(sp["blob"]))
                    w, h = pil_img.size
                    # Tự động resize phù hợp khu vực hiển thị
                    limit = 650
                    if w > limit:
                        pil_img = pil_img.resize((limit, int(h * limit / w)), Image.LANCZOS)
                    photo = ImageTk.PhotoImage(pil_img)
                    self.preview_images.append(photo)
                    self.txt_detail.insert(tk.END, "\n")
                    self.txt_detail.image_create(tk.END, image=photo)
                    self.txt_detail.insert(tk.END, "\n")
                except Exception:
                    self.txt_detail.insert(tk.END, " [Lỗi hiển thị ảnh] ")
            elif stype == "omml":
                # Trình phân tích thô để tạo chuỗi văn bản mô phỏng công thức
                xml_node = sp.get("xml")
                def _xml_to_pseudo_latex(node):
                    if node is None: return ""
                    res = ""
                    tag = node.tag.rsplit('}', 1)[-1]
                    
                    if tag == 't': # Văn bản toán học
                        return node.text or ""
                    elif tag == 'f': # Phân số (fraction)
                        # Thường có m:num và m:den
                        num = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="num"]'))
                        den = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="den"]'))
                        return f"({num}/{den})"
                    elif tag == 'rad': # Căn thức (radical)
                        # Thường có m:e (nội dung) và m:deg (bậc căn)
                        e = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="e"]'))
                        deg = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="deg"]'))
                        if deg and deg.strip(): return f"root({deg}, {e})"
                        return f"sqrt({e})"
                    elif tag == 'sup': # Số mũ
                        e = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="e"]'))
                        sup = "".join(_xml_to_pseudo_latex(c) for c in node.xpath('.//*[local-name()="sup"]'))
                        return f"{e}^{{{sup}}}"
                    
                    # Duyệt các con nếu không phải tag đặc biệt
                    for child in node.iterchildren():
                        res += _xml_to_pseudo_latex(child)
                    return res

                raw_txt = ""
                if xml_node is not None:
                    try: raw_txt = _xml_to_pseudo_latex(xml_node)
                    except: pass
                
                display_txt = f" [{raw_txt}] " if raw_txt else " [Công thức] "
                self.txt_detail.insert(tk.END, display_txt, "bold")

    def show_q(self, event):
        sel = self.lb.curselection()
        if not sel: return
        idx = sel[0]
        q = self.questions[idx]
        
        self.txt_detail.config(state="normal")
        self.txt_detail.delete("1.0", tk.END)
        self.preview_images.clear()
        
        # Header
        diff = q.get('diff_code') or 'N/A'
        self.txt_detail.insert(tk.END, f"Câu số {idx+1} [Độ khó: {diff}]\n\n", "bold_header")
        
        # Thân câu hỏi (Media spans)
        if q.get("stem_media_spans"):
            self.render_spans(q["stem_media_spans"])
        else:
            self.txt_detail.insert(tk.END, q.get("stem_text", ""))
        
        # Các đoạn extra (nếu có)
        for extra in q.get("stem_extra_media_spans", []):
            self.txt_detail.insert(tk.END, "\n")
            self.render_spans(extra)
        
        self.txt_detail.insert(tk.END, "\n\n")
        
        # Phương án
        letters = ["A", "B", "C", "D", "E", "F"]
        options = q.get("options", [])
        correct_idx = q.get("correct_index")
        
        for i, opt in enumerate(options):
            is_correct = (i == correct_idx)
            start_ptr = self.txt_detail.index("end-1c")
            
            self.txt_detail.insert(tk.END, f"{letters[i]}. " if i < len(letters) else "- ")
            
            # opt có thể là dict {"info": {"spans": ...}} hoặc trực tiếp cấu trúc span
            spans = []
            if isinstance(opt, dict) and "info" in opt:
                spans = opt["info"].get("spans", [])
            elif isinstance(opt, dict) and "spans" in opt:
                spans = opt.get("spans", [])
            
            self.render_spans(spans)
            
            if is_correct:
                self.txt_detail.insert(tk.END, "  <-- CHÍNH XÁC")
                self.txt_detail.tag_add("correct", start_ptr, self.txt_detail.index("end-1c"))
            
            self.txt_detail.insert(tk.END, "\n")
        
        self.txt_detail.config(state="disabled")


