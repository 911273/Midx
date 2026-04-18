# -*- coding: utf-8 -*-
import os
import shutil
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from datetime import datetime
import io
import copy
from docx import Document

from docx.shared import Pt, RGBColor, Cm
import utils
import db

try:
    from tkinterdnd2 import DND_FILES
    HAS_DND = True
except ImportError:
    HAS_DND = False

# Giả định tron_de.py nằm cùng thư mục hoặc trong path
try:
    from tron_de import split_questions_from_docx
except ImportError:
    split_questions_from_docx = None

BANK_DIR = "NGAN_HANG_CAU_HOI"

class BankManagerTab:
    def __init__(self, parent):
        self.parent = parent
        os.makedirs(BANK_DIR, exist_ok=True)
        
        # Bố cục chính
        self.parent.rowconfigure(1, weight=1)  # Treeview co giãn
        self.parent.columnconfigure(0, weight=1)
        
        # --- Container cho Toolbar ---
        frm_controls = ttk.Frame(self.parent)
        frm_controls.grid(row=0, column=0, sticky="ew", padx=utils.PAD_M, pady=utils.PAD_S)
        
        # Nhóm quản lý File
        lf_files = ttk.LabelFrame(frm_controls, text="Quản lý File Word", padding=utils.PAD_S)
        lf_files.pack(side=tk.LEFT, fill=tk.Y, padx=(0, utils.PAD_S))
        
        ttk.Button(lf_files, text="🔄 Làm mới", command=self.load_list).pack(side=tk.LEFT, padx=utils.PAD_XS)
        ttk.Button(lf_files, text="➕ Thêm", command=self.add_file).pack(side=tk.LEFT, padx=utils.PAD_XS)
        ttk.Button(lf_files, text="✏️ Đổi tên", command=self.rename_file).pack(side=tk.LEFT, padx=utils.PAD_XS)
        ttk.Button(lf_files, text="🔍 Kiểm tra", command=self.analyze_file).pack(side=tk.LEFT, padx=utils.PAD_XS)
        ttk.Button(lf_files, text="👁️ Xem", command=self.preview_file).pack(side=tk.LEFT, padx=utils.PAD_XS)
        ttk.Button(lf_files, text="❌ Xóa", command=self.delete_file).pack(side=tk.LEFT, padx=utils.PAD_XS)
        ttk.Button(lf_files, text="📥 Xuất mẫu", command=self.export_epu_format).pack(side=tk.LEFT, padx=utils.PAD_XS)

        # Nhóm CSDL
        lf_db = ttk.LabelFrame(frm_controls, text="Cơ sở dữ liệu", padding=utils.PAD_S)
        lf_db.pack(side=tk.LEFT, fill=tk.Y, padx=utils.PAD_S)
        
        ttk.Button(lf_db, text="💾 Import", command=self.import_to_db, bootstyle="info").pack(side=tk.LEFT, padx=utils.PAD_XS)
        ttk.Button(lf_db, text="📂 Quản lý", command=self.search_db).pack(side=tk.LEFT, padx=utils.PAD_XS)
        ttk.Button(lf_db, text="📊 Thống kê", command=self.show_db_stats).pack(side=tk.LEFT, padx=utils.PAD_XS)

        # Nhóm Hệ thống (bên phải)
        lf_sys = ttk.Frame(frm_controls)
        lf_sys.pack(side=tk.RIGHT, fill=tk.Y)
        ttk.Button(lf_sys, text="📝 Mở Word", command=self.open_selected_file).pack(side=tk.TOP, fill=tk.X, pady=1)
        ttk.Button(lf_sys, text="📂 Thư mục", command=self.open_folder).pack(side=tk.TOP, fill=tk.X, pady=1)
        
        # --- Treeview Danh sách File ---
        frm_list = ttk.Frame(self.parent)
        frm_list.grid(row=1, column=0, sticky="nsew", padx=utils.PAD_M, pady=utils.PAD_XS)
        frm_list.rowconfigure(0, weight=1)
        frm_list.columnconfigure(0, weight=1)
        
        cols = ("filename", "size", "modified", "db_status")
        self.tree = ttk.Treeview(frm_list, columns=cols, show="headings", bootstyle="primary", selectmode="extended")
        self.tree.heading("filename", text="Tên file ngân hàng (.docx)")
        self.tree.heading("size", text="Size")
        self.tree.heading("modified", text="Lần sửa cuối")
        self.tree.heading("db_status", text="CSDL")
        
        self.tree.column("filename", width=400, anchor=tk.W)
        self.tree.column("size", width=70, anchor=tk.E)
        self.tree.column("modified", width=130, anchor=tk.CENTER)
        self.tree.column("db_status", width=140, anchor=tk.W)
        
        self.tree.grid(row=0, column=0, sticky="nsew")
        
        vsb = ttk.Scrollbar(frm_list, orient="vertical", command=self.tree.yview)
        vsb.grid(row=0, column=1, sticky="ns")
        self.tree.configure(yscrollcommand=vsb.set)
        
        # Double click để mở file hệ thống
        self.tree.bind("<Double-1>", lambda e: self.open_selected_file())
        
        # Kéo thả file
        if HAS_DND:
            self.tree.drop_target_register(DND_FILES)
            self.tree.dnd_bind('<<Drop>>', self.handle_drop)

        # Khung log
        frm_log = ttk.LabelFrame(self.parent, text="Thông báo hệ thống", padding=utils.PAD_S)
        frm_log.grid(row=2, column=0, sticky="ew", padx=utils.PAD_M, pady=(utils.PAD_XS, utils.PAD_S))
        self.lbl_log = ttk.Label(frm_log, text="Sẵn sàng. Kéo thả file .docx vào danh sách để thêm nhanh.", font=utils.FONT_MAIN)
        self.lbl_log.pack(side=tk.LEFT, padx=utils.PAD_S, pady=utils.PAD_XS, fill=tk.X)
        
        # Lần nạp đầu
        self.load_list()

    def get_selected_paths(self):
        selected = self.tree.selection()
        if not selected:
            return []
        paths = []
        for sel in selected:
            item = self.tree.item(sel)
            fname = item['values'][0]
            paths.append(os.path.join(BANK_DIR, fname))
        return paths

    def get_selected_path(self):
        paths = self.get_selected_paths()
        return paths[0] if paths else None

    def load_list(self):
        for item in self.tree.get_children():
            self.tree.delete(item)
        if not os.path.exists(BANK_DIR):
            return
        files = [f for f in os.listdir(BANK_DIR) if f.lower().endswith('.docx') and not f.startswith('~')]
        for f in files:
            path = os.path.join(BANK_DIR, f)
            size_kb = os.path.getsize(path) / 1024
            mtime = os.path.getmtime(path)
            dt_str = datetime.fromtimestamp(mtime).strftime("%d/%m/%Y %H:%M")
            
            # Kiểm tra trạng thái DB
            db_status = "⬜ Chưa import"
            try:
                sync = db.check_sync_status(path)
                if sync["status"] == "up_to_date":
                    db_status = "✅ Đã đồng bộ"
                elif sync["status"] == "modified":
                    db_status = "⚠️ File đã sửa"
            except Exception:
                pass
            
            self.tree.insert("", "end", values=(f, f"{size_kb:.1f}", dt_str, db_status))

    def add_file(self):
        fnames = filedialog.askopenfilenames(filetypes=[("Word files", "*.docx")])
        if not fnames: return
        added = 0
        for fname in fnames:
            base = os.path.basename(fname)
            dest = os.path.join(BANK_DIR, base)
            if not os.path.exists(dest):
                shutil.copy2(fname, dest)
                added += 1
            else:
                if messagebox.askyesno("Trùng lặp", f"File '{base}' đã tồn tại. Ghi đè?"):
                    shutil.copy2(fname, dest)
                    added += 1
        if added > 0:
            self.load_list()
            messagebox.showinfo("Thành công", f"Đã thêm/cập nhật {added} file vào ngân hàng.")

    def handle_drop(self, event):
        """Xử lý kéo thả file từ Windows Explorer."""
        data = event.data
        # Windows thường bọc đường dẫn có khoảng trắng trong dấu {}
        # Ví dụ: {C:\My Files\test.docx}
        import re
        files = re.findall(r'{(.*?)}', data)
        if not files:
            files = data.split()
            
        added = 0
        for f in files:
            f = f.strip("{}").strip()
            if f.lower().endswith(".docx"):
                base = os.path.basename(f)
                dest = os.path.join(BANK_DIR, base)
                try:
                    shutil.copy2(f, dest)
                    added += 1
                except Exception:
                    pass
        if added > 0:
            self.load_list()
            self.lbl_log.config(text=f"✅ Đã thêm {added} file qua kéo thả.")

    def delete_file(self):
        paths = self.get_selected_paths()
        if not paths:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn ít nhất 1 file để xóa!")
            return
        
        n = len(paths)
        if n == 1:
            msg = f"Bạn có chắc chắn muốn xóa file:\n{os.path.basename(paths[0])}?"
        else:
            msg = f"Bạn có chắc chắn muốn xóa {n} file đang chọn?"

        if messagebox.askyesno("Xác nhận", msg):
            success = 0
            for path in paths:
                try:
                    os.remove(path)
                    success += 1
                except Exception as e:
                    messagebox.showerror("Lỗi", f"Không thể xóa {os.path.basename(path)}: {e}")
            
            self.load_list()
            if success > 0:
                self.lbl_log.config(text=f"✅ Đã xóa {success} file khỏi ngân hàng.")

    def rename_file(self):
        from tkinter import simpledialog
        path = self.get_selected_path()
        if not path:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn 1 file để đổi tên!")
            return
        old_name = os.path.basename(path)
        new_name = simpledialog.askstring("Đổi tên", "Nhập tên mới (bao gồm cả .docx):", initialvalue=old_name)
        if new_name and new_name.strip() != old_name:
            new_name = new_name.strip()
            if not new_name.lower().endswith('.docx'):
                new_name += '.docx'
            new_path = os.path.join(BANK_DIR, new_name)
            if os.path.exists(new_path):
                messagebox.showerror("Lỗi", "Tên file đã tồn tại!")
                return
            try:
                os.rename(path, new_path)
                self.load_list()
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể đổi tên file (có thể file đang được mở): {e}")

    def open_folder(self):
        abs_path = os.path.abspath(BANK_DIR)
        try:
            os.startfile(abs_path)
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))

    def open_selected_file(self):
        path = self.get_selected_path()
        if not path: return
        try:
            os.startfile(os.path.abspath(path))
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))

    def analyze_file(self):
        path = self.get_selected_path()
        if not path:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn 1 file để phân tích!")
            return
        self.lbl_log.config(text="Đang phân tích định dạng, vui lòng đợi...")
        self.parent.update_idletasks()
        
        if not split_questions_from_docx:
            self.lbl_log.config(text="Lỗi: Không thể tìm thấy hàm phân tích (split_questions_from_docx).")
            return
        
        try:
            qs, _, warns = split_questions_from_docx(path)
            msg = f"✅ FILE HỢP LỆ. Tổng số câu hỏi đọc được: {len(qs)} câu.\n"
            
            if warns:
                msg += f"⚠️ CẢNH BÁO: {len(warns)} câu bị lỗi nhận diện (đã bỏ qua).\n"
                msg += "\n".join(warns[:10]) + ("\n..." if len(warns) > 10 else "") + "\n"

            # Thống kê phân loại độ khó nếu có
            diff_counts = {}
            for q in qs:
                diff = q.get('diff_code') or "Không xác định"
                diff_counts[diff] = diff_counts.get(diff, 0) + 1
            
            if diff_counts:
                msg += "Phân loại độ khó: " + ", ".join([f"[{k}]: {v} câu" for k,v in diff_counts.items()])
            
            self.lbl_log.config(text=msg)
            res_msg = f"Phân tích thành công!\nSố câu đọc được: {len(qs)}"
            if warns:
                res_msg += f"\nSố câu lỗi (bỏ qua): {len(warns)}"
            messagebox.showinfo("Phân tích hoàn tất", res_msg)
        except Exception as e:
            msg = f"❌ FILE LỖI: {e}\nHãy kiểm tra lại file Word (đánh dấu đáp án, từ khóa 'Câu X.', v.v...)"
            self.lbl_log.config(text=msg)
            messagebox.showerror("Phân tích lỗi", str(e))

    def preview_file(self):
        path = self.get_selected_path()
        if not path:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn 1 file để xem trước!")
            return
        
        try:
            qs, _, warns = split_questions_from_docx(path)
            if warns:
                self.lbl_log.config(text=f"⚠️ Cảnh báo xem trước: Có {len(warns)} câu hỏi không nhận diện được đáp án.")
            
            # Sử dụng Dialog chuẩn từ utils
            utils.QuestionPreviewDialog(self.parent, qs, title=f"Xem trước: {os.path.basename(path)}")
            




        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể xem trước file: {e}")

    def export_epu_format(self):
        path = self.get_selected_path()
        if not path:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn 1 file để xuất!")
            return
        if not split_questions_from_docx:
            messagebox.showerror("Lỗi", "Không thể tìm thấy hàm phân tích (split_questions_from_docx).")
            return
            
        save_path = filedialog.asksaveasfilename(
            title="Lưu file xuất",
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")],
            initialfile=os.path.basename(path).replace(".docx", "_EPU.docx")
        )
        if not save_path:
            return
            
        self.lbl_log.config(text="Đang xuất file theo mẫu EPU, vui lòng đợi...")
        self.parent.update_idletasks()
        
        try:
            qs, _, _ = split_questions_from_docx(path)
            
            # Chuyển sang dùng utils cho đồng nhất

            doc = Document()
            utils.set_page_layout(doc)
            utils.set_single_line_spacing(doc)
            for i, q in enumerate(qs, 1):
                diff = q.get('diff_code') or "TB"
                p_stem = doc.add_paragraph()
                r_label = p_stem.add_run(f"Câu {i} [<{diff}>]: ")
                r_label.bold = True
                
                # Render nội dung chính
                if q.get("stem_media_spans"):
                    utils.render_spans_into_paragraph(p_stem, q["stem_media_spans"], default_black=False)
                else:
                    stem_text = (q.get("stem_text") or "").strip()
                    if stem_text.startswith(":"): stem_text = stem_text[1:].strip()
                    if stem_text: p_stem.add_run(stem_text)
                            
                for extra_spans in q.get("stem_extra_media_spans", []):
                    if not extra_spans: continue
                    p_ex = doc.add_paragraph()
                    utils.render_spans_into_paragraph(p_ex, extra_spans, default_black=False)
                            
                opts = q["options"]
                for j, opt in enumerate(opts):
                    p_opt = doc.add_paragraph()
                    r_opt_lbl = p_opt.add_run("[<$>] ")
                    if j == q["correct_index"]:
                        r_opt_lbl.font.color.rgb = RGBColor(0xFF, 0, 0)
                    utils.render_spans_into_paragraph(p_opt, opt["info"]["spans"], default_black=False)
                            
            doc.save(save_path)
            self.lbl_log.config(text=f"✅ Xuất thành công: {os.path.basename(save_path)}")
            messagebox.showinfo("Thành công", f"Đã xuất file thành công tại:\n{save_path}")
            
        except Exception as e:
            self.lbl_log.config(text=f"❌ Lỗi khi xuất: {e}")
            messagebox.showerror("Lỗi", str(e))

    # ============= CHỨC NĂNG CƠ SỞ DỮ LIỆU =============
    
    def import_to_db(self):
        """Import file Word đang chọn vào cơ sở dữ liệu."""
        path = self.get_selected_path()
        if not path:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn 1 file để import vào DB!")
            return
        if not split_questions_from_docx:
            messagebox.showerror("Lỗi", "Không thể tìm thấy hàm phân tích.")
            return
        
        # Sử dụng hộp thoại lựa chọn mới
        dlg = ImportChoiceDialog(self.parent, os.path.basename(path))
        self.parent.wait_window(dlg.win)
        
        if not dlg.result:
            return
            
        mode = dlg.result["mode"]
        val = dlg.result["value"]

        self.lbl_log.config(text="Đang import vào cơ sở dữ liệu...")
        self.parent.update_idletasks()
        
        try:
            # Kiểm tra file đã thay đổi chưa
            sync = db.check_sync_status(path)
            
            if mode == "new":
                subject_name = val
                if sync["status"] == "up_to_date":
                    if not messagebox.askyesno("Đã tồn tại", 
                        f"File này đã được import vào DB (nội dung chưa thay đổi).\n"
                        f"Bạn có muốn import lại không?"):
                        self.lbl_log.config(text=sync["message"])
                        return
                    qs, _, _ = split_questions_from_docx(path)
                    result = db.reimport_bank(path, qs, subject_name)
                elif sync["status"] == "modified":
                    if messagebox.askyesno("File đã thay đổi", 
                        f"File này đã thay đổi từ lần import cuối.\n"
                        f"Bạn có muốn import lại (cập nhật) không?"):
                        qs, _, _ = split_questions_from_docx(path)
                        result = db.reimport_bank(path, qs, subject_name)
                    else:
                        return
                else:
                    qs, _, _ = split_questions_from_docx(path)
                    result = db.import_bank_to_db(path, qs, subject_name)
            else:
                # Chế độ Append (Gộp vào bank có sẵn)
                bank_id = val
                qs, _, _ = split_questions_from_docx(path)
                result = db.append_questions_to_bank(bank_id, qs)

            # Hiển thị kết quả
            msg = result.get("message", "Hoàn tất")
            
            # Cảnh báo câu trùng lặp
            if result.get("duplicates_found"):
                msg += "\n\n⚠️ Câu hỏi trùng lặp (nội dung câu đã có trong DB):\n"
                for dup in result["duplicates_found"][:5]:
                    msg += f"  • Câu {dup['qid']}: '{dup['stem_text'][:50]}...' (trùng với: {dup['existing_file']})\n"
                if len(result["duplicates_found"]) > 5:
                    msg += f"  ... và {len(result['duplicates_found']) - 5} câu khác.\n"
            
            # Cảnh báo đáp án trùng lặp (trong chính câu đó)
            if result.get("duplicate_options_found"):
                msg += "\n\n❌ Phát hiện câu hỏi có các đáp án trùng nhau:\n"
                for dup in result["duplicate_options_found"][:5]:
                    msg += f"  • Câu {dup['qid']}: '{dup['stem_text'][:50]}...'\n"
                if len(result["duplicate_options_found"]) > 5:
                    msg += f"  ... và {len(result['duplicate_options_found']) - 5} câu khác.\n"

            self.lbl_log.config(text=msg)
            messagebox.showwarning("Kết quả Import" if (result.get("duplicates_found") or result.get("duplicate_options_found")) else "Kết quả Import", msg)
            self.load_list()  # Cập nhật trạng thái DB trên treeview
            
        except Exception as e:
            self.lbl_log.config(text=f"❌ Lỗi import: {e}")
            messagebox.showerror("Lỗi", str(e))

    def import_all_to_db(self):
        """Import tất cả file Word trong ngân hàng vào DB."""
        if not split_questions_from_docx:
            messagebox.showerror("Lỗi", "Không thể tìm thấy hàm phân tích.")
            return
        
        if not os.path.exists(BANK_DIR):
            return
        
        files = [f for f in os.listdir(BANK_DIR) if f.lower().endswith('.docx') and not f.startswith('~')]
        if not files:
            messagebox.showinfo("Thông báo", "Không có file nào trong ngân hàng.")
            return
        
        subject_name = simpledialog.askstring(
            "Môn học", f"Nhập tên môn học cho {len(files)} file (có thể để trống):",
            initialvalue=""
        )
        if subject_name is None:
            return
        
        if not messagebox.askyesno("Xác nhận", f"Import {len(files)} file vào cơ sở dữ liệu?"):
            return
        
        total_imported = 0
        total_duplicates = 0
        errors = []
        
        for i, f in enumerate(files):
            path = os.path.join(BANK_DIR, f)
            self.lbl_log.config(text=f"Đang import {i+1}/{len(files)}: {f}...")
            self.parent.update_idletasks()
            
            try:
                qs, _, _ = split_questions_from_docx(path)
                sync = db.check_sync_status(path)
                
                if sync["status"] == "up_to_date":
                    continue  # Bỏ qua file đã import và chưa thay đổi
                elif sync["status"] == "modified":
                    result = db.reimport_bank(path, qs, subject_name)
                else:
                    result = db.import_bank_to_db(path, qs, subject_name)
                
                total_imported += result.get("total_imported", 0)
                total_duplicates += len(result.get("duplicates_found", []))
            except Exception as e:
                errors.append(f"{f}: {e}")
        
        msg = f"✅ Hoàn tất import: {total_imported} câu hỏi mới."
        if total_duplicates:
            msg += f"\nPhát hiện {total_duplicates} câu trùng lặp (đã bỏ qua)."
        if errors:
            msg += f"\n\n⚠️ Lỗi ({len(errors)} file):\n" + "\n".join(errors[:5])
        
        self.lbl_log.config(text=msg)
        if errors:
            messagebox.showwarning("Kết quả Import (Có lỗi)", msg)
        else:
            messagebox.showinfo("Kết quả Import", msg)
        self.load_list()

    def search_db(self):
        """Mở cửa sổ Quản lý CSDL chuyên sâu."""
        DatabaseManagerDialog(self.parent)

    def show_db_stats(self):
        """Hiển thị thống kê tổng quan ngân hàng câu hỏi trong DB."""
        try:
            stats = db.get_bank_stats()
            banks = db.get_all_banks()
            
            msg = "📊 THỐNG KÊ NGÂN HÀNG CÂU HỎI\n"
            msg += "=" * 40 + "\n"
            msg += f"Tổng số file đã import: {stats['total_banks']}\n"
            msg += f"Tổng số môn học: {stats['total_subjects']}\n"
            msg += f"Tổng số câu hỏi: {stats['total_questions']}\n\n"
            
            if stats.get("by_difficulty"):
                msg += "Phân bố theo độ khó:\n"
                for diff, count in stats["by_difficulty"].items():
                    msg += f"  [{diff}]: {count} câu\n"
            
            if banks:
                msg += "\nDanh sách ngân hàng:\n"
                for b in banks[:20]:
                    msg += f"  • {b['file_name']} ({b['total_questions']} câu"
                    if b.get('subject_name'):
                        msg += f", Môn: {b['subject_name']}"
                    msg += f", Import: {b['imported_at']})\n"
            
            # Hiển thị trong dialog
            dlg = tk.Toplevel(self.parent)
            utils.setup_dialog(dlg, width_pct=0.6, height_pct=0.6, title="Thống kê Ngân hàng câu hỏi", parent=self.parent)
            
            txt = tk.Text(dlg, wrap="word", font=("Consolas", 11))
            txt.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            txt.insert(tk.END, msg)
            txt.config(state="disabled")
            
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))

class ImportChoiceDialog:
    def __init__(self, parent, file_name):
        self.win = tk.Toplevel(parent)
        utils.setup_dialog(self.win, width_pct=0.4, height_pct=0.4, title="Lựa chọn Import", parent=parent)
        
        self.result = None
        self.banks = db.get_all_banks()
        
        frm = ttk.Frame(self.win, padding=utils.PAD_L)
        frm.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frm, text=f"File: {file_name}", font=utils.FONT_BOLD).pack(anchor="w", pady=(0, 10))
        
        self.var_mode = tk.StringVar(value="new")
        
        # --- Option 1: New Bank ---
        rb_new = ttk.Radiobutton(frm, text="Tạo ngân hàng mới hoàn toàn", variable=self.var_mode, value="new", command=self._toggle)
        rb_new.pack(anchor="w", pady=5)
        
        self.frm_new = ttk.Frame(frm)
        self.frm_new.pack(fill=tk.X, padx=25, pady=(0, 15))
        ttk.Label(self.frm_new, text="Tên môn học (tùy chọn):").pack(side=tk.LEFT)
        self.ent_subject = ttk.Entry(self.frm_new)
        self.ent_subject.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # --- Option 2: Append to existing ---
        rb_append = ttk.Radiobutton(frm, text="Thêm câu hỏi vào ngân hàng có sẵn", variable=self.var_mode, value="append", command=self._toggle)
        rb_append.pack(anchor="w", pady=5)
        
        self.frm_append = ttk.Frame(frm)
        self.frm_append.pack(fill=tk.X, padx=25, pady=(0, 15))
        
        bank_labels = [f"{b['file_name']} ({b['total_questions']} câu - {b.get('subject_name', '')})" for b in self.banks]
        self.cb_banks = ttk.Combobox(self.frm_append, values=bank_labels, state="readonly")
        self.cb_banks.pack(fill=tk.X, expand=True)
        if self.banks:
            self.cb_banks.current(0)
        else:
            rb_append.configure(state="disabled")
            
        # --- Footer ---
        frm_btn = ttk.Frame(frm)
        frm_btn.pack(side=tk.BOTTOM, fill=tk.X, pady=10)
        
        ttk.Button(frm_btn, text="✅ Xác nhận", command=self._confirm, bootstyle="success").pack(side=tk.RIGHT, padx=5)
        ttk.Button(frm_btn, text="❌ Hủy", command=self.win.destroy).pack(side=tk.RIGHT)
        
        self._toggle()

    def _toggle(self):
        m = self.var_mode.get()
        if m == "new":
            self.ent_subject.configure(state="normal")
            self.cb_banks.configure(state="disabled")
        else:
            self.ent_subject.configure(state="disabled")
            self.cb_banks.configure(state="normal")

    def _confirm(self):
        m = self.var_mode.get()
        if m == "append":
            idx = self.cb_banks.current()
            if idx < 0:
                messagebox.showwarning("Cảnh báo", "Vui lòng chọn ngân hàng đích!")
                return
            self.result = {"mode": "append", "value": self.banks[idx]["id"]}
        else:
            self.result = {"mode": "new", "value": self.ent_subject.get().strip()}
        self.win.destroy()

    def check_duplicates(self):
        """Kiểm tra câu hỏi trùng lặp trong toàn bộ DB."""
        try:
            dups = db.find_duplicates()
            if not dups:
                messagebox.showinfo("Kết quả", "Không phát hiện câu hỏi trùng lặp trong CSDL.")
                return
            
            dlg = tk.Toplevel(self.parent)
            dlg.title(f"Câu hỏi trùng lặp ({len(dups)} cặp)")
            utils.set_window_icon(dlg)
            dlg.geometry("900x500")
            dlg.transient(self.parent.winfo_toplevel())
            
            cols = ("id1", "file1", "text1", "id2", "file2")
            tree = ttk.Treeview(dlg, columns=cols, show="headings")
            tree.heading("id1", text="ID câu 1")
            tree.heading("file1", text="File 1")
            tree.heading("text1", text="Nội dung")
            tree.heading("id2", text="ID câu 2")
            tree.heading("file2", text="File 2")
            
            tree.column("id1", width=60, anchor=tk.CENTER)
            tree.column("file1", width=180, anchor=tk.W)
            tree.column("text1", width=350, anchor=tk.W)
            tree.column("id2", width=60, anchor=tk.CENTER)
            tree.column("file2", width=180, anchor=tk.W)
            
            vsb = ttk.Scrollbar(dlg, orient="vertical", command=tree.yview)
            tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
            vsb.pack(side=tk.RIGHT, fill=tk.Y, pady=10)
            tree.config(yscrollcommand=vsb.set)
            
            for d in dups:
                preview = (d["text1"][:60] + "...") if len(d.get("text1", "")) > 60 else d.get("text1", "")
                tree.insert("", "end", values=(
                    d["id1"], d["file1"], preview, d["id2"], d["file2"]
                ))
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))


class DatabaseManagerDialog:
    def __init__(self, parent):
        self.parent = parent
        self.dlg = tk.Toplevel(parent)
        utils.setup_dialog(self.dlg, width_pct=0.85, height_pct=0.85, title="Quản lý Ngân hàng trong CSDL", parent=parent)

        self.current_bank_id = None
        self.preview_images = []
        self.search_lock = False # Tránh đè dữ liệu khi đang load

        # --- Layout chính: 2 cột (Sidebar, Content) ---
        self.dlg.columnconfigure(1, weight=10)
        self.dlg.rowconfigure(0, weight=1)

        # 1. Sidebar: Danh sách Ngân hàng
        frm_left = ttk.Frame(self.dlg, padding=utils.PAD_S)
        frm_left.grid(row=0, column=0, sticky="nsew")
        frm_left.rowconfigure(1, weight=1)

        ttk.Label(frm_left, text="📂 Ngân hàng", font=utils.FONT_BOLD).pack(pady=utils.PAD_S)
        
        # Thêm nút Quản lý Môn học ở Sidebar
        frm_sub_btns = ttk.Frame(frm_left)
        frm_sub_btns.pack(fill=tk.X, pady=2)
        ttk.Button(frm_sub_btns, text="✏️ Đổi tên Môn", command=self.rename_current_subject, bootstyle="link").pack(side=tk.LEFT, expand=True)
        ttk.Button(frm_sub_btns, text="🗑️ Xóa Môn", command=self.delete_current_subject, bootstyle="link").pack(side=tk.LEFT, expand=True)

        self.bank_tree = ttk.Treeview(frm_left, columns=("id", "questions"), show="tree headings", selectmode="browse")
        self.bank_tree.heading("#0", text="Môn học / File")
        self.bank_tree.heading("id", text="ID")
        self.bank_tree.heading("questions", text="Câu")
        self.bank_tree.column("#0", width=220)
        self.bank_tree.column("id", width=35, anchor=tk.CENTER)
        self.bank_tree.column("questions", width=45, anchor=tk.CENTER)
        
        self.bank_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb_b = ttk.Scrollbar(frm_left, orient="vertical", command=self.bank_tree.yview)
        vsb_b.pack(side=tk.RIGHT, fill=tk.Y)
        self.bank_tree.configure(yscrollcommand=vsb_b.set)
        self.bank_tree.bind("<<TreeviewSelect>>", self.on_bank_select)

        # 2. Content: Danh sách Câu hỏi & Preview
        frm_right = ttk.Frame(self.dlg, padding=5)
        frm_right.grid(row=0, column=1, sticky="nsew")
        frm_right.rowconfigure(2, weight=1) # Treeview
        frm_right.rowconfigure(3, weight=1) # Preview
        frm_right.columnconfigure(0, weight=1)

        # Toolbar 1: Thao tác chung
        frm_toolbar = ttk.Frame(frm_right)
        frm_toolbar.grid(row=0, column=0, sticky="ew", pady=2)
        
        self.btn_del_bank = ttk.Button(frm_toolbar, text="❌ Xóa Ngân hàng", command=self.delete_bank, state="disabled")
        self.btn_del_bank.pack(side=tk.LEFT, padx=2)
        
        ttk.Button(frm_toolbar, text="🔗 Gộp Ngân hàng", command=self.merge_banks_ui).pack(side=tk.LEFT, padx=2)
        ttk.Button(frm_toolbar, text="📄 Xuất Word", command=self.export_bank_db_to_word).pack(side=tk.LEFT, padx=2)

        ttk.Separator(frm_toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=5)
        
        ttk.Button(frm_toolbar, text="🪄 Gộp câu trùng", command=self.auto_merge_duplicates, bootstyle="warning-outline").pack(side=tk.LEFT, padx=2)
        ttk.Separator(frm_toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=5)

        ttk.Button(frm_toolbar, text="➕ Soạn mới", command=self.add_new_question).pack(side=tk.LEFT, padx=2)
        self.btn_edit_q = ttk.Button(frm_toolbar, text="✏️ Sửa", command=self.edit_question, state="disabled")
        self.btn_edit_q.pack(side=tk.LEFT, padx=2)
        self.btn_move_q = ttk.Button(frm_toolbar, text="📦 Di chuyển", command=self.move_questions_ui, state="disabled")
        self.btn_move_q.pack(side=tk.LEFT, padx=2)
        self.btn_del_q = ttk.Button(frm_toolbar, text="🗑️ Xóa Câu", command=self.delete_question, state="disabled")
        self.btn_del_q.pack(side=tk.LEFT, padx=2)
        
        ttk.Button(frm_toolbar, text="🔄 Refresh", command=self.refresh_all).pack(side=tk.RIGHT, padx=2)
        
        # Toolbar 2: Tìm kiếm & Tiện ích nâng cao
        frm_search = ttk.Frame(frm_right)
        frm_search.grid(row=1, column=0, sticky="ew", pady=5)

        ttk.Label(frm_search, text="🔍 Tìm kiếm:").pack(side=tk.LEFT, padx=5)
        self.ent_search = ttk.Entry(frm_search, width=40)
        self.ent_search.pack(side=tk.LEFT, padx=5)
        self.ent_search.bind("<Return>", lambda e: self.global_search())
        
        ttk.Button(frm_search, text="Tìm kiếm", command=self.global_search, bootstyle="info").pack(side=tk.LEFT, padx=2)
        
        ttk.Separator(frm_search, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)
        
        ttk.Button(frm_search, text="🪄 Kiểm tra CSDL", command=self.open_audit_center, bootstyle="danger").pack(side=tk.LEFT, padx=2)
        ttk.Button(frm_search, text="🛡️ Sao lưu DB", command=self.backup_db_ui, bootstyle="success-outline").pack(side=tk.LEFT, padx=2)

        # Danh sách câu hỏi
        frm_q_list = ttk.LabelFrame(frm_right, text="Danh sách câu hỏi (Giữ Ctrl/Shift để chọn nhiều)")
        frm_q_list.grid(row=2, column=0, sticky="nsew", pady=5)
        frm_q_list.rowconfigure(0, weight=1)
        frm_q_list.columnconfigure(0, weight=1)

        cols = ("id", "qid", "stem", "diff", "correct", "bank")
        self.q_tree = ttk.Treeview(frm_q_list, columns=cols, show="headings", selectmode="extended")
        self.q_tree.heading("id", text="ID")
        self.q_tree.heading("qid", text="Câu")
        self.q_tree.heading("stem", text="Nội dung")
        self.q_tree.heading("diff", text="Độ khó")
        self.q_tree.heading("correct", text="Đ/A")
        self.q_tree.heading("bank", text="Ngân hàng")
        
        self.q_tree.column("id", width=40, anchor=tk.CENTER)
        self.q_tree.column("qid", width=40, anchor=tk.CENTER)
        self.q_tree.column("stem", width=450, anchor=tk.W)
        self.q_tree.column("diff", width=50, anchor=tk.CENTER)
        self.q_tree.column("correct", width=35, anchor=tk.CENTER)
        self.q_tree.column("bank", width=150, anchor=tk.W)
        
        self.q_tree.grid(row=0, column=0, sticky="nsew")
        vsb_q = ttk.Scrollbar(frm_q_list, orient="vertical", command=self.q_tree.yview)
        vsb_q.grid(row=0, column=1, sticky="ns")
        self.q_tree.configure(yscrollcommand=vsb_q.set)
        self.q_tree.bind("<<TreeviewSelect>>", self.on_question_select)

        # Preview Pane
        frm_preview = ttk.LabelFrame(frm_right, text="Xem trước nội dung")
        frm_preview.grid(row=3, column=0, sticky="nsew", pady=5)
        
        self.txt_preview = tk.Text(frm_preview, wrap="word", font=("Arial", 11), state="disabled")
        self.txt_preview.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        vsb_p = ttk.Scrollbar(frm_preview, orient="vertical", command=self.txt_preview.yview)
        vsb_p.pack(side=tk.RIGHT, fill=tk.Y)
        self.txt_preview.configure(yscrollcommand=vsb_p.set)

        # Tags cho Preview
        self.txt_preview.tag_config("bold_header", font=("Arial", 12, "bold"), foreground="#2C3E50")
        self.txt_preview.tag_config("bold", font=("Arial", 11, "bold"))
        self.txt_preview.tag_config("italic", font=("Arial", 11, "italic"))
        self.txt_preview.tag_config("bold_italic", font=("Arial", 11, "bold", "italic"))
        self.txt_preview.tag_config("correct", foreground="red", font=("Arial", 11, "bold"))

        self.refresh_all()

    def refresh_all(self):
        """Tải lại danh sách Ngân hàng và Môn học."""
        self.bank_tree.delete(*self.bank_tree.get_children())
        self.q_tree.delete(*self.q_tree.get_children())
        self.txt_preview.config(state="normal")
        self.txt_preview.delete("1.0", tk.END)
        self.txt_preview.config(state="disabled")
        self.current_bank_id = None
        self.btn_del_bank.config(state="disabled")
        self.btn_edit_q.config(state="disabled")
        self.btn_move_q.config(state="disabled")
        self.btn_del_q.config(state="disabled")

        try:
            banks = db.get_all_banks()
            all_subjects = db.get_subjects()
            
            # Khởi tạo từ điển với tất cả môn học hiện có trong DB
            subjects_map = {s["name"]: [] for s in all_subjects}
            
            # Đảm bảo "Chưa phân loại" xuất hiện nếu có ngân hàng không thuộc môn nào
            if any(not b.get("subject_name") for b in banks):
                if "Chưa phân loại" not in subjects_map:
                    subjects_map["Chưa phân loại"] = []

            # Gom nhóm ngân hàng vào môn học
            for b in banks:
                sname = b.get("subject_name") or "Chưa phân loại"
                if sname not in subjects_map:
                    subjects_map[sname] = []
                subjects_map[sname].append(b)
            
            # Hiển thị lên cây
            for sname in sorted(subjects_map.keys()):
                sid = self.bank_tree.insert("", "end", text=sname, open=True)
                for b in subjects_map[sname]:
                    self.bank_tree.insert(sid, "end", text=b["file_name"], values=(b["id"], b["total_questions"]))
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể nạp dữ liệu: {e}")

    def on_bank_select(self, event):
        sel = self.bank_tree.selection()
        if not sel: return
        item = self.bank_tree.item(sel[0])
        if not item["values"]: # Đây là node Môn học
            self.current_bank_id = None
            self.btn_del_bank.config(state="disabled")
            self.q_tree.delete(*self.q_tree.get_children())
            return
        
        self.current_bank_id = item["values"][0]
        self.btn_del_bank.config(state="normal")
        self.load_questions(self.current_bank_id)

    def load_questions(self, bank_id):
        self.q_tree.delete(*self.q_tree.get_children())
        try:
            questions = db.search_questions(bank_id=bank_id, limit=1000)
            for q in questions:
                correct_label = ""
                for opt in q.get("options", []):
                    if opt.get("is_correct"):
                        correct_label = opt.get("label", "")
                        break
                stem_preview = q.get("stem_text", "").strip()[:80]
                self.q_tree.insert("", "end", values=(
                    q["id"], q.get("qid_in_file", ""), stem_preview, 
                    q.get("diff_code", ""), correct_label, q.get("file_name", "")
                ))
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể nạp câu hỏi: {e}")

    def global_search(self):
        """Tìm kiếm câu hỏi trên toàn bộ CSDL."""
        kw = self.ent_search.get().strip()
        if not kw:
            if self.current_bank_id: self.load_questions(self.current_bank_id)
            return
            
        self.q_tree.delete(*self.q_tree.get_children())
        self.current_bank_id = None
        self.btn_del_bank.config(state="disabled")
        
        try:
            questions = db.search_questions(keyword=kw, limit=500)
            if not questions:
                messagebox.showinfo("Tìm kiếm", f"Không tìm thấy câu hỏi nào với từ khóa: '{kw}'")
                return
                
            for q in questions:
                correct_label = ""
                for opt in q.get("options", []):
                    if opt.get("is_correct"):
                        correct_label = opt.get("label", "")
                        break
                stem_preview = q.get("stem_text", "").strip()[:80]
                self.q_tree.insert("", "end", values=(
                    q["id"], q.get("qid_in_file", ""), stem_preview, 
                    q.get("diff_code", ""), correct_label, q.get("file_name", "")
                ))
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))

    def on_question_select(self, event):
        sel = self.q_tree.selection()
        if not sel:
            self.btn_edit_q.config(state="disabled")
            self.btn_move_q.config(state="disabled")
            self.btn_del_q.config(state="disabled")
            return
        
        self.btn_edit_q.config(state="normal")
        self.btn_move_q.config(state="normal")
        self.btn_del_q.config(state="normal")
        
        if len(sel) == 1:
            q_id = self.q_tree.item(sel[0])["values"][0]
            self.preview_question(q_id)
        else:
            self.txt_preview.config(state="normal")
            self.txt_preview.delete("1.0", tk.END)
            self.txt_preview.insert(tk.END, f"\n\n   --- Đang chọn {len(sel)} câu hỏi ---", "bold_header")
            self.txt_preview.config(state="disabled")

    def preview_question(self, q_id):
        try:
            q = db.get_question_with_spans(q_id)
            if not q: return

            self.txt_preview.config(state="normal")
            self.txt_preview.delete("1.0", tk.END)
            self.preview_images.clear()

            self.txt_preview.insert(tk.END, f"Câu {q.get('qid_in_file', '')} [ID: {q_id}] [Độ khó: {q.get('diff_code') or 'N/A'}]\n\n", "bold_header")

            # Render Stem
            if q.get("stem_media_spans"):
                self.render_spans_to_text(q["stem_media_spans"])
            else:
                self.txt_preview.insert(tk.END, q.get("stem_text", ""))
            
            for extra in q.get("stem_extra_media_spans", []):
                self.txt_preview.insert(tk.END, "\n")
                self.render_spans_to_text(extra)
            
            self.txt_preview.insert(tk.END, "\n\n")

            # Render Options
            letters = ["A", "B", "C", "D"]
            for i, opt in enumerate(q.get("options", [])):
                is_correct = bool(opt.get("is_correct"))
                start_ptr = self.txt_preview.index("end-1c")
                self.txt_preview.insert(tk.END, f"{letters[i]}. ")
                
                spans = opt.get("info", {}).get("spans", [])
                if spans:
                    self.render_spans_to_text(spans)
                else:
                    self.txt_preview.insert(tk.END, opt.get("option_text", ""))
                
                if is_correct:
                    self.txt_preview.insert(tk.END, "  <-- ĐÚNG")
                    self.txt_preview.tag_add("correct", start_ptr, self.txt_preview.index("end-1c"))
                self.txt_preview.insert(tk.END, "\n")

            self.txt_preview.config(state="disabled")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể xem trước câu hỏi: {e}")

    def render_spans_to_text(self, spans_list):
        from PIL import Image, ImageTk
        import io
        for sp in spans_list:
            if sp["type"] == "text":
                tag = ""
                if sp.get("bold") and sp.get("italic"): tag = "bold_italic"
                elif sp.get("bold"): tag = "bold"
                elif sp.get("italic"): tag = "italic"
                self.txt_preview.insert(tk.END, sp["text"], tag)
            elif sp["type"] == "image":
                try:
                    pil_img = Image.open(io.BytesIO(sp["blob"]))
                    w, h = pil_img.size
                    if w > 700:
                        pil_img = pil_img.resize((700, int(h * 700 / w)), Image.LANCZOS)
                    photo = ImageTk.PhotoImage(pil_img)
                    self.preview_images.append(photo)
                    self.txt_preview.insert(tk.END, "\n")
                    self.txt_preview.image_create(tk.END, image=photo)
                    self.txt_preview.insert(tk.END, "\n")
                except: pass
            elif sp["type"] == "omml":
                # Chuyển đổi công thức sang văn bản mô phỏng để xem nhanh
                summary = utils.get_spans_text_summary([sp])
                self.txt_preview.insert(tk.END, f" {summary} ", "bold")

    def delete_question(self):
        sel = self.q_tree.selection()
        if not sel: return
        q_id = self.q_tree.item(sel[0])["values"][0]
        
        if messagebox.askyesno("Xác nhận", f"Bạn có chắc muốn xóa câu hỏi ID {q_id} khỏi CSDL?"):
            try:
                db.delete_question(q_id)
                self.load_questions(self.current_bank_id)
                self.txt_preview.config(state="normal")
                self.txt_preview.delete("1.0", tk.END)
                self.txt_preview.config(state="disabled")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể xóa câu hỏi: {e}")

    def delete_bank(self):
        if not self.current_bank_id: return
        sel = self.bank_tree.selection()
        bank_name = self.bank_tree.item(sel[0])["text"]

        if messagebox.askyesno("Xác nhận", f"Bạn có chắc muốn xóa TOÀN BỘ ngân hàng:\n'{bank_name}'\nvà tất cả câu hỏi liên quan khỏi CSDL?"):
            try:
                db.delete_bank(self.current_bank_id)
                self.refresh_all()
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể xóa ngân hàng: {e}")

    def add_new_question(self):
        """Mở trình soạn thảo để thêm câu hỏi mới."""
        editor = QuestionEditorDialog(self.dlg, on_save_callback=self.refresh_after_edit)
        # Nếu đang chọn ngân hàng, truyền vào luôn
        if self.current_bank_id:
            editor.set_initial_bank(self.current_bank_id)

    def edit_question(self):
        """Mở trình soạn thảo để sửa câu hỏi đang chọn."""
        sel = self.q_tree.selection()
        if not sel: return
        q_id = self.q_tree.item(sel[0])["values"][0]
        QuestionEditorDialog(self.dlg, question_id=q_id, on_save_callback=self.refresh_after_edit)

    def refresh_after_edit(self):
        """Callback sau khi lưu câu hỏi."""
        if self.current_bank_id:
            self.load_questions(self.current_bank_id)
        else:
            self.refresh_all()

    # ============= LOGIC QUẢN LÝ MỚI =============

    def move_questions_ui(self):
        """Di chuyển các câu hỏi đang chọn sang ngân hàng khác."""
        sel = self.q_tree.selection()
        if not sel: return
        q_ids = [self.q_tree.item(s)["values"][0] for s in sel]
        
        dlg = MoveQuestionsDialog(self.dlg, q_ids, on_success=self.refresh_after_edit)
        self.dlg.wait_window(dlg.win)

    def merge_banks_ui(self):
        """Mở dialog gộp ngân hàng."""
        dlg = MergeBanksDialog(self.dlg, on_success=self.refresh_all)
        self.dlg.wait_window(dlg.win)

    def export_bank_db_to_word(self):
        """Xuất ngân hàng đang chọn (hoặc toàn bộ kết quả tìm kiếm) ra Word."""
        if self.current_bank_id:
            bank_id = self.current_bank_id
            sel_item = self.bank_tree.selection()
            bank_name = self.bank_tree.item(sel_item[0])["text"]
        else:
            messagebox.showwarning("Thông báo", "Vui lòng chọn 1 ngân hàng cụ thể để xuất Word.")
            return

        save_path = filedialog.asksaveasfilename(
            title="Lưu file xuất từ CSDL",
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx")],
            initialfile=f"{bank_name}_Export.docx"
        )
        if not save_path: return

        if db.export_bank_to_docx(bank_id, save_path):
            messagebox.showinfo("Thành công", f"Đã xuất ngân hàng ra file Word thành công tại:\n{save_path}")
        else:
            messagebox.showerror("Lỗi", "Không thể xuất file Word. Kiểm tra lại dữ liệu câu hỏi.")

    def open_audit_center(self):
        """Mở trung tâm kiểm tra CSDL."""
        DatabaseAuditDialog(self.dlg)

    def auto_merge_duplicates(self):
        """Tự động tìm và gộp câu trùng (hash giống nhau)."""
        conflicts = db.get_duplicate_conflicts()
        resolutions = {}
        
        if conflicts:
            if not messagebox.askyesno("Phát hiện xung đột", 
                f"Phát hiện {len(conflicts)} nhóm câu trùng nhưng khác mức độ khó.\n"
                f"Bạn có muốn chọn mức độ khó cho các nhóm này trước khi gộp không?"):
                # Nếu không muốn chọn, hệ thống sẽ gộp theo mặc định (giữ mức độ của câu đầu tiên)
                pass
            else:
                dlg = MergeConflictDialog(self.dlg, conflicts)
                self.dlg.wait_window(dlg.win)
                if not dlg.confirmed:
                    return
                resolutions = dlg.results

        if not messagebox.askyesno("Xác nhận", "Hệ thống sẽ tự động tìm và gộp các câu hỏi trùng nội dung.\n\nBạn có muốn thực hiện không?"):
            return
            
        count = db.merge_duplicate_questions(conflict_resolutions=resolutions)
        messagebox.showinfo("Hoàn tất", f"Đã gộp thành công {count} câu hỏi trùng lặp.")
        self.refresh_all()

    def backup_db_ui(self):
        """Sao lưu file CSDL."""
        now = datetime.now().strftime("%Y%m%d_%H%M%S")
        save_path = filedialog.asksaveasfilename(
            title="Chọn vị trí sao lưu CSDL",
            defaultextension=".db",
            filetypes=[("SQLite DB", "*.db")],
            initialfile=f"epu_exam_backup_{now}.db"
        )
        if not save_path: return
        
        if db.backup_database(save_path):
            messagebox.showinfo("Thành công", f"Đã sao lưu CSDL tại:\n{save_path}")
        else:
            messagebox.showerror("Lỗi", "Không thể sao lưu file CSDL.")

    def rename_current_subject(self):
        """Đổi tên môn học đang chọn ở Sidebar."""
        sel = self.bank_tree.selection()
        if not sel: return
        item = self.bank_tree.item(sel[0])
        if item["values"]: # Đây là node Bank, lấy node cha (Subject)
            parent = self.bank_tree.parent(sel[0])
            if not parent: return
            item = self.bank_tree.item(parent)
            subj_node = parent
        else:
            subj_node = sel[0]

        old_name = item["text"]
        if old_name == "Chưa phân loại": return
        
        new_name = simpledialog.askstring("Đổi tên Môn học", f"Nhập tên mới cho môn '{old_name}':", initialvalue=old_name)
        if new_name and new_name.strip() and new_name.strip() != old_name:
            try:
                # Tìm ID môn học (cần logic tìm ID từ name hoặc lưu ID vào node)
                # Để đơn giản, ta tìm ID từ danh sách subjects
                all_subs = db.get_subjects()
                subject = next((s for s in all_subs if s["name"] == old_name), None)
                if subject:
                    db.rename_subject(subject["id"], new_name.strip())
                    self.refresh_all()
            except Exception as e:
                messagebox.showerror("Lỗi", str(e))

    def delete_current_subject(self):
        """Xóa môn học đang chọn."""
        sel = self.bank_tree.selection()
        if not sel: return
        item = self.bank_tree.item(sel[0])
        # Logic tương tự rename để lấy tên môn
        if item["values"]:
            parent = self.bank_tree.parent(sel[0])
            if not parent: return
            item = self.bank_tree.item(parent)
        
        subj_name = item["text"]
        if subj_name == "Chưa phân loại": return

        if messagebox.askyesno("Xác nhận", f"Xóa môn học '{subj_name}'?\nLưu ý: Các ngân hàng thuộc môn này sẽ bị mất liên kết môn học (nhưng không bị xóa câu hỏi)."):
            try:
                all_subs = db.get_subjects()
                subject = next((s for s in all_subs if s["name"] == subj_name), None)
                if subject:
                    db.delete_subject(subject["id"])
                    self.refresh_all()
            except Exception as e:
                messagebox.showerror("Lỗi", str(e))


class MoveQuestionsDialog:
    def __init__(self, parent, question_ids, on_success=None):
        self.win = tk.Toplevel(parent)
        self.q_ids = question_ids
        self.on_success = on_success
        utils.setup_dialog(self.win, width_pct=0.4, height_pct=0.4, title="Di chuyển câu hỏi", parent=parent)
        
        self.banks = db.get_all_banks()
        
        frm = ttk.Frame(self.win, padding=20)
        frm.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frm, text=f"Đang chọn {len(question_ids)} câu hỏi.", font=utils.FONT_BOLD).pack(anchor="w", pady=5)
        ttk.Label(frm, text="Chọn ngân hàng đích:").pack(anchor="w", pady=(10, 2))
        
        bank_labels = [f"{b['file_name']} (Môn: {b.get('subject_name', 'N/A')})" for b in self.banks]
        self.cb_target = ttk.Combobox(frm, values=bank_labels, state="readonly", width=50)
        self.cb_target.pack(fill=tk.X, pady=5)
        if self.banks: self.cb_target.current(0)
        
        btn_frm = ttk.Frame(frm)
        btn_frm.pack(side=tk.BOTTOM, fill=tk.X, pady=10)
        ttk.Button(btn_frm, text="✅ Xác nhận", command=self.confirm, bootstyle="success").pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frm, text="Hủy", command=self.win.destroy).pack(side=tk.RIGHT)

    def confirm(self):
        idx = self.cb_target.current()
        if idx < 0: return
        target_bank_id = self.banks[idx]["id"]
        
        try:
            db.move_questions(self.q_ids, target_bank_id)
            messagebox.showinfo("Thành công", f"Đã di chuyển {len(self.q_ids)} câu hỏi.")
            if self.on_success: self.on_success()
            self.win.destroy()
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))


class MergeConflictDialog:
    def __init__(self, parent, conflicts):
        self.win = tk.Toplevel(parent)
        self.conflicts = conflicts
        self.results = {} # content_hash -> diff_code
        self.confirmed = False
        utils.setup_dialog(self.win, width_pct=0.6, height_pct=0.7, title="Xử lý xung đột mức độ khó", parent=parent)
        
        frm = ttk.Frame(self.win, padding=20)
        frm.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frm, text="Các nhóm câu hỏi sau trùng nội dung nhưng khác mức độ. Hãy chọn mức độ khó mong muốn:", font=utils.FONT_BOLD, wraplength=700).pack(anchor="w", pady=(0, 10))
        
        # Container cho các hàng xung đột
        self.canvas = tk.Canvas(frm, highlightthickness=0)
        self.sb = ttk.Scrollbar(frm, orient="vertical", command=self.canvas.yview)
        self.scroll_frm = ttk.Frame(self.canvas)
        
        self.scroll_frm.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.create_window((0, 0), window=self.scroll_frm, anchor="nw")
        self.canvas.configure(yscrollcommand=self.sb.set)
        
        self.sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.vars = {} # content_hash -> StringVar
        
        for i, conf in enumerate(conflicts):
            f_row = ttk.LabelFrame(self.scroll_frm, text=f"Nhóm {i+1} ({conf['count']} câu)", padding=10)
            f_row.pack(fill=tk.X, pady=5, padx=5)
            
            ttk.Label(f_row, text=conf["stem_preview"] + "...", font=("Arial", 10, "italic"), foreground="gray").pack(anchor="w")
            
            f_opts = ttk.Frame(f_row)
            f_opts.pack(fill=tk.X, pady=5)
            
            v = tk.StringVar(value=conf["diff_codes"][0])
            self.vars[conf["content_hash"]] = v
            
            ttk.Label(f_opts, text="Chọn mức độ: ").pack(side=tk.LEFT)
            for dcode in ["DE", "TB", "KHO"]:
                rb = ttk.Radiobutton(f_opts, text=dcode, variable=v, value=dcode)
                rb.pack(side=tk.LEFT, padx=10)
                if dcode not in conf["diff_codes"]:
                    # Nếu mức độ này không có trong nhóm gốc, làm mờ đi một chút hoặc cứ để đó
                    pass
        
        btn_frm = ttk.Frame(frm)
        btn_frm.pack(side=tk.BOTTOM, fill=tk.X, pady=10)
        ttk.Button(btn_frm, text="✅ Xác nhận & Gộp", command=self.confirm, bootstyle="success").pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frm, text="Hủy", command=self.win.destroy).pack(side=tk.RIGHT)

    def confirm(self):
        for chash, v in self.vars.items():
            self.results[chash] = v.get()
        self.confirmed = True
        self.win.destroy()


class MergeBanksDialog:
    def __init__(self, parent, on_success=None):
        self.win = tk.Toplevel(parent)
        self.on_success = on_success
        utils.setup_dialog(self.win, width_pct=0.5, height_pct=0.6, title="Gộp Ngân hàng câu hỏi", parent=parent)
        
        self.banks = db.get_all_banks()
        
        frm = ttk.Frame(self.win, padding=20)
        frm.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(frm, text="1. Chọn các ngân hàng NGUỒN (sẽ bị xóa sau khi gộp):").pack(anchor="w", pady=2)
        
        # Danh sách chọn nhiều
        self.lb_src = tk.Listbox(frm, selectmode="multiple", height=8)
        self.lb_src.pack(fill=tk.BOTH, expand=True, pady=5)
        for b in self.banks:
            self.lb_src.insert(tk.END, f"{b['file_name']} ({b['total_questions']} câu)")
            
        ttk.Label(frm, text="2. Chọn ngân hàng ĐÍCH (nơi nhận câu hỏi):").pack(anchor="w", pady=(10, 2))
        self.cb_dest = ttk.Combobox(frm, values=[b["file_name"] for b in self.banks], state="readonly")
        self.cb_dest.pack(fill=tk.X, pady=5)
        if self.banks: self.cb_dest.current(0)
        
        btn_frm = ttk.Frame(frm)
        btn_frm.pack(side=tk.BOTTOM, fill=tk.X, pady=10)
        ttk.Button(btn_frm, text="🔗 Tiến hành Gộp", command=self.confirm, bootstyle="danger").pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frm, text="Hủy", command=self.win.destroy).pack(side=tk.RIGHT)

    def confirm(self):
        src_indices = self.lb_src.curselection()
        dest_idx = self.cb_dest.current()
        
        if not src_indices:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn ít nhất 1 ngân hàng nguồn!")
            return
            
        src_ids = [self.banks[i]["id"] for i in src_indices]
        dest_id = self.banks[dest_idx]["id"]
        
        if dest_id in src_ids:
            messagebox.showerror("Lỗi", "Ngân hàng đích không được nằm trong danh sách ngân hàng nguồn!")
            return
            
        if messagebox.askyesno("Xác nhận", f"Bạn có chắc chắn muốn gộp {len(src_ids)} ngân hàng vào '{self.banks[dest_idx]['file_name']}'?\nCác ngân hàng nguồn sẽ bị xóa khỏi hệ thống."):
            try:
                db.merge_banks(src_ids, dest_id)
                messagebox.showinfo("Thành công", "Đã gộp ngân hàng thành công.")
                if self.on_success: self.on_success()
                self.win.destroy()
            except Exception as e:
                messagebox.showerror("Lỗi", str(e))


class QuestionEditorDialog:
    def __init__(self, parent, question_id=None, on_save_callback=None):
        self.parent = parent
        self.question_id = question_id
        self.on_save_callback = on_save_callback
        
        self.dlg = tk.Toplevel(parent)
        self.dlg.title("Soạn thảo câu hỏi trắc nghiệm")
        # Sử dụng setup_dialog để tự động cân đối kích thước theo màn hình (tối đa 90% chiều cao)
        utils.setup_dialog(self.dlg, width_pct=0.8, height_pct=0.9, parent=parent)

        self.media_map = {} # image_name -> {type, blob, latex_code}
        self.tk_images = {} # image_name -> PhotoImage (để tránh GC)

        # --- Dữ liệu ban đầu ---
        self.subjects = db.get_subjects()
        self.banks = []
        
        self.setup_ui()
        
        if self.question_id:
            self.load_question_data()
        else:
            # Mặc định theo config nếu có
            pass

    def setup_ui(self):
        # 0. Thiết lập Scrollable Container (Canvas + Scrollbar)
        self.canvas = tk.Canvas(self.dlg, highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(self.dlg, orient="vertical", command=self.canvas.yview)
        
        # Khung nội dung thực tế
        self.scroll_frm = ttk.Frame(self.canvas, padding=5)
        
        # Bind sự kiện cập nhật vùng cuộn khi khung nội dung thay đổi size
        self.scroll_frm.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        
        # Tạo window bên trong Canvas để chứa Frame
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scroll_frm, anchor="nw")
        
        # Đồng bộ chiều rộng của khung nội dung với Canvas
        self.canvas.bind("<Configure>", self._on_canvas_configure)
        
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        # Ràng buộc cuộn chuột toàn cục cho toàn bộ widget trong dialog
        self.dlg.bind_all("<MouseWheel>", self._on_mousewheel)

        # Frame chính (nay là khung cuộn)
        main_frm = self.scroll_frm
        
        # 1. Header: Môn học & Ngân hàng
        header_frm = ttk.LabelFrame(main_frm, text="Phân loại", padding=5)
        header_frm.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(header_frm, text="Môn học:").grid(row=0, column=0, padx=5, pady=5)
        self.cb_subject = ttk.Combobox(header_frm, values=[s["name"] for s in self.subjects], width=40, state="readonly")
        self.cb_subject.grid(row=0, column=1, padx=5, pady=5)
        self.cb_subject.bind("<<ComboboxSelected>>", self.on_subject_change)
        
        ttk.Label(header_frm, text="Ngân hàng:").grid(row=0, column=2, padx=5, pady=5)
        self.cb_bank = ttk.Combobox(header_frm, width=40, state="readonly")
        self.cb_bank.grid(row=0, column=3, padx=5, pady=5)
        
        ttk.Label(header_frm, text="Độ khó:").grid(row=0, column=4, padx=5, pady=5)
        self.cb_diff = ttk.Combobox(header_frm, values=["DE", "TB", "KHO"], width=10, state="readonly")
        self.cb_diff.set("TB")
        self.cb_diff.grid(row=0, column=5, padx=5, pady=5)

        # 2. Nội dung câu hỏi
        stem_frm = ttk.LabelFrame(main_frm, text="Nội dung câu hỏi", padding=5)
        stem_frm.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # Toolbar câu hỏi
        tb_stem = ttk.Frame(stem_frm)
        tb_stem.pack(fill=tk.X)
        ttk.Button(tb_stem, text="🖼️ Chèn Ảnh", command=lambda: self.insert_image(self.txt_stem)).pack(side=tk.LEFT, padx=2)
        ttk.Button(tb_stem, text="∑ Chèn LaTeX", command=lambda: self.insert_latex(self.txt_stem)).pack(side=tk.LEFT, padx=2)
        ttk.Label(tb_stem, text="(Ctrl+V để dán ảnh)", font=("Arial", 8, "italic")).pack(side=tk.LEFT, padx=10)

        self.txt_stem = tk.Text(stem_frm, height=8, font=("Arial", 11), undo=True)
        self.txt_stem.pack(fill=tk.BOTH, expand=True, pady=5)
        self.txt_stem.bind("<Control-v>", lambda e: self.paste_image(self.txt_stem))

        # 3. Các đáp án
        opt_frm = ttk.LabelFrame(main_frm, text="Các phương án trả lời (Chọn nút tròn cho đáp án đúng)", padding=5)
        opt_frm.pack(fill=tk.BOTH, expand=True, pady=5)
        
        self.opt_texts = []
        self.opt_vars = tk.IntVar(value=0) # 0=A, 1=B, 2=C, 3=D
        
        labels = ["A", "B", "C", "D"]
        for i in range(4):
            f = ttk.Frame(opt_frm)
            f.pack(fill=tk.BOTH, expand=True, pady=2)
            
            # Left: Radio & Toolbar
            f_ctrl = ttk.Frame(f)
            f_ctrl.pack(side=tk.LEFT, fill=tk.Y, padx=5)
            
            ttk.Radiobutton(f_ctrl, text=labels[i], variable=self.opt_vars, value=i).pack(pady=5)
            ttk.Button(f_ctrl, text="🖼️", width=3, command=lambda i=i: self.insert_image(self.opt_texts[i])).pack()
            ttk.Button(f_ctrl, text="∑", width=3, command=lambda i=i: self.insert_latex(self.opt_texts[i])).pack(pady=2)

            t = tk.Text(f, height=2, font=("Arial", 11), undo=True) # Giảm chiều cao từ 3 xuống 2 để tiết kiệm diện tích
            t.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            t.bind("<Control-v>", lambda e, t=t: self.paste_image(t))
            self.opt_texts.append(t)

        # 4. Footer
        btn_frm = ttk.Frame(main_frm)
        btn_frm.pack(fill=tk.X, pady=10)
        
        ttk.Button(btn_frm, text="💾 LƯU CÂU HỎI", command=self.save, width=20, bootstyle="success").pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frm, text="❌ Hủy bỏ", command=self.dlg.destroy).pack(side=tk.RIGHT, padx=5)

    def on_subject_change(self, event=None):
        s_name = self.cb_subject.get()
        subject = next((s for s in self.subjects if s["name"] == s_name), None)
        if not subject: return
        
        self.banks = db.get_banks_by_subject(subject["id"])
        bank_names = [b["file_name"] for b in self.banks]
        bank_names.append("➕ Tạo ngân hàng mới...")
        self.cb_bank.config(values=bank_names)
        if self.banks:
            self.cb_bank.set(self.banks[0]["file_name"])
        else:
            self.cb_bank.set("➕ Tạo ngân hàng mới...")

    # --- Các hàm hỗ trợ cho việc cuộn ---
    def _on_canvas_configure(self, event):
        """Đồng bộ chiều rộng của khung nội dung với chiều rộng canvas."""
        self.canvas.itemconfig(self.canvas_window, width=event.width)

    def _on_mousewheel(self, event):
        """Xử lý cuộn chuột cho Canvas."""
        if self.canvas.winfo_exists():
            # Trên Windows, delta thường là 120 hoặc -120
            self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def set_initial_bank(self, bank_id):
        # Tìm môn học của bank này
        try:
            with db.get_connection() as conn:
                row = conn.execute("SELECT subject_id, file_name FROM question_banks WHERE id = ?", (bank_id,)).fetchone()
                if row:
                    sub_id = row["subject_id"]
                    sub_name = next((s["name"] for s in self.subjects if s["id"] == sub_id), "")
                    self.cb_subject.set(sub_name)
                    self.on_subject_change()
                    self.cb_bank.set(row["file_name"])
        except: pass

    def load_question_data(self):
        try:
            q = db.get_question_with_spans(self.question_id)
            if not q: return
            
            # Set metadata
            # (Ở đây cần code để map bank_id ngược về subject, tương tự set_initial_bank)
            self.set_initial_bank(q["bank_id"])
            self.cb_diff.set(q.get("diff_code", "TB"))
            
            # Fill Stem
            self.render_spans_into_text_widget(self.txt_stem, q.get("stem_media_spans", []))
            # Lưu ý: Hiện tại Editor chưa hỗ trợ stem_extra_media_spans (đoạn phụ) 
            # để đơn giản hóa giao diện. Ta có thể gộp vào stem chính.
            for extra in q.get("stem_extra_media_spans", []):
                self.txt_stem.insert(tk.END, "\n")
                self.render_spans_into_text_widget(self.txt_stem, extra)

            # Fill Options
            self.opt_vars.set(q.get("correct_index", 0))
            for i, opt in enumerate(q.get("options", [])[:4]):
                spans = opt.get("info", {}).get("spans", [])
                self.render_spans_into_text_widget(self.opt_texts[i], spans)
                
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể tải dữ liệu câu hỏi: {e}")

    def render_spans_into_text_widget(self, widget, spans):
        from PIL import Image, ImageTk
        import io
        for sp in spans:
            if sp["type"] == "text":
                widget.insert(tk.END, sp["text"])
            elif sp["type"] == "image":
                try:
                    blob = sp["blob"]
                    pil_img = Image.open(io.BytesIO(blob))
                    # Resize preview
                    w, h = pil_img.size
                    if w > 400: pil_img = pil_img.resize((400, int(h*400/w)), Image.LANCZOS)
                    photo = ImageTk.PhotoImage(pil_img)
                    
                    img_name = f"img_{id(photo)}"
                    self.tk_images[img_name] = photo
                    self.media_map[img_name] = {
                        "type": "latex" if "latex_code" in sp else "image",
                        "blob": blob,
                        "latex_code": sp.get("latex_code")
                    }
                    widget.image_create(tk.END, image=photo, name=img_name)
                except: pass
            elif sp["type"] == "omml":
                # Chuyển công thức sang văn bản để hiển thị trong trình soạn thảo
                summ = utils.get_spans_text_summary([sp])
                widget.insert(tk.END, summ)

    def insert_image(self, widget):
        path = filedialog.askopenfilename(filetypes=[("Image files", "*.png *.jpg *.jpeg *.gif")])
        if not path: return
        try:
            with open(path, "rb") as f:
                blob = f.read()
            self._display_image_in_widget(widget, blob, "image")
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))

    def paste_image(self, widget):
        """Hỗ trợ Ctrl+V ảnh từ clipboard."""
        try:
            from PIL import ImageGrab, Image
            import io
            im = ImageGrab.grabclipboard()
            if isinstance(im, Image.Image):
                buf = io.BytesIO()
                im.save(buf, format="PNG")
                blob = buf.getvalue()
                self._display_image_in_widget(widget, blob, "image")
                return "break" # Ngăn chặn tk.Text xử lý paste text mặc định
        except: pass

    def insert_latex(self, widget):
        latex = simpledialog.askstring("Chèn công thức", "Nhập mã LaTeX (ví dụ: \\frac{a}{b}, x^2 + y^2 = r^2):")
        if not latex: return
        
        blob = utils.render_latex_to_png(latex)
        if blob:
            self._display_image_in_widget(widget, blob, "latex", latex_code=latex)
        else:
            messagebox.showerror("Lỗi", "Không thể render LaTeX. Hãy kiểm tra lại mã.")

    def _display_image_in_widget(self, widget, blob, mtype, latex_code=None):
        from PIL import Image, ImageTk
        import io
        try:
            pil_img = Image.open(io.BytesIO(blob))
            w, h = pil_img.size
            if w > 400: pil_img = pil_img.resize((400, int(h*400/w)), Image.LANCZOS)
            photo = ImageTk.PhotoImage(pil_img)
            
            img_name = f"img_{id(photo)}"
            self.tk_images[img_name] = photo
            self.media_map[img_name] = {"type": mtype, "blob": blob, "latex_code": latex_code}
            widget.image_create(tk.INSERT, image=photo, name=img_name)
        except Exception as e:
            messagebox.showerror("Lỗi hiển thị", str(e))

    def save(self):
        s_name = self.cb_subject.get()
        if not s_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn môn học!")
            return
        
        subject = next((s for s in self.subjects if s["name"] == s_name), None)
        bank_sel = self.cb_bank.get()
        
        if bank_sel == "➕ Tạo ngân hàng mới...":
            new_bank = simpledialog.askstring("Ngân hàng mới", "Nhập tên ngân hàng câu hỏi mới:")
            if not new_bank: return
            bank_id = db.create_virtual_bank(subject["id"], new_bank)
        else:
            bank = next((b for b in self.banks if b["file_name"] == bank_sel), None)
            if not bank:
                messagebox.showerror("Lỗi", "Không tìm thấy ngân hàng!")
                return
            bank_id = bank["id"]

        # Trích xuất dữ liệu
        stem_spans = utils.extract_spans_from_text_widget(self.txt_stem, self.media_map)
        stem_text = "".join(s["text"] for s in stem_spans if s["type"] == "text")
        
        options_data = []
        labels = ["A", "B", "C", "D"]
        texts_for_check = []
        for i, t in enumerate(self.opt_texts):
            spans = utils.extract_spans_from_text_widget(t, self.media_map)
            otext = "".join(s["text"] for s in spans if s["type"] == "text").strip()
            options_data.append({
                "label": labels[i],
                "text": otext,
                "spans": spans,
                "is_correct": (self.opt_vars.get() == i)
            })
            if otext:
                texts_for_check.append(otext.lower())

        # Kiểm tra đáp án giống nhau
        if len(texts_for_check) != len(set(texts_for_check)):
            if not messagebox.askyesno("Cảnh báo trùng lặp", 
                "Phát hiện có các phương án trả lời có nội dung giống nhau.\n"
                "Bạn có chắc chắn muốn lưu câu hỏi này không?"):
                return

        try:
            db.save_question_manually(
                bank_id=bank_id,
                stem_text=stem_text,
                stem_media_spans=stem_spans,
                stem_extra_media_spans=[], # Gộp hết vào stem chính
                diff_code=self.cb_diff.get(),
                options_data=options_data,
                question_id=self.question_id
            )
            messagebox.showinfo("Thành công", "Đã lưu câu hỏi vào cơ sở dữ liệu.")
            if self.on_save_callback: self.on_save_callback()
            self.dlg.destroy()
        except Exception as e:
            messagebox.showerror("Lỗi lưu DB", str(e))


class DatabaseAuditDialog:
    def __init__(self, parent):
        self.parent = parent
        self.win = tk.Toplevel(parent)
        self.win.title("Trung tâm Kiểm tra & Rà soát CSDL")
        utils.setup_dialog(self.win, width_pct=0.9, height_pct=0.8, parent=parent)
        
        self.issues = []
        self.setup_ui()
        self.refresh_data()
        
    def setup_ui(self):
        # Toolbar
        tb = ttk.Frame(self.win, padding=10)
        tb.pack(side=tk.TOP, fill=tk.X)
        
        ttk.Label(tb, text="🪄 Danh sách các vấn đề được phát hiện trong CSDL:", font=("Arial", 12, "bold")).pack(side=tk.LEFT)
        ttk.Button(tb, text="🔄 Làm mới", command=self.refresh_data).pack(side=tk.RIGHT, padx=5)
        
        # Khung chứa Treeview
        frm_tree = ttk.Frame(self.win, padding=10)
        frm_tree.pack(fill=tk.BOTH, expand=True)
        
        cols = ("type", "bank", "qid", "stem", "detail")
        self.tree = ttk.Treeview(frm_tree, columns=cols, show="headings", bootstyle="primary")
        self.tree.heading("type", text="Loại sự cố")
        self.tree.heading("bank", text="Ngân hàng")
        self.tree.heading("qid", text="Câu")
        self.tree.heading("stem", text="Nội dung câu hỏi")
        self.tree.heading("detail", text="Chi tiết")
        
        self.tree.column("type", width=120)
        self.tree.column("bank", width=150)
        self.tree.column("qid", width=50, anchor=tk.CENTER)
        self.tree.column("stem", width=400)
        self.tree.column("detail", width=200)
        
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb = ttk.Scrollbar(frm_tree, orient="vertical", command=self.tree.yview)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.configure(yscrollcommand=vsb.set)
        
        # Footer Buttons
        btn_frm = ttk.Frame(self.win, padding=10)
        btn_frm.pack(side=tk.BOTTOM, fill=tk.X)
        
        ttk.Button(btn_frm, text="✏️ Sửa câu hỏi", command=self.edit_issue_q, width=15, bootstyle="info").pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frm, text="🗑️ Xóa câu hỏi", command=self.delete_issue_q, width=12, bootstyle="danger").pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frm, text="🔗 Gộp câu trùng", command=self.auto_merge_in_audit, width=12, bootstyle="success").pack(side=tk.LEFT, padx=5)
        
        ttk.Button(btn_frm, text="Đóng", command=self.win.destroy, width=10).pack(side=tk.RIGHT, padx=5)

    def refresh_data(self):
        self.tree.delete(*self.tree.get_children())
        try:
            # Tự động cập nhật lại mã băm để đảm bảo tính chính xác cho các câu hỏi công thức
            db.recompute_all_hashes()
            
            self.issues = db.audit_all_questions()
            for i, iss in enumerate(self.issues):
                self.tree.insert("", tk.END, iid=str(i), values=(
                    iss["type"], iss["file_name"], iss["qid_in_file"], iss["stem_text"][:100], iss["detail"]
                ))
            if not self.issues:
                messagebox.showinfo("Thông báo", "Không phát hiện thấy vấn đề nào trong CSDL!")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể rà soát CSDL: {e}")
            
    def edit_issue_q(self):
        sel = self.tree.selection()
        if not sel: return
        idx = int(sel[0])
        iss = self.issues[idx]
        
        # Mở QuestionEditorDialog để sửa câu hỏi đó
        QuestionEditorDialog(self.win, question_id=iss["id"], on_save_callback=self.refresh_data)

    def delete_issue_q(self):
        sel = self.tree.selection()
        if not sel: return
        idx = int(sel[0])
        iss = self.issues[idx]
        
        if messagebox.askyesno("Xác nhận xóa", f"Bạn có chắc muốn xóa câu hỏi:\n'{iss['stem_text'][:60]}...'?"):
            try:
                db.delete_question(iss["id"])
                self.refresh_data()
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể xóa câu hỏi: {e}")
    def auto_merge_in_audit(self):
        """Gộp câu trùng ngay trong Audit Center."""
        conflicts = db.get_duplicate_conflicts()
        resolutions = {}
        
        if conflicts:
            if messagebox.askyesno("Phát hiện xung đột", 
                f"Phát hiện {len(conflicts)} nhóm câu trùng nhưng khác mức độ khó.\n"
                f"Bạn có muốn chọn mức độ khó cho các nhóm này không?"):
                dlg = MergeConflictDialog(self.win, conflicts)
                self.win.wait_window(dlg.win)
                if not dlg.confirmed: return
                resolutions = dlg.results

        if not messagebox.askyesno("Xác nhận", "Thực hiện gộp tất cả câu hỏi trùng lặp?"):
            return
            
        count = db.merge_duplicate_questions(conflict_resolutions=resolutions)
        messagebox.showinfo("Hoàn tất", f"Đã gộp thành công {count} câu hỏi trùng lặp.")
        self.refresh_data()
